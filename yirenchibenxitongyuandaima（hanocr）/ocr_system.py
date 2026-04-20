# === 依賴庫 ===
import sys, os, cv2, time, uuid, json
import numpy as np
from enum import Enum
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

from PyQt6.QtWidgets import (QApplication, QMainWindow, QPushButton, QVBoxLayout, 
                             QHBoxLayout, QWidget, QLabel, QFileDialog, QListWidget, QListWidgetItem,
                             QGraphicsView, QGraphicsScene, QGraphicsItem, QToolBar, QMenu, QSlider, QSplitter,
                             QMessageBox, QProgressDialog, QTextEdit)
from PyQt6.QtGui import (QImage, QPixmap, QColor, QFont, QAction, QPainter, QTextCursor, 
                         QTextCharFormat, QTextDocument)
from PyQt6.QtCore import Qt, pyqtSignal, QRectF, QObject, QRunnable, QThreadPool, QSize

# 引入 NPU 加速引擎
try:
    from rapidocr_openvino import RapidOCR
    HAS_OCR = True
except ImportError:
    HAS_OCR = False
    print("警告: 未安裝 rapidocr_openvino，OCR 功能將不可用。")

# ==========================================
# 1. 數據模型層 
# ==========================================
class PageStatus(Enum):
    WAITING = 0
    PROCESSING = 1
    DONE = 2

class PageModel:
    def __init__(self, file_path, page_index=0, is_pdf=False):
        self.uuid = str(uuid.uuid4())
        self.file_path = file_path
        self.rel_path = ""
        self.page_index = page_index
        self.is_pdf = is_pdf
        self.status = PageStatus.WAITING
        
        self.content_html = "" 
        self.plain_text = ""
        
        self.img_width = 0
        self.img_height = 0

    def to_dict(self):
        return {
            "uuid": self.uuid,
            "rel_path": self.rel_path,
            "page_index": self.page_index,
            "is_pdf": self.is_pdf,
            "content_html": self.content_html, 
            "plain_text": self.plain_text, 
            "status": self.status.value,
            "img_width": self.img_width,
            "img_height": self.img_height
        }

    @staticmethod
    def from_dict(data, project_root):
        try:
            if os.path.isabs(data["rel_path"]): abs_path = data["rel_path"]
            else: abs_path = os.path.normpath(os.path.join(project_root, data["rel_path"]))
        except: abs_path = ""

        page = PageModel(abs_path, data.get("page_index", 0), data.get("is_pdf", False))
        page.uuid = data.get("uuid", str(uuid.uuid4()))
        page.content_html = data.get("content_html", "")
        page.plain_text = data.get("plain_text", "") 
        page.img_width = data.get("img_width", 0)
        page.img_height = data.get("img_height", 0)
        status_val = data.get("status", 0)
        if page.content_html: status_val = 2 
        page.status = PageStatus(status_val)
        return page

# ==========================================
# 2. 視圖層
# ==========================================
class ZoomGraphicsView(QGraphicsView):
    def __init__(self, scene):
        super().__init__(scene)
        self.setRenderHint(QPainter.RenderHint.Antialiasing)
        self.setDragMode(QGraphicsView.DragMode.ScrollHandDrag) 
        self.setMouseTracking(True)
        self.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)

    def resizeEvent(self, event):
        super().resizeEvent(event)
        self.fit_scene_to_view()

    def fit_scene_to_view(self):
        if self.scene() and not self.scene().itemsBoundingRect().isEmpty():
            self.fitInView(self.scene().sceneRect(), Qt.AspectRatioMode.KeepAspectRatio)
    
    def wheelEvent(self, event):
        if event.modifiers() & Qt.KeyboardModifier.ControlModifier:
            zoom_in = event.angleDelta().y() > 0
            scale_factor = 1.1 if zoom_in else 0.9
            self.scale(scale_factor, scale_factor)
            event.accept()
        else:
            super().wheelEvent(event)

# ==========================================
# 3. 資源提供層
# ==========================================
class ImageProvider:
    def __init__(self, cache_size=20):
        self.cache = {} 
        self.cache_order = [] 
        self.max_size = cache_size

    def get_image(self, page_model: PageModel):
        key = f"{page_model.file_path}_{page_model.page_index}"
        if key in self.cache:
            self._update_lru(key)
            return self.cache[key]

        img = None
        if page_model.is_pdf:
            try:
                doc = fitz.open(page_model.file_path)
                pix = doc.load_page(page_model.page_index).get_pixmap(matrix=fitz.Matrix(2, 2))
                img_array = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.h, pix.w, pix.n)
                if pix.n == 3: img = cv2.cvtColor(img_array, cv2.COLOR_RGB2BGR)
                elif pix.n == 4: img = cv2.cvtColor(img_array, cv2.COLOR_RGBA2BGR)
                else: img = cv2.cvtColor(img_array, cv2.COLOR_GRAY2BGR)
                doc.close()
            except Exception as e: print(f"PDF Error: {e}")
        else:
            try: img = cv2.imdecode(np.fromfile(page_model.file_path, dtype=np.uint8), cv2.IMREAD_COLOR)
            except: pass

        if img is not None:
            page_model.img_height, page_model.img_width = img.shape[:2]
            self.cache[key] = img
            self.cache_order.append(key)
            self._evict()
        return img

    def _update_lru(self, key):
        if key in self.cache_order:
            self.cache_order.remove(key)
            self.cache_order.append(key)

    def _evict(self):
        while len(self.cache_order) > self.max_size:
            del self.cache[self.cache_order.pop(0)]

# ==========================================
# 4. 異步 OCR 任務
# ==========================================
class OCRSignals(QObject):
    finished = pyqtSignal(str, str) # uuid, result_text

class OCRRunnable(QRunnable):
    def __init__(self, page_model, model_paths, img_provider):
        super().__init__()
        self.page = page_model
        self.model_paths = model_paths
        self.img_provider = img_provider
        self.signals = OCRSignals()

class OCRRunnable(QRunnable):
    def __init__(self, page_model, model_paths, img_provider):
        super().__init__()
        self.page = page_model
        self.model_paths = model_paths
        self.img_provider = img_provider
        self.signals = OCRSignals()

    def run(self):
        if not HAS_OCR: 
            print("OCR 引擎未加載，跳過識別")
            return

        try:
            img = self.img_provider.get_image(self.page)
            if img is None: return
            
            engine = RapidOCR(
                det_model_path=self.model_paths['det'],
                rec_model_path=self.model_paths['rec'],
                rec_keys_path=self.model_paths['keys'],
                intra_op_num_threads=1
            )

            ocr_result, _ = engine(img, limit_side_len=2000)
            
            final_text = ""
            if ocr_result:
                # === 古籍排序算法 ===
                def get_center(box):
                    pts = np.array(box[0])
                    return np.mean(pts[:, 0]), np.mean(pts[:, 1])

                items = []
                for item in ocr_result:
                    cx, cy = get_center(item)
                    items.append({'cx': cx, 'cy': cy, 'txt': item[1], 'box': item[0]})

                items.sort(key=lambda x: x['cx'], reverse=True)

                threshold = 50 
                if self.page.img_width > 0:
                    threshold = self.page.img_width / 25.0

                columns = []
                if items:
                    current_col = [items[0]]
                    for i in range(1, len(items)):
                        prev = current_col[-1]
                        curr = items[i]
                        if abs(prev['cx'] - curr['cx']) < threshold:
                            current_col.append(curr)
                        else:
                            columns.append(current_col)
                            current_col = [curr]
                    columns.append(current_col)

                sorted_txt_list = []
                for col in columns:
                    col.sort(key=lambda x: x['cy']) 
                    for line in col:
                        sorted_txt_list.append(line['txt'])

                final_text = "".join(sorted_txt_list)

            self.signals.finished.emit(self.page.uuid, final_text)
            
        except Exception as e:
            print(f"OCR Critical Error {self.page.uuid}: {e}")


class EditorScene(QGraphicsScene):
    def __init__(self, parent=None):
        super().__init__(parent)

# ==========================================
# 主窗口 
# ==========================================
class OCRMainWindow(QMainWindow):

    file_exported = pyqtSignal(str)

    def __init__(self):
        super().__init__()
        self.setWindowTitle("一人讀書·古籍 OCR 系統")
        self.resize(1600, 900)
        self.pages = []
        self.current_page = None
        self.img_provider = ImageProvider()
        
        self.is_ocr_running = False # 標記：識別是否運行中
        
        base = os.path.dirname(os.path.abspath(__file__))
        models_dir = os.path.join(base, "models")
        self.model_paths = {
            'det': os.path.join(models_dir, "det_v5.onnx"),
            'rec': os.path.join(models_dir, "rec_v5.onnx"),
            'keys': os.path.join(models_dir, "temp_dict.txt")
        }
        
        self.thread_pool = QThreadPool()
        self.thread_pool.setMaxThreadCount(4)
        self.init_ui()

    def init_ui(self):
        tb = QToolBar(); self.addToolBar(tb)
        
        # === 導入與打開 ===
        action_open = tb.addAction("📂 打開項目 (.json)", self.open_project)
        
        btn_import = QPushButton("➕ 導入新圖片")
        menu_import = QMenu()
        menu_import.addAction("導入 PDF 文件", lambda: self.import_files("pdf"))
        menu_import.addAction("導入 圖片文件夾", lambda: self.import_files("folder"))
        btn_import.setMenu(menu_import)
        tb.addWidget(btn_import)
        
        tb.addSeparator()
        tb.addAction("💾 保存項目", self.save_project)
        tb.addAction("📤 導出結果 (Word+JSON)", self.export_results)
        
        # === 批量控制 ===
        tb.addSeparator()
        # 修改：連接到 toggle 函數，初始不禁用（邏輯內部判斷）
        self.btn_batch = tb.addAction("🚀 開始批量識別", self.toggle_ocr_process)
        self.btn_batch.setEnabled(False) # 初始禁用，直到有文件

        tb.addSeparator()
        tb.addWidget(QLabel("  編輯操作: "))
        
        act_note = QAction("🔴 標記夾註", self)
        act_note.triggered.connect(self.mark_as_note)
        tb.addAction(act_note)
        
        act_norm = QAction("⚫ 取消夾註", self)
        act_norm.triggered.connect(self.mark_as_normal)
        tb.addAction(act_norm)
        
        act_rev = QAction("🔄 順序置換", self)
        act_rev.triggered.connect(self.reverse_selection)
        tb.addAction(act_rev)

        splitter = QSplitter(Qt.Orientation.Horizontal)
        self.setCentralWidget(splitter)

        left = QWidget(); l_lay = QVBoxLayout(left)
        self.nav_list = QListWidget()
        self.nav_list.itemClicked.connect(self.switch_page)
        l_lay.addWidget(self.nav_list)
        
        bot = QWidget(); b_lay = QHBoxLayout(bot)
        self.slider = QSlider(Qt.Orientation.Horizontal); self.slider.setRange(1, 16); self.slider.setValue(4)
        self.lbl_th = QLabel("並發: 4"); self.slider.valueChanged.connect(lambda v: (self.thread_pool.setMaxThreadCount(v), self.lbl_th.setText(f"並發: {v}")))
        b_lay.addWidget(self.lbl_th); b_lay.addWidget(self.slider)
        l_lay.addWidget(bot)
        splitter.addWidget(left)

        center = QWidget(); c_lay = QVBoxLayout(center)
        self.scene = EditorScene()
        self.view = ZoomGraphicsView(self.scene)
        c_lay.addWidget(QLabel("圖片預覽 (Ctrl+滾輪縮放)")) 
        c_lay.addWidget(self.view)
        splitter.addWidget(center)

        right = QWidget(); r_lay = QVBoxLayout(right)
        self.text_editor = QTextEdit()
        self.text_editor.setFont(QFont("SimSun", 14)) 
        self.text_editor.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.text_editor.customContextMenuRequested.connect(self.show_text_context_menu)
        
        r_lay.addWidget(QLabel("文字編輯 (右鍵可用操作)"))
        r_lay.addWidget(self.text_editor)
        splitter.addWidget(right)
        
        splitter.setSizes([200, 800, 600])

    def show_text_context_menu(self, pos):
        menu = self.text_editor.createStandardContextMenu()
        menu.addSeparator()
        menu.addAction("🔴 標記為夾註 (紅字小號)", self.mark_as_note)
        menu.addAction("⚫ 恢復為正文 (黑字大號)", self.mark_as_normal)
        menu.addAction("🔄 順序置換 (倒序)", self.reverse_selection)
        menu.exec(self.text_editor.mapToGlobal(pos))

    # === 編輯功能實現 ===
    def mark_as_note(self):
        cursor = self.text_editor.textCursor()
        if not cursor.hasSelection(): return
        
        fmt = QTextCharFormat()
        fmt.setForeground(QColor("red"))
        fmt.setFontPointSize(9) 
        fmt.setFontFamily("SimSun")
        cursor.mergeCharFormat(fmt)
        self.text_editor.setFocus()

    def mark_as_normal(self):
        cursor = self.text_editor.textCursor()
        if not cursor.hasSelection(): return
        
        fmt = QTextCharFormat()
        fmt.setForeground(QColor("black"))
        fmt.setFontPointSize(14)
        fmt.setFontFamily("SimSun")
        cursor.mergeCharFormat(fmt)
        self.text_editor.setFocus()

    def reverse_selection(self):
        cursor = self.text_editor.textCursor()
        if not cursor.hasSelection(): return
        
        txt = cursor.selectedText()
        rev_txt = txt[::-1]
        cursor.insertText(rev_txt)

    # === 項目管理功能 (新增：打開項目) ===
    def open_project(self):
        # 1. 詢問是否保存當前工作
        if self.pages:
            reply = QMessageBox.question(self, "提示", "是否先保存當前項目？\n打開新項目將覆蓋當前列表。", 
                                       QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No | QMessageBox.StandardButton.Cancel)
            if reply == QMessageBox.StandardButton.Yes:
                self.save_project()
            elif reply == QMessageBox.StandardButton.Cancel:
                return

        # 2. 選擇文件
        path, _ = QFileDialog.getOpenFileName(self, "打開項目", "", "JSON (*.json)")
        if not path: return

        try:
            with open(path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            self.pages = []
            self.nav_list.clear()
            self.scene.clear()
            self.text_editor.clear()
            self.current_page = None
            
            project_root = os.path.dirname(path)
            
            # 3. 重建數據與UI
            for item_data in data:
                page = PageModel.from_dict(item_data, project_root)
                self.pages.append(page)
                
                if page.is_pdf:
                    label = f"{os.path.basename(page.file_path)} - {page.page_index + 1}"
                else:
                    label = os.path.basename(page.file_path)
                
                # 恢復狀態圖標
                if page.status == PageStatus.DONE:
                    label += " ✅"
                elif page.status == PageStatus.PROCESSING:
                    # 如果上次意外關閉，重置為 WAITING
                    page.status = PageStatus.WAITING
                
                self.add_nav_item(page, label)

            # 4. 恢復按鈕狀態
            self.btn_batch.setEnabled(True)
            self.btn_batch.setText("🚀 繼續批量識別") 
            self.is_ocr_running = False
            
            QMessageBox.information(self, "成功", f"已加載 {len(self.pages)} 個頁面。")

        except Exception as e:
            QMessageBox.critical(self, "錯誤", f"無法打開項目文件: {str(e)}")

    def import_files(self, mode):
        # 導入時清空舊列表
        self.pages = []
        self.nav_list.clear()
        self.scene.clear()
        self.text_editor.clear()
        
        if mode == "pdf":
            path, _ = QFileDialog.getOpenFileName(self, "選擇 PDF", "", "PDF (*.pdf)")
            if not path: return
            try:
                doc = fitz.open(path)
                for i in range(len(doc)):
                    p = PageModel(path, i, True)
                    self.pages.append(p)
                    self.add_nav_item(p, f"{os.path.basename(path)} - {i+1}")
                doc.close()
            except Exception as e: QMessageBox.critical(self, "Err", str(e))
        elif mode == "folder":
            path = QFileDialog.getExistingDirectory(self, "選擇文件夾")
            if not path: return
            files = sorted([f for f in os.listdir(path) if f.lower().endswith(('jpg','png','jpeg'))])
            for f in files:
                p = PageModel(os.path.join(path, f), 0, False)
                self.pages.append(p)
                self.add_nav_item(p, f)
        
        self.btn_batch.setEnabled(True)
        self.btn_batch.setText("🚀 開始批量識別")

    def add_nav_item(self, page, label):
        item = QListWidgetItem(label)
        item.setData(Qt.ItemDataRole.UserRole, page)
        self.nav_list.addItem(item)

    # === 批量 OCR 控制 (新增：開始/停止) ===
    def toggle_ocr_process(self):
        if not self.is_ocr_running:
            self.start_batch_ocr()
        else:
            self.stop_batch_ocr()

    def start_batch_ocr(self):
        # 1. 檢查任務
        has_tasks = False
        for i in range(self.nav_list.count()):
            page = self.nav_list.item(i).data(Qt.ItemDataRole.UserRole)
            if page.status == PageStatus.WAITING:
                has_tasks = True
                break
        
        if not has_tasks:
            QMessageBox.information(self, "提示", "所有頁面均已識別完成！")
            return

        # 2. 設置狀態
        self.is_ocr_running = True
        self.btn_batch.setText("🛑 停止識別") 

        # 3. 提交任務
        for i in range(self.nav_list.count()):
            item = self.nav_list.item(i)
            page = item.data(Qt.ItemDataRole.UserRole)
            
            if page.status == PageStatus.WAITING:
                page.status = PageStatus.PROCESSING
                item.setText(item.text() + " ⏳")
                
                worker = OCRRunnable(page, self.model_paths, self.img_provider)
                worker.signals.finished.connect(self.on_page_ocr_finished)
                self.thread_pool.start(worker)

    def stop_batch_ocr(self):
        # 1. 清空等待隊列
        self.thread_pool.clear() 

        # 2. 重置未開始的任務狀態
        for i in range(self.nav_list.count()):
            item = self.nav_list.item(i)
            page = item.data(Qt.ItemDataRole.UserRole)
            
            # 如果標記為 PROCESSING 但實際上沒出結果，說明被攔截了，重置為 WAITING
            if page.status == PageStatus.PROCESSING and not page.content_html:
                page.status = PageStatus.WAITING
                txt = item.text().replace(" ⏳", "")
                item.setText(txt)

        # 3. 恢復 UI
        self.is_ocr_running = False
        self.btn_batch.setText("🚀 繼續批量識別")

    def on_page_ocr_finished(self, uid, text_content):
        page = next((p for p in self.pages if p.uuid == uid), None)
        if not page: return
        
        base_html = f"""<!DOCTYPE HTML PUBLIC "-//W3C//DTD HTML 4.0//EN" "http://www.w3.org/TR/REC-html40/strict.dtd">
<html><head><meta name="qrichtext" content="1" /><style type="text/css">
p, li {{ white-space: pre-wrap; }}
</style></head><body style=" font-family:'SimSun'; font-size:14pt; font-weight:400; font-style:normal;">
<p style=" margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;">{text_content}</p>
</body></html>"""
        
        page.content_html = base_html
        page.plain_text = text_content
        page.status = PageStatus.DONE
        
        for i in range(self.nav_list.count()):
            it = self.nav_list.item(i)
            if it.data(Qt.ItemDataRole.UserRole).uuid == uid:
                it.setText(it.text().replace(" ⏳", "") + " ✅")
        
        if self.current_page and self.current_page.uuid == uid:
            self.load_page_into_editor(page)

        # 檢查是否全部完成 (線程池空且是運行狀態)
        if self.thread_pool.activeThreadCount() == 0 and self.is_ocr_running:
            remaining = any(p.status == PageStatus.WAITING for p in self.pages)
            if not remaining:
                self.is_ocr_running = False
                self.btn_batch.setText("🚀 重新批量識別")
                QMessageBox.information(self, "完成", "批量識別已全部完成！")

    def switch_page(self, item):
        new_page = item.data(Qt.ItemDataRole.UserRole)
        if self.current_page: self.save_current_editor_to_model()
        self.load_page_into_editor(new_page)

    def save_current_editor_to_model(self):
        if not self.current_page: return
        self.current_page.content_html = self.text_editor.toHtml()
        self.current_page.plain_text = self.text_editor.toPlainText()

    def load_page_into_editor(self, page):
        self.current_page = page
        self.scene.clear()
        
        img = self.img_provider.get_image(page)
        if img is not None:
            h, w = img.shape[:2]
            qimg = QImage(img.data, w, h, 3*w, QImage.Format.Format_RGB888)
            self.scene.addPixmap(QPixmap.fromImage(qimg))
            self.scene.setSceneRect(0, 0, w, h)
            self.view.fit_scene_to_view()
        
        if page.content_html:
            self.text_editor.setHtml(page.content_html)
        else:
            self.text_editor.clear()
            self.text_editor.setFont(QFont("SimSun", 14))

    def save_project(self):
        path, _ = QFileDialog.getSaveFileName(self, "保存項目", "", "JSON (*.json)")
        if not path: return
        self.save_current_editor_to_model()
        
        data = []
        root = os.path.dirname(path)
        for p in self.pages:
            try: p.rel_path = os.path.relpath(p.file_path, root)
            except: p.rel_path = p.file_path
            data.append(p.to_dict())
        with open(path, 'w', encoding='utf-8') as f: json.dump(data, f, ensure_ascii=False, indent=2)

    def export_results(self):
        self.save_current_editor_to_model() 
        
        word_path, _ = QFileDialog.getSaveFileName(self, "導出 Word", "", "Word (*.docx)")
        if not word_path: return
        
        base_name = os.path.splitext(word_path)[0]
        json_path = base_name + "_info.json"
        
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'SimSun'
        style.font.size = Pt(14)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

        export_data = []
        current_global_index = 0
        
        temp_doc = QTextDocument()
        
        for p in self.pages:
            if p.is_pdf:
                fname = f"{os.path.basename(p.file_path)}-{p.page_index+1}"
            else:
                fname = os.path.basename(p.file_path)

            p_text = p.plain_text
            start_idx = current_global_index
            end_idx = start_idx + len(p_text)
            current_global_index = end_idx
            
            export_data.append({
                "file_name": fname,
                "content_html": p.content_html,
                "word_index": {
                    "start": start_idx,
                    "end": end_idx
                }
            })
            
            temp_doc.setHtml(p.content_html)
            paragraph = doc.add_paragraph()
            
            block = temp_doc.begin() 
            while block.isValid():
                iterator = block.begin() 
                while not iterator.atEnd():
                    fragment = iterator.fragment() 
                    if fragment.isValid():
                        text = fragment.text()
                        if text:
                            fmt = fragment.charFormat()
                            text = text.replace('\x00', '')
                            run = paragraph.add_run(text)
                            
                            run.font.name = 'SimSun'
                            run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
                            
                            color = fmt.foreground().color()
                            if color.red() == 255 and color.green() == 0 and color.blue() == 0:
                                run.font.color.rgb = RGBColor(255, 0, 0)
                                run.font.size = Pt(9)
                            else:
                                run.font.color.rgb = RGBColor(0, 0, 0)
                                run.font.size = Pt(14)
                    
                    iterator += 1 
                block = block.next() 
        
        try:
            doc.save(word_path)
            with open(json_path, 'w', encoding='utf-8') as f:
                json.dump(export_data, f, ensure_ascii=False, indent=2)
            QMessageBox.information(self, "成功", f"導出完成！\nWord: {word_path}\nJSON: {json_path}")
            self.file_exported.emit(word_path) 

        except Exception as e:
            QMessageBox.critical(self, "錯誤", f"導出失敗: {str(e)}")

if __name__ == "__main__":
    app = QApplication(sys.argv)
    w = OCRMainWindow() 
    w.show()
    sys.exit(app.exec())