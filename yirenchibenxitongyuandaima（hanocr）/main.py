import sys
import os
import fitz  # PyMuPDF: 用於讀取和解析 PDF 文件的 XML 結構。
import ctypes
import re
import gzip
import json
import pandas as pd
from collections import Counter, defaultdict
import difflib # 用於計算文本相似度
import copy    # 用於深拷貝數據實現撤銷

# [文檔處理模塊]
# python-docx: 用於讀取和解析 .docx 文件的 XML 結構。
# 此庫主要負責「讀取」底本與校本的文本內容。
# Document: 文檔對象模型入口。
# qn (Qualified Name): 用於處理 Open XML (OOXML) 的命名空間，這是精確控制 Word 字體（如設置中文字體）的關鍵。
# DocxColor: 用於在不調用 COM 的情況下，修改文檔底層 XML 的顏色屬性。
from docx import Document
from docx.oxml.ns import qn 
from docx.shared import RGBColor as DocxColor

# [核心算法模塊]
# Biopython: 生物信息學計算庫。
# 這裡並非用於生物學，而是利用其高度優化的序列比對算法 (Sequence Alignment)。
# pairwise2: 實現了 Needleman-Wunsch 全局比對算法與 Smith-Waterman 局部比對算法。（低速，為了實現大文本校對，更換爲同樣運算邏輯，但是更快的用c寫的Align）
# 本系統利用它來計算底本與校本之間的「最優編輯距離」，這是自動校勘的數學核心。 
from Bio import Align

# [圖形用戶界面 (GUI) 模塊]
# PyQt6: 基於 Qt 框架的 Python 綁定，採用信號與槽 (Signal & Slot) 機制實現事件驅動編程。
# 
# QtWidgets: 包含所有可視化控件。
# - QMainWindow, QWidget: 窗口基類。
# - QVBoxLayout, QHBoxLayout: 佈局管理器，負責響應式排版。
# - QTextEdit: 用於顯示富文本 (Rich Text) 的預覽區。
# - QTableWidget: 用於展示結構化的異文數據表。
# - QUndoStack, QUndoCommand: 實現「撤銷/重做」功能的命令模式 (Command Pattern) 基礎設施。
from PyQt6.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, 
                             QHBoxLayout, QTextEdit, QLabel, QPushButton, 
                             QTableWidget, QTableWidgetItem, QHeaderView, QMessageBox,
                             QTabWidget, QFileDialog, QComboBox, QListWidget, QGroupBox,
                             QProgressBar, QSplitter, QFrame, QCheckBox, QScrollArea, QLineEdit, QMenu, QDialog, QFormLayout, QListWidgetItem, QSlider, QRadioButton, QButtonGroup, QGraphicsView, QGraphicsScene, QDockWidget, QToolBar, 
                             QVBoxLayout, QWidget, QPushButton, QLabel, QSplitter, QSizePolicy)

# QtCore: 核心非 GUI 功能。
# - QThread: 實現多線程 (Multithreading)，將耗時的校勘運算移出主線程，防止界面凍結。
# - pyqtSignal: 線程間通信機制。
# QtGui: 繪圖與字體處理。
# - QRawFont: 提供對底層字體文件的直接訪問，用於判斷字體是否支持某個特定的 Unicode 擴展區漢字（解決生僻字顯示問題的核心類）。
# - QFontDatabase: 管理應用程序加載的自定義字體文件。
from PyQt6.QtCore import Qt, QThread, pyqtSignal, QSize, QEvent, QCoreApplication, Qt, pyqtSignal, QSize, QTimer
from PyQt6.QtGui import QFont, QColor, QTextCharFormat, QColor as RGBColor, QIcon, QFontDatabase, QScreen, QFontMetrics, QRawFont, QUndoCommand, QUndoStack, QKeySequence, QAction, QPixmap, QPainter, QImage

from PyQt6.QtWidgets import QStyledItemDelegate
from PyQt6.QtGui import QTextDocument, QAbstractTextDocumentLayout

import numpy as np  # 用於矩陣運算
from PyQt6.QtGui import QPainter, QImage, QColor, qRgb

class AutoAdaptDelegate(QStyledItemDelegate):
    """
    [新增組件] 自適應換行代理
    功能：接管 QTreeWidget 的繪製和尺寸計算，解決文字重疊，實現無限高度和舒適留白。
    """
    def __init__(self, parent=None, vertical_padding=20, horizontal_padding=10):
        super().__init__(parent)
        self.v_padding = vertical_padding  # 上下總留白
        self.h_padding = horizontal_padding # 左右總留白

    def paint(self, painter, option, index):
        # 覆蓋默認繪製：讓內容區域內縮，從而產生「留白」效果
        # 通過調整 option.rect，讓文字不要畫在邊框上
        option.rect.adjust(self.h_padding // 2, self.v_padding // 2, 
                           -(self.h_padding // 2), -(self.v_padding // 2))
        super().paint(painter, option, index)

    def sizeHint(self, option, index):
        # 核心邏輯：計算無限撐大的高度
        text = index.data(Qt.ItemDataRole.DisplayRole)
        if not text:
            return super().sizeHint(option, index)

        # 1. 獲取當前列的實際寬度
        # 注意：option.rect.width() 在某些初始化時刻可能為0，需要容錯
        column_width = option.rect.width()
        if column_width <= 0:
            # 嘗試從父控件獲取列寬 (這裡假設 parent 是 treeWidget)
            tree = self.parent()
            if tree and hasattr(tree, "header"):
                column_width = tree.header().sectionSize(index.column())
            else:
                return super().sizeHint(option, index)

        # 2. 減去左右留白，得到文字可用的淨寬度
        text_width = column_width - self.h_padding

        # 3. 使用 QTextDocument 模擬排版，計算精確的換行高度
        document = QTextDocument(text)
        document.setDefaultFont(option.font)
        document.setTextWidth(text_width) # 設定約束寬度

        # 4. 返回高度：文字內容高度 + 上下留白
        height = int(document.size().height() + self.v_padding)
        
        return QSize(column_width, height)

# --- 導入可視化庫 ---
# [模塊設計意圖]：可選依賴 (Optional Dependency)。
# 為了增強系統的魯棒性，將數據可視化模塊設為可選。
# 即使用戶環境缺失 matplotlib，系統的核心校勘功能仍能正常運行，僅圖表功能失效。
try:
    import matplotlib
    matplotlib.use('QtAgg') 
    from matplotlib.backends.backend_qtagg import FigureCanvasQTAgg as FigureCanvas
    from matplotlib.figure import Figure
    import matplotlib.pyplot as plt
    import numpy as np
    # 設置中文字體 (優先嘗試常見中文字體)
    plt.rcParams['font.sans-serif'] = ['SimSun', 'KaiTi', 'Microsoft YaHei', 'SimHei'] 
    plt.rcParams['axes.unicode_minus'] = False
    HAS_MATPLOTLIB = True
except ImportError:
    HAS_MATPLOTLIB = False

# --- 導入 Word 自動化庫 ---
# [模塊說明]：負責複雜格式的 Word 文檔導出。
# 使用 Windows COM (Component Object Model) 接口，直接調用底層 Word 應用程序。
# 此技術選型是為了繞過 python-docx 在處理 Unicode 擴展區漢字字體回退（Font Fallback）時的限制。
try:
    import win32com.client as win32
    # 專門導入 dynamic 模塊，這是關鍵
    import win32com.client.dynamic 
    import pythoncom
    HAS_PYWIN32 = True
except ImportError:
    HAS_PYWIN32 = False

# [新增] MDict 依賴庫
from readmdict import MDX, MDD  # 需 pip install readmdict
from PyQt6.QtWebEngineCore import QWebEngineUrlSchemeHandler, QWebEngineUrlScheme, QWebEngineProfile, QWebEngineUrlRequestJob
from PyQt6.QtWebEngineWidgets import QWebEngineView
from PyQt6.QtCore import QBuffer, QIODevice, QUrl

from ocr_system import OCRMainWindow

# ==========================================
# 0. 路徑處理函數 (兼容打包與IDE運行)
# ==========================================
def get_resource_path(relative_path):
    """
    [功能說明]：資源路徑標準化函數。
    [設計意圖]：解決 Python 腳本在開發環境（IDE）與打包環境（PyInstaller EXE）下
    臨時文件目錄結構不一致的問題，確保異體字字典和字體文件能被正確尋址。
    """
    # 1. 如果是打包後的 EXE 環境
    if hasattr(sys, '_MEIPASS'):
        return os.path.join(sys._MEIPASS, relative_path)
    
    # 2. 如果是正常 Python 腳本運行環境
    # 使用 __file__ 獲取當前腳本的絕對路徑目錄
    current_dir = os.path.dirname(os.path.abspath(__file__))
    return os.path.join(current_dir, relative_path)

# ==========================================
# 1. 全局樣式表 (QSS) - 新中式風格
# ==========================================
# [設計意圖]：UI/UX 設計採用「新中式」美學。
# 配色方案參考傳統中國色（宣紙白、徽墨黑、硃砂紅、竹青等），
# 旨在為古籍整理工作者提供符合人文審美且低視覺疲勞的操作環境。
STYLESHEET = """
/* 全局背景：宣紙白 */
QMainWindow, QWidget {
    background-color: #FAF9F6; 
    color: #2B2B2B; /* 徽墨黑 */
    font-family: "SimSun", "Songti SC", serif;
}

/* 標題字體：楷體 */
QGroupBox::title {
    font-family: "KaiTi", "BiauKai", serif;
    font-size: 16px;
    font-weight: bold;
    color: #5D4037; /* 深褐 */
    subcontrol-origin: margin;
    subcontrol-position: top left;
    padding: 0 5px;
    left: 10px;
}

/* GroupBox 樣式：去框，留底線，雅致 */
QGroupBox {
    font-size: 14px;
    border: 1px solid #E0E0E0;
    border-radius: 6px;
    margin-top: 15px; /* 避開標題 */
    background-color: #FFFFFF; /* 內容區純白 */
}

/* 按鈕通用樣式：扁平、圓角、雅致 */
QPushButton {
    background-color: #ECEAE4; /* 米灰 */
    border: 1px solid #D1CFC8;
    border-radius: 4px;
    padding: 6px 15px;
    font-family: "KaiTi", serif;
    font-size: 14px;
    color: #333;
}
QPushButton:hover {
    background-color: #E0DED7;
    border-color: #B0B0B0;
}
QPushButton:pressed {
    background-color: #D6D4CD;
}

/* 主操作按鈕（開始校勘）：朱砂紅 */
QPushButton#primary_btn {
    background-color: #B74639; /* 朱砂 */
    color: #FFFFFF;
    border: none;
    font-weight: bold;
    font-size: 16px;
}
QPushButton#primary_btn:hover {
    background-color: #C85446;
}
QPushButton#primary_btn:disabled {
    background-color: #DBCBCB;
    color: #F0F0F0;
}

/* 導出按鈕：竹青、靛藍 */
QPushButton#export_excel { color: #2E5C8A; font-weight: bold; border: 1px solid #2E5C8A; background-color: transparent; }
QPushButton#export_excel:hover { background-color: #EBF4FA; }

QPushButton#export_word_anno { color: #5C7A62; font-weight: bold; border: 1px solid #5C7A62; background-color: transparent; }
QPushButton#export_word_anno:hover { background-color: #F0F7F2; }

QPushButton#export_word_note { color: #B74639; font-weight: bold; border: 1px solid #B74639; background-color: transparent; }
QPushButton#export_word_note:hover { background-color: #FDF3F2; }

/* 列表和表格：紙張感 */
QListWidget, QTableWidget, QTextEdit {
    background-color: #FFFEFA;
    border: 1px solid #D0D0D0;
    selection-background-color: #E0E8F0; /* 淡藍選中 */
    selection-color: #000;
}

/* ==============================================
   一級 Tab (主導航) - 莫蘭迪·陶土紅【印章風格】
   ============================================== */
QTabWidget::pane {
    border: 1px solid #D6D2CE; /* 邊框淡暖灰 */
    background: #FFFFFF;
    border-radius: 4px;
    top: -1px; 
}

QTabBar::tab {
    background: #F2F0EB; /* 未選中：暖米灰 */
    border: 1px solid #E0DCD8;
    color: #9E9490;      /* 未選中文字：淺暖灰 */
    padding: 8px 25px;   /* 寬敞大氣 */
    margin-right: 2px;
    border-top-left-radius: 4px;
    border-top-right-radius: 4px;
    font-family: "KaiTi", serif;
    font-size: 15px;
}

QTabBar::tab:selected {
    /* 背景全填充：陶土紅 */
    background-color: #A86C65; 
    
    /* 文字反白，形成強對比 */
    color: #FFFFFF; 
    
    /* 邊框同色，底部封死，形成獨立印章感 */
    border: 1px solid #A86C65; 
    border-bottom: 1px solid #A86C65; 
    
    font-weight: bold;
}

QTabBar::tab:hover {
    background-color: #FCFBFA;
    color: #8C7A74;
}

/* ==============================================
   二級 Tab (內層導航) - 莫蘭迪·竹青綠【印章風格】
   使用 ID 選擇器 #SubTabs 進行覆蓋
   ============================================== */
QTabWidget#SubTabs::pane {
    border: 1px solid #CFD9D6; /* 邊框改為冷調淡綠灰 */
}

QTabWidget#SubTabs QTabBar::tab {
    background: #F4F6F5;  /* 未選中：冷調極淡綠 */
    border: 1px solid #E1E6E4;
    color: #99A69E;       /* 未選中文字：淺綠灰 */
    font-size: 14px;      /* 字體稍小一號 */
    padding: 6px 15px;    /* 間距稍緊湊 */
    font-family: "SimSun", serif; 
}

QTabWidget#SubTabs QTabBar::tab:selected {
    /* 背景全填充：竹青綠 */
    background-color: #7D9689; 
    
    /* 文字反白 */
    color: #FFFFFF;
    
    /* 邊框同色 */
    border: 1px solid #7D9689;
    border-bottom: 1px solid #7D9689;
    
    font-weight: bold;
}

QTabWidget#SubTabs QTabBar::tab:hover {
    background-color: #F9FAFA;
    color: #6E8279;
}

/* 進度條：纖細 */
QProgressBar {
    border: none;
    background-color: #E0E0E0;
    border-radius: 2px;
    height: 6px;
    text-align: center;
}
QProgressBar::chunk {
    background-color: #B74639;
    border-radius: 2px;
}

/* 滾動條美化 (可選) */
QScrollBar:vertical {
    border: none;
    background: #F0F0F0;
    width: 8px;
    margin: 0px;
}
QScrollBar::handle:vertical {
    background: #C0C0C0;
    min-height: 20px;
    border-radius: 4px;
}
"""

# ==========================================
#  撤銷/恢復命令類 (Undo/Redo Commands)
# ==========================================
# [架構說明]：採用命令模式 (Command Pattern) 實現用戶操作的狀態管理。
# 每個操作（修改類型、編輯內容、刪除）都被封裝為一個獨立的對象，
# 包含 `redo`（執行）和 `undo`（回滾）方法，由 QUndoStack 統一調度。
class NoteTypeChangeCommand(QUndoCommand):
    """
    [命令類]：處理校勘記類型的變更操作 (如 訛 -> 脫)。
    [學術邏輯]：當類型變更時（例如從『訛』變為『衍』），校勘記的術語模板（Template）也會隨之改變。
    因此在 `redo` 時需要重新生成符合對應類型的校勘記，而在 `undo` 時恢復舊數據。
    """
    def __init__(self, window, row_idx, new_type):
        super().__init__()
        self.window = window
        self.row_idx = row_idx
        self.new_type = new_type
        
        self.record = self.window.note_records[row_idx]
        self.old_type = self.record['type']
        self.old_content = self.record['content']
        
        # 計算新內容：如果不是手動修改過，則自動重新生成
        if not self.record['is_manual']:
            self.new_content = self.window.generate_note_content(
                new_type, 
                self.record['base_char'], 
                self.record['wit_char'], 
                self.record['wit_names'], 
                self.record['prev_char']
            )
        else:
            self.new_content = self.old_content 

        self.setText(f"修改類型: {self.old_type}->{self.new_type}")

    def redo(self):
        # 1. 更新數據
        self.record['type'] = self.new_type
        self.record['content'] = self.new_content
        # 2. 更新界面 (表格 + 預覽)
        self.window.refresh_ui_row(self.row_idx)

    def undo(self):
        # 1. 還原數據
        self.record['type'] = self.old_type
        self.record['content'] = self.old_content
        # 2. 還原界面
        self.window.refresh_ui_row(self.row_idx)

class NoteContentEditCommand(QUndoCommand):
    """
    [命令類]：處理校勘記內容的手動編輯。
    [邏輯]：記錄編輯前後的文本狀態，並標記 `is_manual` 屬性，
    防止自動模板生成邏輯覆蓋用戶的人工校訂成果。
    """
    def __init__(self, window, row_idx, new_text):
        super().__init__()
        self.window = window
        self.row_idx = row_idx
        self.new_text = new_text
        
        self.record = self.window.note_records[row_idx]
        self.old_text = self.record['content']
        self.old_manual_state = self.record['is_manual']
        
        self.setText("編輯校勘記內容")

    def redo(self):
        self.record['content'] = self.new_text
        self.record['is_manual'] = True # 標記為手動修改
        self.window.refresh_preview_only() # 只刷新預覽，不刷新表格防止焦點丟失

    def undo(self):
        self.record['content'] = self.old_text
        self.record['is_manual'] = self.old_manual_state
        self.window.refresh_ui_row(self.row_idx)

class NoteDeleteCommand(QUndoCommand):
    """
    [命令類]：處理校勘記的刪除操作。
    [邏輯]：物理刪除數據行，Undo 時將數據對象重新插入原索引位置。
    """
    def __init__(self, window, row_idx):
        super().__init__()
        self.window = window
        self.row_idx = row_idx
        self.record_to_del = self.window.note_records[row_idx]
        self.setText("刪除校勘記")

    def redo(self):
        self.window.delete_note_data(self.row_idx)

    def undo(self):
        self.window.insert_note_data(self.row_idx, self.record_to_del)

# ==========================================
# 2. 核心算法引擎 
# ==========================================
class CollationEngine:
    """
    [類別說明]：校勘核心引擎。
    [職責]：
    1. 管理異體字數據庫 (Variant Database)。
    2. 提供字形比較邏輯 (Variant Relation Check)。
    3. 執行基於生物信息學算法的文本對齊 (Global Alignment)。
    """
    def __init__(self, variant_db_path=None):
        self.variant_db = {}
        if variant_db_path and os.path.exists(variant_db_path):
            try:
                with gzip.open(variant_db_path, 'rt', encoding='utf-8') as f:
                    self.variant_db = json.load(f)
            except Exception as e:
                print(f"異體字數據庫加載失敗: {e}")
                # === 【插入點 1：初始化自定義字典】 ===
        self.custom_variants = {} # 用於存 A=B
        self.custom_excludes = set() # 用於存 A!=B
        # === 【插入以下代碼：初始化新引擎】 ===
        self.aligner = Align.PairwiseAligner()
        self.aligner.mode = 'global'
        # 參數設置：匹配=1, 不匹配=-1, 開口=-5, 延長=-0.5 (與您原來的參數一致)
        self.aligner.match_score = 1
        self.aligner.mismatch_score = -1
        self.aligner.open_gap_score = -5
        self.aligner.extend_gap_score = -0.5
        # === 【插入點 2：新增一個方法，請插在 __init__ 方法結束後，is_variant_relation 之前】 ===
    def update_custom_dict(self, variants, excludes):
        """更新內存中的自定義規則"""
        self.custom_variants = variants
        self.custom_excludes = excludes

    def is_variant_relation(self, char1, char2):
        """
        [算法]：異體字判定算法。
        [學術定義]：判斷兩個字是否隸屬於同一個「正體字群組 (Standard Group)」。
        [邏輯]：若 char1 與 char2 在數據庫中共享至少一個共同的正體字標準，則視為異體關係。
        """
        if char1 == char2: return False
        # 1. 優先檢查【強制訛誤/反向屏蔽】
        if (char1, char2) in self.custom_excludes or (char2, char1) in self.custom_excludes:
            return False

        # 2. 檢查【臨時異體】(雙向檢查)
        if self.custom_variants.get(char1) == char2 or self.custom_variants.get(char2) == char1:
            return True
        
        d1 = self.variant_db.get(char1)
        d2 = self.variant_db.get(char2)
        if not d1 or not d2: return False
        
        # 核心判斷：是否擁有共同的正體字群組
        common_standards = set(d1.get('standards', [])) & set(d2.get('standards', []))
        return len(common_standards) > 0

    def get_char_attr(self, char1, char2):
        """獲取一對字的屬性和情況說明"""
        d1 = self.variant_db.get(char1, {})
        d2 = self.variant_db.get(char2, {})
        is_s1 = d1.get('is_standard', False)
        is_s2 = d2.get('is_standard', False)
        
        # 判定「二字均可作正體」
        if is_s1 and is_s2:
            return "", "", "二字均可作正體！請回查原文判斷！" # 屬性欄留空，用高亮顯示
        if not is_s1 and not is_s2:
            return "異體", "異體", "均異"
        
        attr1 = "正體" if is_s1 else "異體"
        attr2 = "正體" if is_s2 else "異體"
        return attr1, attr2, "存在正異對照"

    def align_paragraph(self, base_text, wit_text):
        if not base_text or not wit_text: return []
        # 參數設計意圖: match=1, mismatch=-1, gap_open=-5, gap_extend=-0.5
    # 1. 調用新引擎 (不會崩潰)
        alignments = self.aligner.align(base_text, wit_text)
        alignment = alignments[0]

        # 2. 解析結果 (從 FASTA 格式提取序列)
        lines = format(alignment, "fasta").splitlines()
        seqA = lines[1] # 底本序列 (含Gap)
        seqB = lines[3] # 校本序列 (含Gap)
        
        raw_ops = []
        base_index = 0 
        wit_index = 0 
        
        for a, b in zip(seqA, seqB):
            op_data = {'base': a, 'wit': b, 'base_idx': base_index, 'wit_idx': wit_index}
            if a == b: 
                op_data['op'] = 'Match'
                base_index += 1
                wit_index += 1
            elif a == '-': 
                op_data['op'] = 'Ins'
                wit_index += 1
            elif b == '-': 
                op_data['op'] = 'Del'
                base_index += 1
            else: 
                op_data['op'] = 'Sub'
                base_index += 1
                wit_index += 1
            raw_ops.append(op_data)

        final_blocks = []
        buffer = []
        
        def flush_buffer():
            if not buffer: return
            
            # 1. 提取原始文本
            base_chars = [x['base'] for x in buffer if x['base'] != '-']
            wit_chars = [x['wit'] for x in buffer if x['wit'] != '-']
            base_str = "".join(base_chars)
            wit_str = "".join(wit_chars)
            
            start_idx = buffer[0]['base_idx']
            start_wit_idx = buffer[0]['wit_idx']
            
            # --- [算法核心優化]：基於 ID 歸一化的集合分離法 ---
            
            # A. 歸一化：轉為 ID 序列
            base_ids = []
            for ch in base_str:
                entry = self.variant_db.get(ch)
                # 若是異體字取標準ID，否則取字符本身
                vid = entry['standards'][0] if entry and entry.get('standards') else ch
                base_ids.append(vid)
                
            wit_ids = []
            for ch in wit_str:
                entry = self.variant_db.get(ch)
                vid = entry['standards'][0] if entry and entry.get('standards') else ch
                wit_ids.append(vid)
            
            # B. 集合運算：分離 Common (潛在倒文/異體) 和 Noise (訛/脫/衍)
            from collections import Counter
            c_base = Counter(base_ids)
            c_wit = Counter(wit_ids)
            
            # 計算交集 (Multiset Intersection)
            intersection = c_base & c_wit 
            common_total = sum(intersection.values())
            
            # --- 判斷邏輯 ---
            is_inversion = False
            
            if common_total >= 2:
                base_common_ids = []   
                base_common_chars = [] 
                noise_base_chars = []  
                
                temp_int_base = intersection.copy()
                for ch, vid in zip(base_str, base_ids):
                    if temp_int_base[vid] > 0:
                        base_common_ids.append(vid)
                        base_common_chars.append(ch)
                        temp_int_base[vid] -= 1
                    else:
                        noise_base_chars.append(ch)
                        
                wit_common_ids = []
                wit_common_chars = []
                noise_wit_chars = []
                
                temp_int_wit = intersection.copy()
                for ch, vid in zip(wit_str, wit_ids):
                    if temp_int_wit[vid] > 0:
                        wit_common_ids.append(vid)
                        wit_common_chars.append(ch)
                        temp_int_wit[vid] -= 1
                    else:
                        noise_wit_chars.append(ch)
                
                # C. 倒文判定：比較 ID 序列的順序
                if base_common_ids != wit_common_ids:
                    is_inversion = True
                    
                    # 1. 輸出【倒文】主條目
                    final_blocks.append({
                        'type': '倒',
                        'base_clean': "".join(base_common_chars),
                        'wit_clean': "".join(wit_common_chars),
                        'base_clean_len': len(base_common_chars), 
                        'wit_clean_len': len(wit_common_chars),
                        'idx': start_idx,
                        'wit_idx': start_wit_idx
                    })
                    
                    # 2. 輸出【倒中異】條目 (雙向索引修正版)
                    for vid in intersection.keys():
                        # [關鍵修正 1]: 同時獲取 Wit 字符的相對索引 (w_rel_idx)
                        # b_occurrences 格式: [(字符, 相對位置), ...]
                        b_occurrences = [(ch, i) for i, (ch, v) in enumerate(zip(base_str, base_ids)) if v == vid]
                        w_occurrences = [(ch, i) for i, (ch, v) in enumerate(zip(wit_str, wit_ids)) if v == vid]
                        
                        # 遍歷底本中該 ID 的所有出現
                        for b_char, b_rel_idx in b_occurrences:
                            matched = False
                            # 遍歷校本中該 ID 的所有出現
                            for w_char, w_rel_idx in w_occurrences:
                                # 檢查是否為異體 (ID相同但寫法不同)
                                if b_char != w_char and self.is_variant_relation(b_char, w_char):
                                    final_blocks.append({
                                        'type': '異',
                                        'base_clean': b_char,
                                        'wit_clean': w_char,
                                        'base_clean_len': 1,
                                        'wit_clean_len': 1,
                                        
                                        # [核心修正 2]: Base 和 Wit 都必須加上各自的相對偏移量
                                        # 這樣才能精確指向倒文發生後，該字實際所在的位置
                                        'idx': start_idx + b_rel_idx, 
                                        'wit_idx': start_wit_idx + w_rel_idx 
                                    })
                                    matched = True
                                    
                                    # [細節]: 匹配到一個後立即 break，
                                    # 防止 "A...A" vs "A'...A'" 這種情況產生笛卡爾積式的重複匹配
                                    # 在倒文亂序的情況下，我們優先匹配順序靠前的，這是一種合理的模糊處理
                                    break 
                            
                            if matched:
                                break # 每個底本字符只匹配一次

                    # 3. 輸出【噪聲】條目
                    n_base = "".join(noise_base_chars)
                    n_wit = "".join(noise_wit_chars)
                    
                    if n_base or n_wit:
                        n_type = ""
                        if not n_base: n_type = '衍'
                        elif not n_wit: n_type = '脫'
                        else: n_type = '訛' 
                        
                        final_blocks.append({
                            'type': n_type,
                            'base_clean': n_base,
                            'wit_clean': n_wit,
                            'base_clean_len': len(n_base),
                            'wit_clean_len': len(n_wit),
                            'idx': start_idx,
                            'wit_idx': start_wit_idx
                        })

            # --- 常規分支 (非倒文) ---
            if not is_inversion:
                err_type = ""
                if not wit_str: err_type = '脫'
                elif not base_str: err_type = '衍'
                else: err_type = '訛'
                
                final_blocks.append({
                    'type': err_type,
                    'base_clean_len': len(base_str),
                    'wit_clean_len': len(wit_str),
                    'base_clean': base_str,
                    'wit_clean': wit_str,
                    'idx': start_idx,
                    'wit_idx': start_wit_idx
                })

            buffer.clear()

        # 主循環
        for item in raw_ops:
            if item['op'] == 'Match':
                flush_buffer()
                final_blocks.append({
                    'type': 'Match',
                    'base_clean': item['base'],
                    'wit_clean': item['wit'],
                    'base_clean_len': 1,
                    'wit_clean_len': 1,
                    'idx': item['base_idx'],
                    'wit_idx': item['wit_idx']
                })
            else:
                buffer.append(item)
        
        flush_buffer()
        return final_blocks

    def process_full_text(self, base_full, wit_full):
        return self.align_paragraph(base_full, wit_full)


# ==========================================
# [新增模塊] 形近字分析引擎 (IDS + Pixel IoU)
# ==========================================
class VisualCheckEngine:
    """
    [核心算法]：形近字計算引擎
    1. IDS 結構分析：計算部件重合率 (Jaccard Index)
    2. 像素 IoU 分析：將漢字渲染為位圖，計算像素重合度 (解決 OCR 形近錯誤)
    """
    def __init__(self, ids_path, font_map, embedded_fonts):
        self.ids_db = {}
        self.ids_path = ids_path
        self.font_map = font_map
        self.embedded_fonts = embedded_fonts # 用於回退渲染
        
        # 自定義規則庫
        self.custom_replacements = {} # A->B (單向)
        self.custom_equivalents = set() # (A, B) (雙向)

        self.load_ids()

    def load_ids(self):
        """加載 IDS 數據"""
        if not os.path.exists(self.ids_path):
            print("未找到 ids.txt")
            return
        
        try:
            with open(self.ids_path, 'r', encoding='utf-8') as f:
                for line in f:
                    parts = line.strip().split('\t')
                    if len(parts) >= 3:
                        char = parts[1]
                        ids_seq = parts[2]
                        # 簡單清洗：去掉 IDS 操作符，只保留部件
                        # 操作符範圍 U+2FF0 - U+2FFF
                        clean_components = set()
                        for c in ids_seq:
                            if not ('\u2ff0' <= c <= '\u2fff'):
                                clean_components.add(c)
                        self.ids_db[char] = clean_components
        except Exception as e:
            print(f"IDS 加載失敗: {e}")

    def update_rules(self, replacements, equivalents):
        """更新用戶自定義的 OCR 混淆規則"""
        self.custom_replacements = replacements
        self.custom_equivalents = equivalents

    def get_component_similarity(self, c1, c2):
        """算法 1: 部件 Jaccard 相似度"""
        s1 = self.ids_db.get(c1, set(c1))
        s2 = self.ids_db.get(c2, set(c2))
        
        intersection = len(s1 & s2)
        union = len(s1 | s2)
        
        if union == 0: return 0.0
        return intersection / union

    def get_pixel_iou(self, c1, c2, font_family="SimSun"):
        """
        [算法 2: 像素級 IoU (Intersection over Union)]
        使用 QPainter 在內存中繪製字符，轉為 Numpy 矩陣進行位運算。
        """
        size = 64 # 渲染分辨率 64x64
        
        def char_to_matrix(char, font_name):
            # 創建空白畫布 (Grayscale)
            img = QImage(size, size, QImage.Format.Format_Grayscale8)
            img.fill(0) # 黑色背景
            
            painter = QPainter(img)
            # 設置字體
            font = QFont(font_name, 48) # 字號大一點以減少鋸齒影響
            painter.setFont(font)
            painter.setPen(QColor(255, 255, 255)) # 白色文字
            
            # 居中繪製
            rect = painter.boundingRect(0, 0, size, size, Qt.AlignmentFlag.AlignCenter, char)
            painter.drawText(rect, Qt.AlignmentFlag.AlignCenter, char)
            painter.end()
            
            # 轉換為 Numpy 數組
            ptr = img.bits()
            ptr.setsize(size * size)
            arr = np.array(ptr).reshape(size, size)
            # 二值化：大於 50 視為筆畫
            return arr > 50

        # 智能選擇字體：如果有生僻字，嘗試使用嵌入字體
        target_font = font_family
        # 這裡簡單調用一個邏輯：如果是生僻字，可能需要切換字體
        # 實際項目中可複用 MainWindow 的 get_font_for_char 邏輯
        
        try:
            mat1 = char_to_matrix(c1, target_font)
            mat2 = char_to_matrix(c2, target_font)
            
            # 計算 IoU
            intersection = np.logical_and(mat1, mat2).sum()
            union = np.logical_or(mat1, mat2).sum()
            
            if union == 0: return 0.0
            return intersection / union
        except:
            return 0.0

    def check_similarity(self, base_char, wit_char, weights=(0.4, 0.6), threshold=0.5):
        """
        綜合計算
        :param weights: (IDS權重, 像素權重) 默認像素權重更高，因為 OCR 錯誤多為形似
        """
        # 1. 優先檢查自定義規則 (OCR 容錯)
        if self.custom_replacements.get(base_char) == wit_char:
            return 1.0, "規則庫:單向替換"
        if (base_char, wit_char) in self.custom_equivalents or (wit_char, base_char) in self.custom_equivalents:
            return 1.0, "規則庫:雙向互通"

        # 2. 計算得分
        ids_score = self.get_component_similarity(base_char, wit_char)
        
        # 如果 IDS 差異極大，可能根本不需要算像素（優化性能）
        # 但對於 "日/曰" 這種 IDS 完全不同但像素極高的，不能跳過
        
        pixel_score = self.get_pixel_iou(base_char, wit_char)
        
        final_score = (ids_score * weights[0]) + (pixel_score * weights[1])
        
        desc = []
        if ids_score > 0.5: desc.append(f"部件似({ids_score:.2f})")
        if pixel_score > 0.6: desc.append(f"字形似({pixel_score:.2f})")
        
        return final_score, " + ".join(desc)
    
# ==========================================
# [更新模塊] 廣韻數據直讀加載器 
# ==========================================
import os
from collections import defaultdict, Counter

class GuangYunLoader:
    """
    [核心數據加載器 - 最終完整版]
    功能：
    1. 讀取韻典網 (ytenx) 15個源數據文件。
    2. 內置嚴格的《廣韻》同用/獨用規則 (基於用戶提供標準)。
    3. 提供分層級的音近計算算法 (Level 1-4)。
    4. 統計反切上下字頻次。
    5. 結構化存儲擬音數據。
    """
    
    # 【廣韻同用規則表】
    # 依據用戶提供的標準設置。未列出者默認為「獨用」。
    # 註：代碼中使用源數據文件的 "眞" 字，對應通用寫法 "真"。
    TONGYONG_GROUPS = [
        ["東", "冬"],             # 通攝：鍾獨用
        ["支", "脂", "之"],       # 止攝：微獨用
        ["魚", "虞", "模"],       # 遇攝：三韻同用 (修正：模加入)
        ["佳", "皆"],             # 蟹攝：二韻同用 (新增)
        ["灰", "咍"],             # 蟹攝：二韻同用
        ["眞", "諄", "臻"],       # 臻攝：三韻同用 (眞=真)
        ["元", "魂", "痕"],       # 臻攝：三韻同用 (修正：元加入，文/欣獨用)
        ["寒", "桓"],             # 山攝
        ["刪", "山"],             # 山攝
        ["先", "仙"],             # 山攝
        ["蕭", "宵"],             # 效攝 (新增：肴/豪獨用)
        ["歌", "戈"],             # 果攝
        ["陽", "唐"],             # 宕攝 (江獨用)
        ["庚", "耕", "清"],       # 梗攝 (青獨用)
        ["蒸", "登"],             # 曾攝
        ["尤", "侯"],             # 流攝 (新增：幽獨用)
        ["覃", "談"],             # 咸攝
        ["鹽", "添"],             # 咸攝
        ["咸", "銜"],             # 咸攝
        ["嚴", "凡"],             # 咸攝
        # 獨用韻：江、微、齊、廢、文、欣、肴、豪、麻、青、幽、侵
    ]

    def __init__(self, resource_dir="resources"):
        self.resource_dir = resource_dir
        self.ready = False
        
        # --- 數據庫存儲容器 ---
        self.cjeng_db = {}      # 聲母
        self.gheh_db = {}       # 韻系 (攝)
        self.yonh_mux_db = {}   # 韻母
        self.yonh_miuk_db = {}  # 韻目 (四聲)
        self.sieux_yonh_db = {} # 小韻 (核心)
        self.dzih_db = {}       # 字頭解釋
        
        # 輔助索引
        self.char_map = defaultdict(list) # 字 -> [小韻ID列表]
        self.cjeng_list = []    # 有序聲母列表
        self.she_list = []      # 攝列表
        self.gheh_list = [] # 有序韻系列表 (用於下拉菜單)
        self.rhyme_groups = defaultdict(list) # 攝 -> [韻母列表]
        self.all_chars = []     # 全量字表 (用於字典列表)

        # 統計數據
        self.init_stats = defaultdict(Counter)  # 聲母 -> {反切上字: 次數}
        self.final_stats = defaultdict(Counter) # 韻母 -> {反切下字: 次數}

        # 擬音數據
        self.scholars = []      # 學者名列表
        self.ngix_init = {}     # {聲母名: {學者: 擬音}}
        self.ngix_final = {}    # {韻母名: {學者: 擬音}}
        
        # 同用組索引 (Gheh -> GroupID)
        self.tongyong_map = {}
        self._build_tongyong_map()

        if os.path.exists(resource_dir):
            self.load_all_data()

    def _build_tongyong_map(self):
        """將二維列表轉換為查找字典，方便快速判斷 Level 4C"""
        for idx, group in enumerate(self.TONGYONG_GROUPS):
            for gheh in group:
                self.tongyong_map[gheh] = idx

    def _read_file(self, filename, skip_header=True):
        path = os.path.join(self.resource_dir, filename)
        if not os.path.exists(path):
            print(f"[警告] 文件缺失: {filename}")
            return [], None
        
        with open(path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        if not lines: return [], None
        
        # 處理表頭
        first_line = lines[0].strip()
        header = None
        start_idx = 0
        if first_line.startswith('#'):
            header = first_line[1:].strip().split(' ')
            if skip_header: start_idx = 1
            
        result = []
        for i in range(start_idx, len(lines)):
            line = lines[i].strip()
            if line: result.append(line)
        return result, header

    def load_all_data(self):
        print("正在加載廣韻數據庫...")
        
        # 1. 聲母 (CjengMux)
        lines, _ = self._read_file("CjengMux.txt")
        for line in lines:
            parts = line.split(' ')
            if len(parts) >= 3:
                name = parts[0]
                self.cjeng_db[name] = {
                    'name': name,
                    'type': parts[1], # 聲類(脣音等)
                    'order': int(parts[2])
                }
                self.cjeng_list.append(name)
        self.cjeng_list.sort(key=lambda x: self.cjeng_db[x]['order'])

        # 2. 韻系/攝 (YonhGheh)
        lines, _ = self._read_file("YonhGheh.txt")
        for line in lines:
            parts = line.split(' ')
            if len(parts) >= 2:
                gheh, she = parts[0], parts[1]
                self.gheh_db[gheh] = she
                if she not in self.she_list: self.she_list.append(she)
                if gheh not in self.gheh_list: self.gheh_list.append(gheh)

        # 3. 韻母 (YonhMux)
        lines, _ = self._read_file("YonhMux.txt")
        for line in lines:
            parts = line.split(' ')
            if len(parts) >= 5:
                name = parts[0]
                gheh_ref = parts[1]
                # 處理對立韻母 (第6列，如果有的話)
                tuaih = parts[5] if len(parts) > 5 else None
                
                self.yonh_mux_db[name] = {
                    'name': name,
                    'gheh': gheh_ref,
                    'she': self.gheh_db.get(gheh_ref, '?'),
                    'deng': parts[2],
                    'kaihe': parts[3],
                    'is_kai': (parts[3] == '開'),
                    'shuru': parts[4], # 舒/促/入
                    'tuaih': tuaih
                }
                she = self.gheh_db.get(gheh_ref, '?')
                self.rhyme_groups[she].append(self.yonh_mux_db[name])

        # 4. 韻目 (YonhMiuk)
        lines, _ = self._read_file("YonhMiuk.txt")
        for line in lines:
            parts = line.split(' ')
            if len(parts) >= 3:
                name = parts[0]
                tone = int(parts[2])
                self.yonh_miuk_db[name] = {'tone': tone, 'gheh': parts[1]}

        # 5. 擬音數據
        # 聲母擬音
        lines, header = self._read_file("CjengMuxNgixQim.txt", skip_header=True)
        if header: self.scholars = header[1:] # 跳過第一列"字"
        for line in lines:
            parts = line.split(' ')
            if parts[0] in self.cjeng_db:
                self.ngix_init[parts[0]] = {k: v for k, v in zip(self.scholars, parts[1:])}

        # 韻母擬音
        lines, _ = self._read_file("YonhMuxNgixQim.txt", skip_header=True)
        for line in lines:
            parts = line.split(' ')
            if parts[0] in self.yonh_mux_db:
                self.ngix_final[parts[0]] = {k: v for k, v in zip(self.scholars, parts[1:])}

        # 6. 小韻 (SieuxYonh)
        lines, _ = self._read_file("SieuxYonh.txt")
        for line in lines:
            parts = line.split(' ')
            if len(parts) >= 6:
                sid = parts[0]
                char = parts[1]
                init = parts[2]
                final = parts[3]
                rime = parts[4] # 韻目(包含聲調信息)
                fanqie = parts[5]
                
                # 統計反切
                if len(fanqie) >= 2:
                    upper, lower = fanqie[0], fanqie[1]
                    self.init_stats[init][upper] += 1
                    self.final_stats[final][lower] += 1
                
                final_attr = self.yonh_mux_db.get(final, {})
                tone = self.yonh_miuk_db.get(rime, {}).get('tone', 0)
                
                # 構建舊版兼容字符串 (僅用於兼容，新版已用 dict)
                recon_i_str = " | ".join(self.ngix_init.get(init, {}).values())
                recon_f_str = " | ".join(self.ngix_final.get(final, {}).values())

                self.sieux_yonh_db[sid] = {
                    'id': sid,
                    'char': char,
                    'initial': init,
                    'final': final,
                    'rime': rime,
                    'fanqie': fanqie,
                    'tone': tone,
                    'she': final_attr.get('she', '?'),
                    'gheh': final_attr.get('gheh', '?'),
                    'deng': final_attr.get('deng', '?'),
                    'kaihe': final_attr.get('kaihe', '?'),
                    'is_kai': final_attr.get('is_kai', False),
                    'shuru': final_attr.get('shuru', ''),
                    'tuaih': final_attr.get('tuaih', ''),
                    'recon_i': recon_i_str,
                    'recon_f': recon_f_str
                }

        # 7. 字表 (Dzih)
        lines, _ = self._read_file("Dzih.txt")
        count = 1
        for line in lines:
            parts = line.split(' ')
            if len(parts) >= 2:
                char = parts[0]
                sid = parts[1]
                self.char_map[char].append(sid)
                defn = " ".join(parts[2:]) if len(parts) > 2 else ""
                self.dzih_db[f"{char}_{sid}"] = defn
                
                # 扁平化存儲，用於列表展示
                sy = self.sieux_yonh_db.get(sid)
                if sy:
                    self.all_chars.append({
                        'uid': count,
                        'char': char,
                        'sid': sid,
                        'initial': sy['initial'],
                        'final': sy['final'],
                        'tone': sy['tone'],
                        'fanqie': sy['fanqie'],
                        'def': defn
                    })
                    count += 1

        # 8. 推導音 (Dauh)
        lines, _ = self._read_file("Dauh.txt")
        for line in lines:
            parts = line.split(' ')
            if len(parts) >= 3:
                sid = parts[0]
                if sid in self.sieux_yonh_db:
                    self.sieux_yonh_db[sid]['dauh_zz'] = parts[1] 
                    self.sieux_yonh_db[sid]['dauh_pt'] = parts[2]

        self.ready = True
        print(f"廣韻數據加載完成。全字表: {len(self.all_chars)}")

    # --- 接口 1: 基礎查詢 ---
    def get_sounds(self, char):
        if not self.ready: return []
        sids = self.char_map.get(char, [])
        results = []
        for sid in sids:
            if sid in self.sieux_yonh_db:
                info = self.sieux_yonh_db[sid].copy()
                info['def'] = self.dzih_db.get(f"{char}_{sid}", "")
                results.append(info)
        return results

    # --- 接口 2: 分層級音韻相似度判定 (定稿版) ---
    def check_similarity(self, char_a, char_b):
        """
        [算法核心] 依據用戶定義的四層級規則
        返回: (score, description_string)
        """
        sounds_a = self.get_sounds(char_a)
        sounds_b = self.get_sounds(char_b)
        if not sounds_a or not sounds_b: return 0.0, ""
        
        best_score = 0.0
        best_desc = ""
        
        # 窮舉比較所有讀音組合
        for sa in sounds_a:
            for sb in sounds_b:
                score = 0.0
                desc = ""
                
                # --- 屬性提取 ---
                init_a, init_b = sa['initial'], sb['initial']
                final_a, final_b = sa['final'], sb['final']
                tone_a, tone_b = sa['tone'], sb['tone']
                gheh_a, gheh_b = sa['gheh'], sb['gheh']
                she_a, she_b = sa['she'], sb['she']
                type_a = self.cjeng_db.get(init_a, {}).get('type')
                type_b = self.cjeng_db.get(init_b, {}).get('type')
                
                # Level 1: 精確匹配 (小韻同音)
                if init_a == init_b and final_a == final_b and tone_a == tone_b:
                    score = 1.0
                    desc = "L1:小韻同音"
                
                # Level 2: 聲韻全同 / 四聲相承 (嚴格禁入聲跨調)
                # 規則：聲母相同，韻系相同，聲調不同，且兩者均非入聲(tone!=4)
                elif init_a == init_b and gheh_a == gheh_b:
                    if tone_a == 4 or tone_b == 4:
                        pass # 涉及入聲，不歸入 L2
                    else:
                        score = 0.9
                        desc = "L2:四聲相承"
                        
                else:
                    # Level 3: 嚴格雙聲 / 嚴格疊韻
                    is_shuang = (init_a == init_b)
                    
                    # 嚴格疊韻: 韻系(Gheh)相同 且 舒促性質一致
                    is_shu_a = (tone_a != 4)
                    is_shu_b = (tone_b != 4)
                    is_dieyun = (gheh_a == gheh_b and is_shu_a == is_shu_b)
                    
                    if is_shuang and is_dieyun:
                        score = 0.8; desc = "L3:雙聲疊韻"
                    elif is_shuang:
                        score = 0.6; desc = "L3:雙聲"
                    elif is_dieyun:
                        score = 0.6; desc = "L3:疊韻"
                        
                    # Level 4: 音轉關係
                    else:
                        l4_desc = []
                        
                        # 4A. 陰陽入對轉 (利用 tuaih 字段)
                        tuaih_a = sa.get('tuaih')
                        tuaih_b = sb.get('tuaih')
                        if (tuaih_a and tuaih_a == final_b) or (tuaih_b and tuaih_b == final_a):
                            l4_desc.append("陰陽入對轉")
                            
                        # 4B. 旁轉 (聲旁轉 or 韻旁轉)
                        # B1. 聲旁轉: 聲母不同但發音部位(type)相同
                        if init_a != init_b and type_a and type_b and type_a == type_b:
                            l4_desc.append(f"聲旁轉({type_a})")
                        
                        # B2. 韻旁轉: 同攝但不同韻系
                        if she_a == she_b and gheh_a != gheh_b:
                            l4_desc.append(f"韻旁轉({she_a}攝)")
                            
                        # 4C. 通轉 (異攝同用)
                        gid_a = self.tongyong_map.get(gheh_a)
                        gid_b = self.tongyong_map.get(gheh_b)
                        if she_a != she_b and gid_a is not None and gid_a == gid_b:
                            l4_desc.append("古音通轉(同用)")
                            
                        if l4_desc:
                            score = 0.4
                            desc = "L4:" + "+".join(l4_desc)

                if score > best_score:
                    best_score = score
                    best_desc = desc

        return best_score, best_desc

    # --- 接口 3: 等韻圖數據生成 ---
    # [GuangYunLoader]
    def get_rime_table_data(self, gheh_name):
        """
        生成指定【韻系】的等韻圖數據矩陣，並查找具體的【韻目】名稱
        """
        if not self.ready: return []
        
        # 1. 構建一個臨時查找表：通過 (tone) -> 找到 具體韻名(例如 "董")
        # self.yonh_miuk_db 結構: {'東': {'tone':1, 'gheh':'東'}, '董': {'tone':2, 'gheh':'東'}...}
        tone_to_name = {}
        for r_name, r_info in self.yonh_miuk_db.items():
            if r_info['gheh'] == gheh_name:
                tone_to_name[r_info['tone']] = r_name

        # 2. 找出該韻系下的所有韻母
        relevant_finals = [m for m in self.yonh_mux_db.values() if m['gheh'] == gheh_name]
        if not relevant_finals: return []

        # 3. 構建單元格映射
        cell_map = {} 
        for sid, sy in self.sieux_yonh_db.items():
            if sy['gheh'] == gheh_name:
                try: deng = int(sy['deng'][0]) if sy['deng'] and sy['deng'][0].isdigit() else 0
                except: deng = 0
                key = (sy['tone'], deng, sy['is_kai'], sy['initial'])
                if key not in cell_map: cell_map[key] = sy

        # 4. 生成矩陣
        tables = []
        for tone in range(1, 5):
            # 【關鍵修改】獲取具體的韻目名稱，如果沒找到（比如該韻系沒有入聲），則顯示空或系名
            specific_name = tone_to_name.get(tone, "") 
            
            tone_data = {
                'tone': tone, 
                'rime_name': specific_name,  # <--- 把 "董"、"送" 傳出去
                'grids': []
            }
            
            kai_group = [] 
            he_group = []  
            for deng in range(1, 5):
                row_kai = []
                row_he = []
                for init in self.cjeng_list:
                    k_kai = (tone, deng, True, init)
                    sy_kai = cell_map.get(k_kai)
                    row_kai.append(sy_kai['char'] if sy_kai else "")
                    
                    k_he = (tone, deng, False, init)
                    sy_he = cell_map.get(k_he)
                    row_he.append(sy_he['char'] if sy_he else "")
                    
                kai_group.append(row_kai)
                he_group.append(row_he)
                
            tone_data['grids'] = [kai_group, he_group]
            tables.append(tone_data)
            
        return tables

# ==========================================
# [更新組件] 廣韻數據庫瀏覽器 (v14 顯示修復版)
# ==========================================
from PyQt6.QtWidgets import (QDialog, QVBoxLayout, QHBoxLayout, QTabWidget, QWidget, 
                             QTableWidget, QTableWidgetItem, QHeaderView, QSplitter, 
                             QLabel, QComboBox, QScrollArea, QGroupBox, QPushButton, 
                             QLineEdit, QMenu, QToolButton)
from PyQt6.QtCore import Qt
from PyQt6.QtGui import QColor, QAction
import re

class GuangYunDatabaseWindow(QWidget):
    def __init__(self, loader, parent=None):
        super().__init__(None)
        self.loader = loader
        self.setWindowTitle("《廣韻》數據庫")
        self.resize(1300, 850)
        
        self.setWindowFlags(Qt.WindowType.Window)

        # 字表分頁狀態
        self.char_page = 0
        self.page_size = 50
        self.filtered_chars = self.loader.all_chars # 默認顯示所有
        
        # 擬音顯示狀態 (Set of scholar names)
        self.visible_scholars = set() 
        
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout(self)
        self.tabs = QTabWidget()
        self.tabs.setStyleSheet("QTabBar::tab { min-width: 100px; padding: 10px; font-size: 14px; }")
        
        self.tabs.addTab(self.create_char_tab(), "《廣韻》字表")
        self.tabs.addTab(self.create_initial_tab(), "聲母及反切上字")
        self.tabs.addTab(self.create_final_tab(), "韻母及反切下字")
        self.tabs.addTab(self.create_rime_table_tab(), "等韻圖")
        
        layout.addWidget(self.tabs)

    # --- Tab 1: 字表 (仿韻典網檢索) ---
    def create_char_tab(self):
        widget = QWidget()
        lay = QVBoxLayout(widget)
        
        # 頂部搜索
        search_lay = QHBoxLayout()
        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText("輸入漢字、聲母、韻母或反切進行搜索...")
        self.search_input.returnPressed.connect(self.perform_search)
        btn_search = QPushButton("搜索")
        btn_search.clicked.connect(self.perform_search)
        search_lay.addWidget(self.search_input)
        search_lay.addWidget(btn_search)
        lay.addLayout(search_lay)
        
        # 表格
        self.char_table = QTableWidget()
        self.char_table.setEditTriggers(QTableWidget.EditTrigger.NoEditTriggers)
        cols = ["ID", "字", "聲母", "韻母", "聲調", "小韻號", "反切", "解釋"]
        self.char_table.setColumnCount(len(cols))
        self.char_table.setHorizontalHeaderLabels(cols)
        self.char_table.setWordWrap(True)
        header = self.char_table.horizontalHeader()
        header.setSectionResizeMode(0, QHeaderView.ResizeMode.ResizeToContents) # ID
        header.setSectionResizeMode(1, QHeaderView.ResizeMode.ResizeToContents) # 字
        header.setSectionResizeMode(2, QHeaderView.ResizeMode.ResizeToContents) # 聲母
        header.setSectionResizeMode(3, QHeaderView.ResizeMode.ResizeToContents) # 韻母
        header.setSectionResizeMode(4, QHeaderView.ResizeMode.ResizeToContents) # 聲調
        header.setSectionResizeMode(5, QHeaderView.ResizeMode.ResizeToContents) # 小韻號
        header.setSectionResizeMode(6, QHeaderView.ResizeMode.ResizeToContents) # 反切
        header.setSectionResizeMode(7, QHeaderView.ResizeMode.Stretch)          # 解釋 (拉伸)
        self.char_table.setAlternatingRowColors(True)
        lay.addWidget(self.char_table)
        
        # 底部翻頁
        page_lay = QHBoxLayout()
        self.btn_prev = QPushButton("上一頁")
        self.btn_prev.clicked.connect(self.prev_page)
        self.btn_next = QPushButton("下一頁")
        self.btn_next.clicked.connect(self.next_page)
        self.lbl_page = QLabel("1 / 1")
        
        page_lay.addStretch()
        page_lay.addWidget(self.btn_prev)
        page_lay.addWidget(self.lbl_page)
        page_lay.addWidget(self.btn_next)
        page_lay.addStretch()
        lay.addLayout(page_lay)
        
        self.render_char_table()
        return widget

    def perform_search(self):
        txt = self.search_input.text().strip()
        if not txt:
            self.filtered_chars = self.loader.all_chars
        else:
            # 簡單過濾
            self.filtered_chars = [
                c for c in self.loader.all_chars 
                if txt in c['char'] or txt in c['initial'] or txt in c['final'] or txt in c['fanqie']
            ]
        self.char_page = 0
        self.render_char_table()

    def render_char_table(self):
        total = len(self.filtered_chars)
        max_page = (total - 1) // self.page_size + 1 if total > 0 else 1
        self.lbl_page.setText(f"{self.char_page + 1} / {max_page} (共 {total} 字)")
        
        start = self.char_page * self.page_size
        end = min(start + self.page_size, total)
        data = self.filtered_chars[start:end]
        
        self.char_table.setRowCount(len(data))
        tone_map = {1:'平', 2:'上', 3:'去', 4:'入'}
        
        for i, row in enumerate(data):
            # 1. ID
            self.char_table.setItem(i, 0, QTableWidgetItem(str(row['uid'])))
            
            # 2. 字 (居中)
            item_char = QTableWidgetItem(row['char'])
            item_char.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            self.char_table.setItem(i, 1, item_char)
            
            # 3. 聲母
            self.char_table.setItem(i, 2, QTableWidgetItem(row['initial']))
            # 4. 韻母
            self.char_table.setItem(i, 3, QTableWidgetItem(row['final']))
            # 5. 聲調
            t_str = tone_map.get(row['tone'], str(row['tone']))
            self.char_table.setItem(i, 4, QTableWidgetItem(t_str))
            
            # 6. 【新增】小韻號
            item_sid = QTableWidgetItem(str(row['sid']))
            item_sid.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            self.char_table.setItem(i, 5, item_sid)
            
            # 7. 反切
            self.char_table.setItem(i, 6, QTableWidgetItem(row['fanqie']))
            
            # 8. 【優化】解釋 (去重、自動換行)
            raw_def = row['def']
            # 使用正則表達式：去除開頭的數字和空格 (例如 "17 醜兒" -> "醜兒")
            clean_def = re.sub(r'^\d+\s*', '', raw_def)
            
            item_def = QTableWidgetItem(clean_def)
            # 設置頂部對齊，這樣多行文字會從上面開始排，不會在中間尷尬地懸浮
            item_def.setTextAlignment(Qt.AlignmentFlag.AlignLeft | Qt.AlignmentFlag.AlignTop)
            self.char_table.setItem(i, 7, item_def)

        # 【關鍵修改 3】數據填充完畢後，讓表格根據內容自動調整行高
        # 這會確保長文本全部顯示，沒有省略號
        self.char_table.resizeRowsToContents()

    def prev_page(self):
        if self.char_page > 0:
            self.char_page -= 1
            self.render_char_table()

    def next_page(self):
        total = len(self.filtered_chars)
        if (self.char_page + 1) * self.page_size < total:
            self.char_page += 1
            self.render_char_table()

    # --- Tab 2: 聲母表 (含反切統計 & 擬音控制) ---
    def create_initial_tab(self):
        widget = QWidget()
        lay = QVBoxLayout(widget)
        
        # 工具欄
        tool_lay = QHBoxLayout()
        btn_col = QToolButton()
        btn_col.setText("顯示擬音...")
        btn_col.setPopupMode(QToolButton.ToolButtonPopupMode.InstantPopup)
        menu = QMenu()
        for s in self.loader.scholars:
            act = QAction(s, menu, checkable=True)
            act.triggered.connect(self.update_init_columns)
            menu.addAction(act)
        btn_col.setMenu(menu)
        tool_lay.addWidget(btn_col)
        tool_lay.addStretch()
        lay.addLayout(tool_lay)
        
        self.init_table = QTableWidget()
        self.init_table.setEditTriggers(QTableWidget.EditTrigger.NoEditTriggers)
        self.render_init_table()
        lay.addWidget(self.init_table)
        return widget

    def update_init_columns(self):
        sender = self.sender()
        if sender:
            if sender.isChecked(): self.visible_scholars.add(sender.text())
            else: self.visible_scholars.discard(sender.text())
        self.render_init_table()

    def render_init_table(self):
        # 動態列：聲類, 聲母, 反切上字, [擬音1, 擬音2...]
        headers = ["聲類", "聲母", "反切上字 (頻次)"] + list(self.visible_scholars)
        self.init_table.setColumnCount(len(headers))
        self.init_table.setHorizontalHeaderLabels(headers)
        self.init_table.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeMode.Stretch)
        
        data = self.loader.cjeng_list
        self.init_table.setRowCount(len(data))
        
        for i, name in enumerate(data):
            info = self.loader.cjeng_db[name]
            
            # 1. 基礎信息
            self.init_table.setItem(i, 0, QTableWidgetItem(info['type']))
            self.init_table.setItem(i, 1, QTableWidgetItem(name))
            
            # 2. 反切統計
            stats = self.loader.init_stats.get(name, {})
            # 按頻次降序
            sorted_stats = sorted(stats.items(), key=lambda x: x[1], reverse=True)
            txt = "、".join([f"{k}({v})" for k, v in sorted_stats])
            self.init_table.setItem(i, 2, QTableWidgetItem(txt))
            
            # 3. 擬音
            recons = self.loader.ngix_init.get(name, {})
            for j, s_name in enumerate(self.visible_scholars):
                val = recons.get(s_name, "")
                self.init_table.setItem(i, 3 + j, QTableWidgetItem(val))
                
        self.init_table.resizeRowsToContents()

    # --- Tab 3: 韻母表 (含反切統計 & 擬音) ---
    def create_final_tab(self):
        widget = QWidget()
        lay = QVBoxLayout(widget)
        
        # 類似的擬音控制
        tool_lay = QHBoxLayout()
        btn_col = QToolButton()
        btn_col.setText("顯示擬音...")
        btn_col.setPopupMode(QToolButton.ToolButtonPopupMode.InstantPopup)
        menu = QMenu()
        for s in self.loader.scholars:
            act = QAction(s, menu, checkable=True)
            act.triggered.connect(self.update_final_columns)
            menu.addAction(act)
        btn_col.setMenu(menu)
        tool_lay.addWidget(btn_col)
        tool_lay.addStretch()
        lay.addLayout(tool_lay)
        
        self.final_table = QTableWidget()
        self.final_table.setEditTriggers(QTableWidget.EditTrigger.NoEditTriggers)
        self.render_final_table()
        lay.addWidget(self.final_table)
        return widget

    def update_final_columns(self):
        sender = self.sender()
        if sender:
            if sender.isChecked(): self.visible_scholars.add(sender.text())
            else: self.visible_scholars.discard(sender.text())
        self.render_final_table()

    def render_final_table(self):
        """
        渲染韻母表 (修正版：嚴格遵循廣韻韻目原序)
        """
        headers = ["攝", "韻母", "等", "呼", "調", "反切下字"] + list(self.visible_scholars)
        self.final_table.setColumnCount(len(headers))
        self.final_table.setHorizontalHeaderLabels(headers)
        self.final_table.horizontalHeader().setSectionResizeMode(5, QHeaderView.ResizeMode.Stretch)
        
        # 1. 準備排序權重
        # (A) 韻系順序：直接使用 loader 中加載的列表順序 (這是最關鍵的修正)
        if hasattr(self.loader, 'gheh_list'):
            gheh_order = {name: i for i, name in enumerate(self.loader.gheh_list)}
        else:
            gheh_order = {} # 保底

        # (B) 其他輔助權重
        tone_map = {'平': 1, '上': 2, '去': 3, '入': 4}
        kai_map = {'開': 1, '合': 2}
        
        # 2. 獲取數據
        finals = list(self.loader.yonh_mux_db.values())
        
        # 3. 定義符合音韻學的排序鍵
        def sort_key(info):
            # Key 1: 韻系原序 (東, 冬, 鍾...)
            # 如果找不到(比如數據錯亂)，就放到最後(999)
            k1 = gheh_order.get(info['gheh'], 999)
            
            # Key 2: 聲調 (平上去入)
            k2 = tone_map.get(info['shuru'], 5)
            
            # Key 3: 呼 (開口在前，合口在後)
            k3 = kai_map.get(info['kaihe'], 3)
            
            # Key 4: 等 (1,2,3,4)
            d_str = str(info['deng'])
            if d_str.isdigit():
                k4 = int(d_str)
            else:
                k4 = {'一':1, '二':2, '三':3, '四':4}.get(d_str, 9)
            
            return (k1, k2, k3, k4)

        # 執行排序
        finals.sort(key=sort_key)

        # 4. 渲染表格
        self.final_table.setRowCount(len(finals))
        for i, info in enumerate(finals):
            name = info['name']
            
            # 居中輔助
            def make_item(text):
                it = QTableWidgetItem(str(text))
                it.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
                return it

            self.final_table.setItem(i, 0, make_item(info['she']))
            self.final_table.setItem(i, 1, make_item(name))
            self.final_table.setItem(i, 2, make_item(info['deng']))
            self.final_table.setItem(i, 3, make_item(info['kaihe']))
            self.final_table.setItem(i, 4, make_item(info['shuru']))
            
            # 反切下字
            stats = self.loader.final_stats.get(name, {})
            sorted_stats = sorted(stats.items(), key=lambda x: x[1], reverse=True)
            txt = "、".join([f"{k}({v})" for k, v in sorted_stats])
            self.final_table.setItem(i, 5, QTableWidgetItem(txt))
            
            # 擬音
            recons = self.loader.ngix_final.get(name, {})
            for j, s_name in enumerate(self.visible_scholars):
                self.final_table.setItem(i, 6 + j, QTableWidgetItem(recons.get(s_name, "")))
        
        self.final_table.resizeRowsToContents()

    # --- Tab 4: 等韻圖 (復用之前的邏輯) ---
    def create_rime_table_tab(self):
        widget = QWidget()
        lay = QVBoxLayout(widget)
        
        ctrl_lay = QHBoxLayout()
        ctrl_lay.addWidget(QLabel("選擇韻系："))
        self.combo_she = QComboBox()
        if hasattr(self.loader, 'gheh_list'):
            self.combo_she.addItems(self.loader.gheh_list)
        else:
            self.combo_she.addItems(self.loader.she_list) # 保底
            
        self.combo_she.currentTextChanged.connect(self.render_rime_table)
        ctrl_lay.addWidget(self.combo_she)
        ctrl_lay.addStretch()
        lay.addLayout(ctrl_lay)
        
        self.scroll = QScrollArea()
        self.scroll.setWidgetResizable(True)
        self.rime_content = QWidget()
        self.rime_layout = QVBoxLayout(self.rime_content)
        self.scroll.setWidget(self.rime_content)
        lay.addWidget(self.scroll)

        if self.combo_she.count() > 0:
            self.render_rime_table(self.combo_she.itemText(0))

        return widget

    def render_rime_table(self, gheh_name):
        while self.rime_layout.count():
            item = self.rime_layout.takeAt(0)
            if item.widget(): item.widget().deleteLater()
            
        data = self.loader.get_rime_table_data(gheh_name)
        if not data:
            self.rime_layout.addWidget(QLabel("無數據"))
            return

        initials = self.loader.cjeng_list
        tone_map = {1:'平聲', 2:'上聲', 3:'去聲', 4:'入聲'}
        
        for tone_block in data:
            # 【關鍵修改】使用具體的韻目名稱 (如 "董")，而不是 generic 的 gheh_name
            real_name = tone_block['rime_name']
            
            # 如果該聲調有對應的字 (real_name 不為空)，顯示 "上聲 (董)"
            # 如果是廢韻或者沒有這個聲調，real_name 為空，則只顯示 "上聲" 或不顯示
            if real_name:
                title = f"{tone_map.get(tone_block['tone'], '')} ({real_name})" # 顯示為：上聲 (董)
            else:
                title = f"{tone_map.get(tone_block['tone'], '')} ({gheh_name}系 - 無字)"

            gb = QGroupBox(title)
            gb_lay = QVBoxLayout(gb)
            
            table = QTableWidget()
            table.setEditTriggers(QTableWidget.EditTrigger.NoEditTriggers)
            table.setColumnCount(len(initials) + 2) 
            headers = ["呼", "等"] + initials
            table.setHorizontalHeaderLabels(headers)
            table.verticalHeader().setVisible(False)
            table.setAlternatingRowColors(True)

            table.horizontalHeader().setDefaultAlignment(Qt.AlignmentFlag.AlignCenter)
            
            for grp_idx, grp_data in enumerate(tone_block['grids']): 
                for deng_idx, row_chars in enumerate(grp_data):
                    r = table.rowCount()
                    table.insertRow(r)
                    table.setItem(r, 0, QTableWidgetItem("開" if grp_idx==0 else "合"))
                    table.setItem(r, 1, QTableWidgetItem(str(deng_idx + 1)))
                    for c_i, char in enumerate(row_chars):
                        if char: 
                            it = QTableWidgetItem(char)
                            it.setBackground(QColor("#E8F8F5") if grp_idx==0 else QColor("#FEF9E7"))
                            table.setItem(r, c_i+2, it)
                        else:
                            table.setItem(r, c_i+2, QTableWidgetItem("")) 
                            
            # 1. 先根據內容調整列寬
            table.resizeColumnsToContents()
            
            # 2. 【關鍵步驟】計算表格內容的總寬度
            # horizontalHeader().length() 是所有列寬之和
            # +40 是預留垂直滾動條和邊框的寬度，防止出滾動條
            total_width = table.horizontalHeader().length() + 40
            
            # 3. 強制設置表格寬度，避免它自動拉伸填滿右邊空白
            table.setFixedWidth(total_width)
            table.setFixedHeight(300) # 高度保持固定或自適應皆可
            
            # 4. 【關鍵步驟】添加到佈局時，指定水平居中
            gb_lay.addWidget(table, alignment=Qt.AlignmentFlag.AlignHCenter)
            
            self.rime_layout.addWidget(gb)




# ==========================================
# 3. 音近通假分析面板 (PhoneticAnalysisPanel) - 最終版
# ==========================================
from PyQt6.QtWidgets import (QWidget, QVBoxLayout, QGroupBox, QHBoxLayout, QPushButton, 
                             QLabel, QCheckBox, QTableWidget, QHeaderView, QTableWidgetItem, QMessageBox)
from PyQt6.QtCore import Qt
from PyQt6.QtGui import QColor, QFont

class PhoneticAnalysisPanel(QWidget):
    """
    [音近通假分析面板 - v13 UI 終極優化版]
    特性：
    1. UI 淨化：移除擬音列。
    2. 家族式歸一化：支持異體字家族借音查找 (標藍)。
    3. 矩陣掃描：支持多字對多字 (M:N) 的內部比對。
    4. 可視化增強：
       - 音韻地位增加聲母顯示。
       - 判定列顯示雙行 HTML 上下文，紅字高亮匹配字。
    """
    def __init__(self, phonetic_loader, collation_engine):
        super().__init__()
        self.loader = phonetic_loader
        self.engine = collation_engine # 用於訪問 variant_db
        self.all_data = {}
        self.current_wit = ""
        self.db_window = None
        self.cached_results = []
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout(self)
        
        # --- 1. 頂部控制欄 ---
        ctrl_box = QGroupBox("分析配置")
        ctrl_lay = QHBoxLayout(ctrl_box)
        
        # 按鈕
        self.btn_db = QPushButton("《廣韻》字典")
        self.btn_db.clicked.connect(self.open_database)
        self.btn_charts = QPushButton("等韻圖")
        self.btn_charts.clicked.connect(self.open_charts)
        ctrl_lay.addWidget(self.btn_db)
        ctrl_lay.addWidget(self.btn_charts)
        
        ctrl_lay.addSpacing(20)
        
        # 過濾等級
        ctrl_lay.addWidget(QLabel("顯示等級："))
        self.cbs = {}
        for l in ["L1 聲韻調俱同", "L2 四聲相承", "L3 雙聲/疊韻/雙聲疊韻", "L4 通轉/對轉/旁轉"]:
            cb = QCheckBox(l)
            cb.setChecked(True)
            cb.clicked.connect(self.render_filtered_results)
            self.cbs[l] = cb
            ctrl_lay.addWidget(cb)
            
        ctrl_lay.addStretch()
        
        # 運行按鈕
        self.btn_run = QPushButton(" ▶ 開始通假分析 ")
        self.btn_run.setStyleSheet("background-color: #B74639; color: white; font-weight: bold; font-family: 'KaiTi'; font-size: 14px;")
        self.btn_run.clicked.connect(self.run_analysis)
        ctrl_lay.addWidget(self.btn_run)
        
        layout.addWidget(ctrl_box)
        
        # --- 2. 結果表格 ---
        self.table = QTableWidget()
        self.table.setEditTriggers(QTableWidget.EditTrigger.NoEditTriggers)
        self.table.setColumnCount(6)
        # [優化] 調整列寬比例
        self.table.setHorizontalHeaderLabels(["相似度", "判定說明", "底本字", "校本字", "底本音韻地位", "校本音韻地位"])

        self.table.horizontalHeader().setDefaultAlignment(Qt.AlignmentFlag.AlignCenter)
        header = self.table.horizontalHeader()
        
        # 1. 小列：設置為固定寬度 (Fixed)
        header.setSectionResizeMode(0, QHeaderView.ResizeMode.Fixed)
        self.table.setColumnWidth(0, 80)  # 相似度：60px 足夠顯示 "100.0"
        
        header.setSectionResizeMode(2, QHeaderView.ResizeMode.Fixed)
        self.table.setColumnWidth(2, 80)  # 底本字：50px
        
        header.setSectionResizeMode(3, QHeaderView.ResizeMode.Fixed)
        self.table.setColumnWidth(3, 80)  # 校本字：50px
        
        # 2. 信息列：給予較寬的固定空間，防止文字換行或擠壓
        # 音韻地位內容較長 (如 "【見】【東三】...")，給 230px 左右比較合適
        header.setSectionResizeMode(4, QHeaderView.ResizeMode.Fixed)
        self.table.setColumnWidth(4, 340) 
        
        header.setSectionResizeMode(5, QHeaderView.ResizeMode.Fixed)
        self.table.setColumnWidth(5, 340)
        
        # 3. 說明列：佔用【剩餘】的所有空間 (Stretch)
        # 由於上述列已經佔用了約 620px，剩下的空間給說明列就不會顯得過於空曠了
        header.setSectionResizeMode(1, QHeaderView.ResizeMode.Stretch)
        # ===================================
        
        self.table.setAlternatingRowColors(True)
        # 設置行高自適應
        self.table.verticalHeader().setSectionResizeMode(QHeaderView.ResizeMode.ResizeToContents)
        
        layout.addWidget(self.table)
        
        # --- 3. 狀態欄 ---
        self.lbl_status = QLabel("就緒！請選擇校本並點擊開始！")
        self.lbl_status.setStyleSheet("color: #666; margin-top: 5px;")
        layout.addWidget(self.lbl_status)

    def load_data(self, wit_name, full_data, *args):
        self.current_wit = wit_name
        self.all_data = full_data
        self.table.setRowCount(0)
        self.lbl_status.setText(f"當前分析對象：{wit_name}")

    def _show_db_window(self, tab_index):
        if not self.loader.ready:
            QMessageBox.warning(self, "錯誤", "數據未加載")
            return

        # 如果窗口不存在，或者已經被用戶關閉（deleted），則重新創建
        if self.db_window is None or not self.db_window.isVisible():
            # 這裡不傳 self 作為 parent，確保它是獨立窗口
            self.db_window = GuangYunDatabaseWindow(self.loader) 
        
        # 設置 Tab
        self.db_window.tabs.setCurrentIndex(tab_index)
        
        # 顯示窗口 (非模態)
        self.db_window.show()
        
        # 如果窗口被最小化了，恢復它
        if self.db_window.isMinimized():
            self.db_window.showNormal()
            
        # 將窗口置頂激活
        self.db_window.activateWindow()
        self.db_window.raise_()

    def open_database(self):
        if not self.loader.ready:
            QMessageBox.warning(self, "錯誤", "數據未加載")
            return

        # 如果窗口不存在，或者被關閉了，就創建新的
        if self.db_window is None or not self.db_window.isVisible():
            self.db_window = GuangYunDatabaseWindow(self.loader)
        
        self.db_window.tabs.setCurrentIndex(0) # 跳轉到字表 Tab
        self.db_window.show()           # 【關鍵】改成 show()
        self.db_window.activateWindow() # 【關鍵】把窗口提到最前

    def open_charts(self):
        if not self.loader.ready:
            QMessageBox.warning(self, "錯誤", "數據未加載")
            return

        # 同樣的邏輯：確保只有一個窗口實例
        if self.db_window is None or not self.db_window.isVisible():
            self.db_window = GuangYunDatabaseWindow(self.loader)
            
        self.db_window.tabs.setCurrentIndex(3) # 跳轉到等韻圖 Tab
        self.db_window.show()           # 【關鍵】改成 show()
        self.db_window.activateWindow() # 【關鍵】把窗口提到最前

    # =======================================================
    #  核心功能 A：異體字家族歸一化
    # =======================================================
    def get_sounds_normalized(self, char):
        # 1. 查本人
        sounds = self.loader.get_sounds(char)
        if sounds:
            return sounds, char, False
            
        # 2. 查家族
        if not self.engine or char not in self.engine.variant_db:
            return [], char, False
            
        entry = self.engine.variant_db[char]
        candidates = []
        
        standards = entry.get('standards', [])
        for s in standards:
            if s != char and s not in candidates: candidates.append(s)
        variants = entry.get('variants', [])
        for v in variants:
            if v != char and v not in candidates: candidates.append(v)
        for std in standards:
            if std in self.engine.variant_db:
                std_vars = self.engine.variant_db[std].get('variants', [])
                for v in std_vars:
                    if v != char and v not in candidates: candidates.append(v)

        for cand in candidates:
            cand_sounds = self.loader.get_sounds(cand)
            if cand_sounds:
                return cand_sounds, cand, True
        
        return [], char, False

    # =======================================================
    #  核心功能 B：矩陣掃描與 UI 渲染
    # =======================================================
    def run_analysis(self):
        # 1. 前置檢查
        if not self.loader.ready:
            QMessageBox.warning(self, "錯誤", "廣韻數據庫未加載！")
            return
        if not self.current_wit or self.current_wit not in self.all_data:
            return

        results = self.all_data[self.current_wit]['results']
        
        # 2. 寬鬆過濾
        errors = [
            r for r in results 
            if r['type'] == '訛' 
            and 0 < len(r['base_clean']) <= 10 
            and 0 < len(r['wit_clean']) <= 10
        ]
        
        if not errors:
            self.lbl_status.setText("當前校本無符合條件的訛誤塊。")
            return
        
        # 3. 初始化緩存，準備計算
        self.cached_results = [] 
        total_found = 0
        
        # 4. 循環計算 (這裡刪除了所有過濾邏輯和表格操作)
        for r in errors:
            base_str = r['base_clean']
            wit_str = r['wit_clean']
            
            for i, b_char_raw in enumerate(base_str):
                for j, w_char_raw in enumerate(wit_str):
                    # 獲取讀音及借字狀態
                    sounds_a, b_disp, b_is_borrowed = self.get_sounds_normalized(b_char_raw)
                    sounds_b, w_disp, w_is_borrowed = self.get_sounds_normalized(w_char_raw)
                    
                    if not sounds_a or not sounds_b: continue
                    
                    for sa in sounds_a:
                        for sb in sounds_b:
                            score, desc = self._check_single_pair(sa, sb)
                            if score == 0: continue # 分數為0跳過
                            
                            # 生成 HTML 上下文 (用於顯示)
                            b_html = ""
                            for idx, c in enumerate(base_str):
                                b_html += f"<span style='color:#B74639; font-weight:bold; font-size:16px;'>{c}</span>" if idx == i else c
                            w_html = ""
                            for idx, c in enumerate(wit_str):
                                w_html += f"<span style='color:#B74639; font-weight:bold; font-size:16px;'>{c}</span>" if idx == j else c
                            
                            html_text = f"<div style='line-height:1.4; margin:4px; text-align:center;'><b>{desc}</b><br><span style='color:#555;'>{b_html} ↔ {w_html}</span></div>"
                            
                            # 生成音韻地位文本
                            pos_a = f"【{sa['initial']}】【{sa['final']}】 {sa['she']}攝 {sa['deng']}等 {sa['kaihe']} {self._get_tone_name(sa['tone'])}"
                            pos_b = f"【{sb['initial']}】【{sb['final']}】 {sb['she']}攝 {sb['deng']}等 {sb['kaihe']} {self._get_tone_name(sb['tone'])}"

                            # 【關鍵】將所有 UI 需要的信息都存入字典
                            self.cached_results.append({
                                'level': desc.split(':')[0], # 用於過濾等級
                                'score': score,
                                'html': html_text,
                                # 以下是為了您說的顏色和 Tooltip 功能保留的數據：
                                'b_disp': b_disp,         # 顯示字 (可能是歸一化後的)
                                'b_raw': b_char_raw,      # 原字 (用於 Tooltip)
                                'b_bor': b_is_borrowed,   # 是否借字 (用於變藍色)
                                'w_disp': w_disp, 
                                'w_raw': w_char_raw, 
                                'w_bor': w_is_borrowed,
                                'pos_a': pos_a,
                                'pos_b': pos_b
                            })
                            total_found += 1

        self.lbl_status.setText(f"分析完成。共計算出 {total_found} 組潛在關係。")
        # 5. 計算完，交給顯示函數
        self.render_filtered_results()

    def _get_tone_name(self, t):
        return {1:'平', 2:'上', 3:'去', 4:'入'}.get(t, '')

    # =======================================================
    #  核心功能 C：L1-L4 判定邏輯 (含 TONGYONG 表)
    # =======================================================
    def _check_single_pair(self, sa, sb):
        """
        單對讀音判定邏輯
        """
        init_a, init_b = sa['initial'], sb['initial']
        final_a, final_b = sa['final'], sb['final']
        tone_a, tone_b = sa['tone'], sb['tone']
        gheh_a, gheh_b = sa['gheh'], sb['gheh']
        she_a, she_b = sa['she'], sb['she']
        type_a = self.loader.cjeng_db.get(init_a, {}).get('type')
        type_b = self.loader.cjeng_db.get(init_b, {}).get('type')
        
        # L1: 精確匹配
        if init_a == init_b and final_a == final_b and tone_a == tone_b:
            return 1.0, "L1:小韻同音"
        
        # L2: 四聲相承 (嚴格禁入聲)
        if init_a == init_b and gheh_a == gheh_b:
            if tone_a != 4 and tone_b != 4:
                return 0.9, "L2:四聲相承"
                
        # L3: 嚴格雙聲 / 嚴格疊韻
        is_shuang = (init_a == init_b)
        # 疊韻：韻系相同 且 舒促性質一致 (同舒或同入)
        is_shu_a = (tone_a != 4)
        is_shu_b = (tone_b != 4)
        is_dieyun = (gheh_a == gheh_b and is_shu_a == is_shu_b)
        
        if is_shuang and is_dieyun: return 0.8, "L3:雙聲疊韻"
        if is_shuang: return 0.6, "L3:雙聲"
        if is_dieyun: return 0.6, "L3:疊韻"
        
        # L4: 音轉關係
        desc = []
        # 4A. 對轉 (利用數據中的 tuaih 字段)
        tuaih_a = sa.get('tuaih')
        tuaih_b = sb.get('tuaih')
        if (tuaih_a and tuaih_a == final_b) or (tuaih_b and tuaih_b == final_a):
            desc.append("陰陽入對轉")
            
        # 4B. 旁轉
        if init_a != init_b and type_a and type_b and type_a == type_b:
            desc.append(f"聲旁轉({type_a})")
        if she_a == she_b and gheh_a != gheh_b:
            desc.append(f"韻旁轉({she_a}攝)")
            
        # 4C. 通轉 (同用) - 使用 loader 中的同用表
        gid_a = self.loader.tongyong_map.get(gheh_a)
        gid_b = self.loader.tongyong_map.get(gheh_b)
        if she_a != she_b and gid_a is not None and gid_a == gid_b:
            desc.append("古音通轉")
            
        if desc: return 0.4, "L4:" + "+".join(desc)
        
        return 0.0, ""
    
    def render_filtered_results(self):
        """
        根據當前勾選框狀態，從緩存中讀取數據並顯示
        """
        self.table.setRowCount(0)
        
        # 1. 獲取當前激活的等級
        active_levels = []
        for k, cb in self.cbs.items():
            if cb.isChecked():
                active_levels.append(k.split(' ')[0]) # 取 "L1"
        
        # 2. 過濾並填充表格
        visible_count = 0
        
        # 暫時關閉排序功能以提高插入速度
        self.table.setSortingEnabled(False)
        
        for data in self.cached_results:
            if data['level'] not in active_levels:
                continue
                
            row = self.table.rowCount()
            self.table.insertRow(row)
            
            # Col 0: 分數
            item_score = QTableWidgetItem(f"{data['score']:.1f}")
            if data['score'] >= 0.8: item_score.setBackground(QColor("#D5F5E3"))
            item_score.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            self.table.setItem(row, 0, item_score)
            
            # Col 1: 說明 (HTML)
            lbl_desc = QLabel(data['html'])
            lbl_desc.setTextFormat(Qt.TextFormat.RichText)
            lbl_desc.setWordWrap(True)
            lbl_desc.setAlignment(Qt.AlignmentFlag.AlignCenter)
            self.table.setCellWidget(row, 1, lbl_desc)
            
            # Col 2: 底本字
            item_b = QTableWidgetItem(data['b_disp'])
            if data['b_bor']: item_b.setForeground(QColor("#2E5C8A"))
            item_b.setToolTip(f"原字：{data['b_raw']}" + (" (使用《廣韻》字形)" if data['b_bor'] else ""))
            item_b.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            self.table.setItem(row, 2, item_b)
            
            # Col 3: 校本字
            item_w = QTableWidgetItem(data['w_disp'])
            if data['w_bor']: item_w.setForeground(QColor("#2E5C8A"))
            item_w.setToolTip(f"原字：{data['w_raw']}" + (" (使用《廣韻》字形)" if data['w_bor'] else ""))
            item_w.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            self.table.setItem(row, 3, item_w)
            
            # Col 4 & 5: 音韻地位
            item_pa = QTableWidgetItem(data['pos_a'])
            item_pa.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            self.table.setItem(row, 4, item_pa)
            
            item_pb = QTableWidgetItem(data['pos_b'])
            item_pb.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            self.table.setItem(row, 5, item_pb)
            
            visible_count += 1
            
        self.table.setSortingEnabled(True)
        
        # 更新狀態欄
        self.lbl_status.setText(f"勾選音韻等級下共 {visible_count} 條；所有存在音近關係字總計 {len(self.cached_results)} 條。")
    
# ==========================================
# 3. 校勘工作線程 
# ==========================================
class WorkerThread(QThread):
    progress_update = pyqtSignal(str) 
    finished = pyqtSignal(dict, str, str, dict, list, dict, dict) 
    error = pyqtSignal(str)

    def __init__(self, base_path, wit_paths, engine, use_variant_filter):
        super().__init__()
        self.base_path = base_path
        self.wit_paths = wit_paths 
        self.engine = engine
        self.use_variant_filter = use_variant_filter

    def read_docx(self, path):
        try:
            doc = Document(path)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except Exception as e:
            return ""

    def clean_filename(self, path):
        name = os.path.splitext(os.path.basename(path))[0]
        if not name.endswith("本"): name += "本"
        return name

    def clean_and_map(self, original_text):
        """
        [功能說明]：文本清洗與坐標映射。
        去除標點符號，同時建立一個映射表 (mapping)，
        記錄清洗後每個字符在原始文本中的位置，確保後續校勘結果能精確回溯到原文位置。
        """
        pattern = r"[，。、；：？！“”‘’「」『』（）《》〈〉【】〖〗—….,;?!:\"'(){}<>\[\]\-~\s\n\t]"
        clean_text = ""
        mapping = [] 
        for i, char in enumerate(original_text):
            if not re.match(pattern, char):
                clean_text += char
                mapping.append(i) 
        mapping.append(len(original_text))
        return clean_text, mapping

    # === 多對多異體字最佳路徑解算器 ===
    def solve_mn_block(self, base_seq, wit_seq, base_start_idx, wit_start_idx):
        """
        [算法核心]：基於異體字加權的局部動態規劃 (Local DP with Variant Weighting)。
        
        [問題背景]：
        在處理 M:N 的文本差異塊時（例如底本3字，校本4字），傳統的 LCS 算法只能識別相同字。
        本算法引入了「異體字等同」的權重概念，試圖在差異塊內部尋找隱藏的異體對應關係。
        
        [算法邏輯]：
        1. 構建 DP 表：如果字符相同 OR 是異體關係，則 dp[i][j] = dp[i-1][j-1] + 1。
        2. 路徑回溯：優先匹配異體字路徑。
        3. 回退檢查 (Backtracking Check)：
           代碼檢查拆分結果 `has_common_anchor`。
           如果拆分出的結果全是「脫」和「衍」（即沒有找到任何錨點），
           說明這段 M:N 的差異在語義上完全無關（純噪聲）。
           此時拒絕拆分，保持原樣為一整塊「訛」，防止算法產生過度碎片化 (Over-segmentation)。
        """
        n = len(base_seq)
        m = len(wit_seq)
        dp = [[0] * (m + 1) for _ in range(n + 1)]
        
        for i in range(1, n + 1):
            for j in range(1, m + 1):
                b_char = base_seq[i-1]
                w_char = wit_seq[j-1]
                is_match = (b_char == w_char) or self.engine.is_variant_relation(b_char, w_char)
                if is_match:
                    dp[i][j] = dp[i-1][j-1] + 1
                else:
                    dp[i][j] = max(dp[i-1][j], dp[i][j-1])
        
        raw_ops = []
        i, j = n, m
        while i > 0 and j > 0:
            b_char = base_seq[i-1]
            w_char = wit_seq[j-1]
            is_match = (b_char == w_char) or self.engine.is_variant_relation(b_char, w_char)
            
            if is_match and dp[i][j] == dp[i-1][j-1] + 1:
                op_type = 'Match' if b_char == w_char else '異'
                raw_ops.append({
                    'type': op_type, 
                    'base': b_char, 'wit': w_char, 
                    'idx': base_start_idx + i - 1, 
                    'wit_idx': wit_start_idx + j - 1
                })
                i -= 1
                j -= 1
            elif dp[i-1][j] >= dp[i][j-1]:
                # 脫文：底本有字，對校本無
                raw_ops.append({
                    'type': '脫', 
                    'base': b_char, 'wit': '', 
                    'idx': base_start_idx + i - 1, 
                    'wit_idx': wit_start_idx + j 
                })
                i -= 1
            else:
                # 衍文：底本無字，對校本有
                raw_ops.append({
                    'type': '衍', 
                    'base': '', 'wit': w_char, 
                    'idx': base_start_idx + i, 
                    'wit_idx': wit_start_idx + j - 1
                })
                j -= 1
        
        # 處理剩餘的邊緣情況
        while i > 0:
            raw_ops.append({
                'type': '脫', 
                'base': base_seq[i-1], 'wit': '', 
                'idx': base_start_idx + i - 1, 
                'wit_idx': wit_start_idx + j 
            })
            i -= 1
        while j > 0:
            raw_ops.append({
                'type': '衍', 
                'base': '', 'wit': wit_seq[j-1], 
                'idx': base_start_idx + i, 
                'wit_idx': wit_start_idx + j - 1
            })
            j -= 1
            
        raw_ops.reverse()
        
        merged_res = []
        if not raw_ops: return []
        
        current_block = raw_ops[0].copy()
        current_block['base_clean_len'] = 1 if current_block['base'] else 0
        current_block['wit_clean_len'] = 1 if current_block['wit'] else 0

        for k in range(1, len(raw_ops)):
            next_op = raw_ops[k]
            if current_block['type'] == next_op['type'] and current_block['type'] in ['脫', '衍']:
                current_block['base'] += next_op['base']
                current_block['wit'] += next_op['wit']
                current_block['base_clean_len'] += 1 if next_op['base'] else 0
                current_block['wit_clean_len'] += 1 if next_op['wit'] else 0
            else:
                current_block['base_clean'] = current_block.pop('base')
                current_block['wit_clean'] = current_block.pop('wit')
                merged_res.append(current_block)
                current_block = next_op.copy()
                current_block['base_clean_len'] = 1 if current_block['base'] else 0
                current_block['wit_clean_len'] = 1 if current_block['wit'] else 0
        
        current_block['base_clean'] = current_block.pop('base')
        current_block['wit_clean'] = current_block.pop('wit')
        merged_res.append(current_block)
        
        return merged_res

    def run(self):
        try:
            if not self.base_path or not os.path.exists(self.base_path): raise Exception("無法找到底本文件！")
            base_raw = self.read_docx(self.base_path)
            if not base_raw: raise Exception("底本內容為空或讀取失敗！")
            base_clean, base_map = self.clean_and_map(base_raw)
            all_results = {}
            wit_originals = {}
            wit_maps = {}
            wit_cleans = {}

            for idx, wit_path in enumerate(self.wit_paths):
                wit_name = self.clean_filename(wit_path)
                self.progress_update.emit(f"正在處理: {wit_name} ...")
                wit_raw = self.read_docx(wit_path)
                wit_clean, wit_map = self.clean_and_map(wit_raw)
                wit_originals[wit_name] = wit_raw
                wit_maps[wit_name] = wit_map
                wit_cleans[wit_name] = wit_clean
                
                initial_results = self.engine.process_full_text(base_clean, wit_clean)
                final_res = []
                var_res = []
                
                if self.use_variant_filter:
                    for r in initial_results:
                        if r['type'] == '訛':
                            bc_seq = list(r['base_clean'])
                            wc_seq = list(r['wit_clean'])
                            
                            # === 分支 1: 長度相等 (1對1) ===
                            if len(bc_seq) == len(wc_seq):
                                temp_base = []
                                temp_wit = []
                                temp_start_idx = -1
                                temp_start_wit_idx = -1

                                for i in range(len(bc_seq)):
                                    b_char = bc_seq[i]
                                    w_char = wc_seq[i]
                                    curr_idx = r['idx'] + i
                                    curr_wit_idx = r['wit_idx'] + i
                                    is_variant = self.engine.is_variant_relation(b_char, w_char)
                                    
                                    if is_variant:
                                        if temp_base:
                                            final_res.append({'type': '訛', 'base_clean': "".join(temp_base), 'wit_clean': "".join(temp_wit), 'base_clean_len': len(temp_base), 'wit_clean_len': len(temp_wit), 'idx': temp_start_idx, 'wit_idx': temp_start_wit_idx})
                                            temp_base = []
                                            temp_wit = []
                                        
                                        sub_r = r.copy()
                                        sub_r.update({'type': '異', 'base_clean': b_char, 'wit_clean': w_char, 'idx': curr_idx, 'wit_idx': curr_wit_idx, 'base_clean_len': 1, 'wit_clean_len': 1})
                                        var_res.append(sub_r)
                                        final_res.append(sub_r) # 異體字保留在主列表中
                                    else:
                                        if not temp_base:
                                            temp_start_idx = curr_idx
                                            temp_start_wit_idx = curr_wit_idx
                                        temp_base.append(b_char)
                                        temp_wit.append(w_char)
                                if temp_base:
                                    final_res.append({'type': '訛', 'base_clean': "".join(temp_base), 'wit_clean': "".join(temp_wit), 'base_clean_len': len(temp_base), 'wit_clean_len': len(temp_wit), 'idx': temp_start_idx, 'wit_idx': temp_start_wit_idx})

                            # === 分支 2: 長度不等 (M:N) - 調用 LCS 解算器 ===
                            else:
                                mn_blocks = self.solve_mn_block(bc_seq, wc_seq, r['idx'], r['wit_idx'])
                                
                                # --- 【核心邏輯】回退檢查 ---
                                # 檢查這一組拆分結果中，是否包含任何 "Match"(相同) 或 "異"(異體)
                                # 如果裡面全是 "脫" 和 "衍"，說明這兩個字串完全沒關係
                                has_common_anchor = any(b['type'] in ['Match', '異'] for b in mn_blocks)
                                
                                if has_common_anchor:
                                    # A. 找到了異體字或相同字 -> 接受拆分結果
                                    for blk in mn_blocks:
                                        if blk['type'] == '異':
                                            var_res.append(blk)
                                            final_res.append(blk) 
                                        else:
                                            final_res.append(blk)
                                else:
                                    # B. 完全沒關係 (純噪聲) -> 拒絕拆分，保持原樣為 "訛"
                                    final_res.append(r)
                                    # 如果算法引擎直接返回了 '異' (來自倒文或噪聲剝離)
                        elif r['type'] == '異':
                            var_res.append(r)   # <--- 關鍵修復：抄送給異體表
                            final_res.append(r) # 加入主表
                        else:
                            final_res.append(r)
                else:
                    final_res = initial_results

                stats = Counter([r['type'] for r in final_res if r['type'] != 'Match'])
                all_results[wit_name] = {'results': final_res, 'variants': var_res, 'stats': stats}
            
            self.finished.emit(all_results, base_raw, base_clean, wit_originals, base_map, wit_maps, wit_cleans)
        except Exception as e:
            import traceback
            traceback.print_exc()
            self.error.emit(str(e))

# ==========================================
# 4. Word 腳注導出線程 
# ==========================================
class WordExportThread(QThread):
    progress_signal = pyqtSignal(int)
    status_signal = pyqtSignal(str)
    finished_signal = pyqtSignal(str)
    error_signal = pyqtSignal(str)

    def __init__(self, output_path, base_full_text, insertion_points, available_fonts, font_map):
        super().__init__()
        self.output_path = output_path
        self.base_full_text = base_full_text
        self.insertion_points = insertion_points 
        self.available_fonts = available_fonts 
        self.font_map = font_map            

    # --- 檢測單字字體 ---
    def get_best_font(self, char):
        """
        [輔助功能]：字體匹配邏輯。
        遍歷可用字體列表，檢測哪個字體包含該字符的字形。
        優先解決生僻字的顯示問題。
        """
        if not char: return None
        char_code = ord(char)
        try:
            if char_code <= 0xFFFF:
                fm_default = QFontMetrics(QFont("SimSun"))
                if fm_default.inFont(char): return None 
        except: pass
        
        for font_name in self.available_fonts:
            try:
                raw_font = QRawFont.fromFont(QFont(font_name))
                if raw_font.supportsCharacter(char_code):
                    return self.font_map.get(font_name, font_name)
            except: continue
        return None

    # --- 核心：智能寫入函數---
    def smart_insert_com(self, word_range, text):
        """
        [技術核心]：基於 COM 的智能文本寫入。
        [問題解決]：Word 的 InsertAfter 方法對於超大字符集支持不佳，往往顯示為方框。
        [解決方案]：
        逐字掃描文本，根據字符所需的最佳字體（如『全宋體』或『花園明朝』）將文本切分為多個片段。
        對每個片段分別調用 COM 接口寫入，並顯式設置該 Range 的字體屬性。
        這確保了古籍中的生僻字在導出文檔中能被正確渲染。
        """
        if not text: return
        # 將文字按字體分組，例如: [('宋體', '天'), ('全宋體-1', '𠮷')]
        segments = []
        current_font = None
        current_buffer = []
        for char in text:
            f_name = self.get_best_font(char)
            if f_name is None: f_name = "Default"
            
            if f_name != current_font:
                if current_buffer:
                    segments.append((current_font, "".join(current_buffer)))
                current_font = f_name
                current_buffer = [char]
            else:
                current_buffer.append(char)
        if current_buffer:
            segments.append((current_font, "".join(current_buffer)))

        # 寫入 Word
        for font_name, content in segments:
            # 確保始終在範圍末尾追加
            word_range.Collapse(0) # 0 = wdCollapseEnd
            word_range.InsertAfter(content)
            if font_name != "Default":
                word_range.Font.Name = font_name
                word_range.Font.NameFarEast = font_name 
            else:
                # 普通字體，保持樣式默認 (宋體)
                word_range.Font.Name = "Times New Roman"
                word_range.Font.NameFarEast = "宋體"
            
            word_range.Collapse(0)

    def run(self):
        pythoncom.CoInitialize()
        word = None
        doc = None

        try:
            self.status_signal.emit("正在啟動 Word ...")
            word = win32com.client.dynamic.Dispatch('Word.Application')
            word.Visible = False
            word.DisplayAlerts = False
            word.ScreenUpdating = False 
            
            doc = word.Documents.Add()
            
            # 設置默認樣式
            try:
                style_normal = doc.Styles(-1) 
                style_normal.Font.Name = "宋體"        
                style_normal.Font.NameAscii = "Times New Roman" 
                style_normal.Font.Size = 12
                style_normal.ParagraphFormat.LineSpacingRule = 4 
                style_normal.ParagraphFormat.LineSpacing = 20
            except: pass

            try:
                style_fn = doc.Styles(-30) # Footnote Text
                style_fn.Font.Name = "宋體"
                style_fn.Font.Size = 10.5
            except: pass
            
            self.status_signal.emit("正在寫入...")
            
            sorted_indices = sorted(self.insertion_points.keys())
            cursor = 0
            total_points = len(sorted_indices)
            
            # 獲取全文字檔範圍
            rng = doc.Content 

            for i, note_idx in enumerate(sorted_indices):
                if i % 10 == 0:
                    self.progress_signal.emit(int((i / (total_points + 1)) * 100))
                
                # === 1. 寫入正文塊 ===
                text_chunk = self.base_full_text[cursor : note_idx]
                if text_chunk:
                    # 確保 rng 在文檔末尾
                    rng = doc.Content
                    rng.Collapse(0)
                    self.smart_insert_com(rng, text_chunk)
                
                # === 2. 插入腳注 ===
                notes = self.insertion_points[note_idx]
                for note_content in notes:
                    # 確保 rng 在剛寫入的正文後面
                    rng = doc.Content
                    rng.Collapse(0) 
                    
                    # 創建一個空內容的腳注
                    fn = doc.Footnotes.Add(Range=rng, Text="") 
                    
                    # 獲取腳注內部的 Range，並智能寫入內容
                    self.smart_insert_com(fn.Range, note_content)
                
                cursor = note_idx
            
            # === 3. 寫入剩餘正文 ===
            remaining_text = self.base_full_text[cursor:]
            if remaining_text:
                rng = doc.Content
                rng.Collapse(0)
                self.smart_insert_com(rng, remaining_text)
            
            self.progress_signal.emit(100)
            self.status_signal.emit("正在嵌入字體並保存...")
            
            # 【關鍵】開啟嵌入字體
            doc.EmbedTrueTypeFonts = True
            
            abs_path = os.path.abspath(self.output_path)
            doc.SaveAs(FileName=abs_path, FileFormat=16)
            self.finished_signal.emit(abs_path)
            
        except Exception as e:
            import traceback
            traceback.print_exc()
            self.error_signal.emit(str(e))
        
        finally:
            if doc:
                try: doc.Close(False) 
                except: pass
            if word:
                try: 
                    word.ScreenUpdating = True
                    word.Quit() 
                except: pass
            pythoncom.CoUninitialize()

# ==========================================
# 處理複雜訛誤模塊
# ==========================================
class SplitDialog(QDialog):
    """
    [重構版]：動態拆分編輯器
    支持添加任意數量的拆分段，以處理「脫+訛+脫」等複雜情況。
    """
    def __init__(self, base_text, wit_text, initial_data=None, parent=None):
        super().__init__(parent)
        self.setWindowTitle("複雜訛誤精細化切分")
        self.resize(600, 500)
        self.base_origin = base_text
        self.wit_origin = wit_text
        self.segments_data = [] # 最終結果列表
        
        self.init_ui(initial_data)

    def init_ui(self, initial_data):
        layout = QVBoxLayout(self)
        
        # 1. 頂部信息
        top_group = QGroupBox("原始異文塊信息")
        top_lay = QFormLayout(top_group)
        top_lay.addRow("原始底本:", QLabel(f"【{self.base_origin}】 (長度: {len(self.base_origin)})"))
        top_lay.addRow("原始校本:", QLabel(f"【{self.wit_origin}】 (長度: {len(self.wit_origin)})"))
        layout.addWidget(top_group)
        
        # 2. 動態編輯區 (放入滾動區域，防止條目過多撐爆窗口)
        self.scroll = QScrollArea()
        self.scroll.setWidgetResizable(True)
        self.scroll_content = QWidget()
        self.rows_layout = QVBoxLayout(self.scroll_content)
        self.rows_layout.setAlignment(Qt.AlignmentFlag.AlignTop)
        self.scroll.setWidget(self.scroll_content)
        
        # 標題頭
        header_lay = QHBoxLayout()
        header_lay.addWidget(QLabel("底本片段"), stretch=1)
        header_lay.addWidget(QLabel("校本片段"), stretch=1)
        header_lay.addWidget(QLabel("預判類型"), stretch=0)
        header_lay.addWidget(QLabel("操作"), stretch=0)
        
        edit_group = QGroupBox("拆分詳情 (系統將自動校驗長度是否守恆)")
        edit_layout = QVBoxLayout(edit_group)
        edit_layout.addLayout(header_lay)
        edit_layout.addWidget(self.scroll)
        
        # 工具欄
        tool_bar = QHBoxLayout()
        btn_add = QPushButton(" + 添加一段 ")
        btn_add.clicked.connect(lambda: self.add_segment_row("", ""))
        tool_bar.addWidget(btn_add)
        tool_bar.addStretch()
        edit_layout.addLayout(tool_bar)
        
        layout.addWidget(edit_group)

        # 3. 初始化數據填充
        if initial_data:
            # 如果有之前的編輯記錄，回填
            for seg in initial_data:
                self.add_segment_row(seg['base'], seg['wit'])
        else:
            # 默認：填充為一整塊 (純訛狀態)
            self.add_segment_row(self.base_origin, self.wit_origin)

        # 4. 底部驗證與按鈕
        self.lbl_status = QLabel("正在編輯...")
        self.lbl_status.setStyleSheet("color: #666;")
        layout.addWidget(self.lbl_status)
        
        btn_box = QHBoxLayout()
        btn_ok = QPushButton("確定並保存")
        btn_ok.clicked.connect(self.check_and_accept)
        btn_cancel = QPushButton("取消")
        btn_cancel.clicked.connect(self.reject)
        
        btn_box.addStretch()
        btn_box.addWidget(btn_ok)
        btn_box.addWidget(btn_cancel)
        layout.addLayout(btn_box)

    def add_segment_row(self, base_val, wit_val):
        """添加一行編輯控件"""
        row_widget = QWidget()
        row_lay = QHBoxLayout(row_widget)
        row_lay.setContentsMargins(0, 2, 0, 2)
        
        input_b = QLineEdit(base_val)
        input_w = QLineEdit(wit_val)
        lbl_type = QLabel("待定")
        lbl_type.setFixedWidth(40)
        lbl_type.setStyleSheet("color: #888; font-size: 12px;")
        
        # 自動更新類型顯示
        def update_type():
            b, w = input_b.text(), input_w.text()
            if not b and not w: t = "無效"
            elif not b: t = "衍"
            elif not w: t = "脫"
            elif b == w: t = "同"
            else: t = "訛"
            lbl_type.setText(t)
            
            # 特殊顏色
            color_map = {'衍': 'green', '脫': 'blue', '訛': 'red', '同': 'gray'}
            lbl_type.setStyleSheet(f"color: {color_map.get(t, 'black')}; font-weight: bold;")
        
        input_b.textChanged.connect(update_type)
        input_w.textChanged.connect(update_type)
        update_type() # 初始化觸發一次
        
        btn_del = QPushButton("×")
        btn_del.setFixedSize(24, 24)
        btn_del.setStyleSheet("color: red; font-weight: bold;")
        btn_del.clicked.connect(lambda: self.remove_row(row_widget))
        
        row_lay.addWidget(input_b)
        row_lay.addWidget(input_w)
        row_lay.addWidget(lbl_type)
        row_lay.addWidget(btn_del)
        
        self.rows_layout.addWidget(row_widget)

    def remove_row(self, widget):
        widget.setParent(None)
        widget.deleteLater()

    def check_and_accept(self):
        """校驗拆分後的總和是否等於原文"""
        temp_segments = []
        full_base = ""
        full_wit = ""
        
        # 遍歷所有行
        for i in range(self.rows_layout.count()):
            widget = self.rows_layout.itemAt(i).widget()
            if not widget: continue
            
            # 獲取 LineEdit
            # 佈局順序: base(0), wit(1), label(2), btn(3)
            # 為了保險，我們查找子控件
            inputs = widget.findChildren(QLineEdit)
            if len(inputs) != 2: continue
            
            b_text = inputs[0].text()
            w_text = inputs[1].text()
            
            # 過濾空行
            if not b_text and not w_text: continue
            
            full_base += b_text
            full_wit += w_text
            
            # 判斷類型
            t = "訛"
            if not b_text: t = "衍"
            elif not w_text: t = "脫"
            elif b_text == w_text: t = "Match" # 允許用戶拆出 Match (雖然少見但合理)
            
            temp_segments.append({'base': b_text, 'wit': w_text, 'type': t})

        # 校驗
        if full_base != self.base_origin:
            self.lbl_status.setText(f"錯誤：底本拆分總和【{full_base}】與原文不符！")
            self.lbl_status.setStyleSheet("color: red; font-weight: bold;")
            return
            
        if full_wit != self.wit_origin:
            self.lbl_status.setText(f"錯誤：校本拆分總和【{full_wit}】與原文不符！")
            self.lbl_status.setStyleSheet("color: red; font-weight: bold;")
            return
        
        self.segments_data = temp_segments
        self.accept()

class ComplexAuditPanel(QWidget):
    """
    [重構版]：複雜訛誤審覈面板 (支持多級撤銷 + 審覈狀態過濾)
    """
    data_changed_signal = pyqtSignal()
    audit_finished_signal = pyqtSignal(bool) # True=完成/解鎖, False=未完成/鎖定

    def __init__(self):
        super().__init__()
        self.all_data = {}
        self.current_wit = ""
        self.edit_cache = {} 
        self.history_stack = [] # 撤銷棧：存儲應用前的 results 列表
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout(self)
        
        # 頂部工具欄
        top_bar = QHBoxLayout()
        info = QLabel("【文本長度不一致訛文模塊診斷】點擊「編輯」修正或確認文本錯誤類型，點擊「應用當前頁面修改」可改變異文對照、校勘記等處記録信息。\n！注意！每次點擊「開始校勘」都會使此部分進行的修改失效！請慎重點擊！\n！！注意！！建議先完成「自定義規則」並重新運行校勘程序後再進行複雜訛誤分割！自定義規則後再次運行程序會導致此前的修正或確認失效！")
        info.setStyleSheet("color: #5D4037; font-family: 'KaiTi';")
        top_bar.addWidget(info)
        top_bar.addStretch()
        
        # 撤銷按鈕 (默認禁用)
        self.btn_undo = QPushButton(" ↶ 撤銷上一步應用 ")
        self.btn_undo.setEnabled(False)
        self.btn_undo.setStyleSheet("""
            QPushButton { background-color: #E0E0E0; color: #333; font-family: 'KaiTi'; border-radius: 4px; padding: 6px 12px; }
            QPushButton:hover { background-color: #D5D5D5; }
            QPushButton:disabled { color: #AAA; background-color: #F0F0F0; }
        """)
        self.btn_undo.clicked.connect(self.undo_last_apply)
        top_bar.addWidget(self.btn_undo)

        # 應用按鈕 (莫蘭迪豆灰綠)
        self.btn_apply = QPushButton(" 應用當前頁面修改 ")
        self.btn_apply.setToolTip("將當前列表中的所有編輯過的條目標記為已審覈。")
        self.btn_apply.setStyleSheet("""
            QPushButton { 
                background-color: #889C8C; color: #FFFFFF; font-family: 'KaiTi'; font-weight: bold; 
                padding: 6px 15px; border-radius: 4px; border: 1px solid #758A7B;
            }
            QPushButton:hover { background-color: #99AD9D; }
            QPushButton:pressed { background-color: #758A7B; }
        """)
        self.btn_apply.clicked.connect(self.apply_changes)
        top_bar.addWidget(self.btn_apply)
        
        layout.addLayout(top_bar)

        # 表格
        self.table = QTableWidget()
        self.table.setColumnCount(5)
        self.table.setHorizontalHeaderLabels(["底本原文回查", "底本文字", "校本文字", "當前處理狀態", "操作"])
        self.table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.Stretch)
        self.table.horizontalHeader().setSectionResizeMode(3, QHeaderView.ResizeMode.Stretch)
        self.table.horizontalHeader().setSectionResizeMode(4, QHeaderView.ResizeMode.ResizeToContents)
        layout.addWidget(self.table)

    def load_data(self, wit_name, full_data, base_clean_text, wit_clean_text):
        # [核心修正 1]：判斷是否切換了校本
        # 只有當用戶真的切換到另一個文件時，才清空撤銷歷史
        # 如果只是當前文件應用了修改導致的刷新，則保留歷史
        if wit_name != self.current_wit:
            self.history_stack = [] 
            self.btn_undo.setEnabled(False)
        
        self.current_wit = wit_name
        self.full_data = full_data
        self.wit_data = full_data[wit_name]
        self.base_text = base_clean_text
        self.wit_text = wit_clean_text
        
        # [核心修正 2]：根據當前歷史棧的狀態，恢復按鈕的可用性
        # 如果棧裡有東西（比如剛應用完），按鈕就會保持亮起
        self.btn_undo.setEnabled(bool(self.history_stack))
        
        self.edit_cache = {} 
        self.render_table()

    def render_table(self):
        self.table.setRowCount(0)
        results = self.wit_data['results']
        
        # 篩選候選者
        # 條件：訛 + 長度不一 + 【未被審覈】
        candidates = []
        for i, r in enumerate(results):
            if r.get('is_audited', False): continue # 過濾掉已審覈的
            
            if r['type'] == '訛' and r['base_clean_len'] != r['wit_clean_len']:
                candidates.append((i, r))
        
        # [核心邏輯]：檢查是否為空 -> 解鎖下一步
        if not candidates:
            # 只有當確實沒有候選者，且結果列表不為空（避免剛初始化時誤判）時解鎖
            self.audit_finished_signal.emit(True)
            # 顯示完成提示 (可選，在表格里顯示)
        else:
            self.audit_finished_signal.emit(False)

        for row_idx, (orig_idx, r) in enumerate(candidates):
            self.table.insertRow(row_idx)
            self.edit_cache[row_idx] = {
                'status': 'original', 'segments': [], 
                'orig_rec': r, 'orig_list_idx': orig_idx
            }

            # 0. 上下文
            start_pos = r['idx']
            end_pos = start_pos + r['base_clean_len']
            ctx_pre = self.base_text[max(0, start_pos - 10) : start_pos]
            ctx_post = self.base_text[end_pos : min(len(self.base_text), end_pos + 10)]
            ctx_html = f"<span style='color:#999'>{ctx_pre}</span><span style='color:#B74639; font-weight:bold;'>【{r['base_clean']}】</span><span style='color:#999'>{ctx_post}</span>"
            lbl_ctx = QLabel(ctx_html)
            self.table.setCellWidget(row_idx, 0, lbl_ctx)
            
            # [修改點]：設置只讀 (ItemIsEditable = False)
            # 1. 底本 (存隱藏數據)
            item_base = QTableWidgetItem(r['base_clean'])
            item_base.setData(Qt.ItemDataRole.UserRole, orig_idx)
            item_base.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            item_base.setFlags(item_base.flags() & ~Qt.ItemFlag.ItemIsEditable) # 只讀
            self.table.setItem(row_idx, 1, item_base)
            
            # 2. 校本
            item_wit = QTableWidgetItem(r['wit_clean'])
            item_wit.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            item_wit.setFlags(item_wit.flags() & ~Qt.ItemFlag.ItemIsEditable) # 只讀
            self.table.setItem(row_idx, 2, item_wit)
            
            # 3. 狀態
            self.update_status_cell(row_idx)
            
            # 4. 操作
            self.update_action_cell(row_idx)

    def update_status_cell(self, row_idx):
        cache = self.edit_cache[row_idx]
        if cache['status'] == 'original':
            text = "當前判定：訛"
            color = QColor(255, 255, 255) 
        else:
            segs = cache['segments']
            preview = []
            for s in segs:
                if s['type'] == '訛': preview.append(f"訛({s['base']}:{s['wit']})")
                elif s['type'] == '衍': preview.append(f"衍({s['wit']})")
                elif s['type'] == '脫': preview.append(f"脫({s['base']})")
                else: preview.append(s['type'])
            text = "已修改: " + " + ".join(preview)
            color = QColor("#F5F5DC") 
        
        item = QTableWidgetItem(text)
        item.setBackground(color)
        item.setToolTip(text)
        item.setFlags(item.flags() & ~Qt.ItemFlag.ItemIsEditable) # 只讀
        self.table.setItem(row_idx, 3, item)
        
        for c in [1, 2]:
            it = self.table.item(row_idx, c)
            if it: it.setBackground(color)

    def update_action_cell(self, row_idx):
        widget = QWidget()
        lay = QHBoxLayout(widget)
        lay.setContentsMargins(4, 2, 4, 2)
        lay.setSpacing(6) 
        
        cache = self.edit_cache[row_idx]
        
        btn_edit = QPushButton("編輯")
        btn_edit.setFixedSize(60, 26) 
        btn_edit.clicked.connect(lambda: self.open_split_dialog(row_idx))
        lay.addWidget(btn_edit)
        
        if cache['status'] == 'modified':
            btn_reset = QPushButton("重置")
            btn_reset.setFixedSize(60, 26)
            btn_reset.setStyleSheet("color: #C0392B; font-weight: bold;") 
            btn_reset.clicked.connect(lambda: self.reset_row(row_idx))
            lay.addWidget(btn_reset)
            
        self.table.setCellWidget(row_idx, 4, widget)
        self.table.resizeColumnToContents(4)

    def reset_row(self, row_idx):
        self.edit_cache[row_idx]['status'] = 'original'
        self.edit_cache[row_idx]['segments'] = []
        self.update_status_cell(row_idx)
        self.update_action_cell(row_idx)

    def open_split_dialog(self, row_idx):
        cache = self.edit_cache[row_idx]
        r = cache['orig_rec']
        init_data = cache['segments'] if cache['status'] == 'modified' else None
        
        dlg = SplitDialog(r['base_clean'], r['wit_clean'], init_data, self)
        if dlg.exec():
            self.edit_cache[row_idx]['segments'] = dlg.segments_data
            self.edit_cache[row_idx]['status'] = 'modified'
            self.update_status_cell(row_idx)
            self.update_action_cell(row_idx)

    def apply_changes(self):
        """
        [邏輯修正版] 應用修改：
        嚴格只應用「已修改 (modified)」的條目。
        未修改的條目將保留在表格中，等待用戶後續處理。
        """
        # 1. 檢查是否有已修改的行 (只看黃色行)
        modified_rows = [k for k, v in self.edit_cache.items() if v['status'] == 'modified']
        
        if not modified_rows:
            QMessageBox.information(self, "提示", "當前沒有「已編輯」的條目可應用。\n請先點擊「編輯」對條目進行處理（或確認）。")
            return
            
        reply = QMessageBox.question(self, "確認應用", 
                                     f"即將應用 {len(modified_rows)} 條已修改的記錄。\n應用後，這些條目將從列表中移除。\n未修改的條目將保留。",
                                     QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.No: return

        # 2. 【入棧】：深拷貝當前結果列表，存入歷史 (用於撤銷)
        current_snapshot = copy.deepcopy(self.wit_data['results'])
        self.history_stack.append(current_snapshot)
        self.btn_undo.setEnabled(True)

        # 3. 計算數據更新
        results = self.wit_data['results']
        replacements = {} # { 原始列表索引 : [新記錄列表] }
        
        # 僅收集 edit_cache 中標記為 modified 的數據
        for row_idx in modified_rows:
            cache = self.edit_cache[row_idx]
            
            # 從第 1 列獲取隱藏的原始索引
            item = self.table.item(row_idx, 1)
            if not item: continue
            orig_idx = item.data(Qt.ItemDataRole.UserRole)
            
            segments = cache['segments']
            orig_rec = cache['orig_rec']
            new_records = []
            
            curr_b_idx = orig_rec['idx']
            curr_w_idx = orig_rec['wit_idx']
            
            for seg in segments:
                new_rec = orig_rec.copy()
                new_rec['type'] = seg['type']
                new_rec['base_clean'] = seg['base']
                new_rec['wit_clean'] = seg['wit']
                new_rec['base_clean_len'] = len(seg['base'])
                new_rec['wit_clean_len'] = len(seg['wit'])
                new_rec['idx'] = curr_b_idx
                new_rec['wit_idx'] = curr_w_idx
                new_rec['is_audited'] = True # 標記為已審覈 (將從表中消失)
                
                curr_b_idx += len(seg['base'])
                curr_w_idx += len(seg['wit'])
                
                new_records.append(new_rec)
            
            replacements[orig_idx] = new_records

        # 4. 重構 results 列表
        final_results = []
        for i, r in enumerate(results):
            if i in replacements:
                # 情況A: 是用戶修改過的 -> 替換為新記錄 (且帶有 audited 標記)
                final_results.extend(replacements[i])
            else:
                # 情況B: 用戶沒動過的 -> 保持原樣 (不加 audited 標記，下次渲染還會出現)
                final_results.append(r)

        # 5. 寫回數據
        self.wit_data['results'] = final_results
        
        from collections import Counter
        self.wit_data['stats'] = Counter([r['type'] for r in final_results if r['type'] != 'Match'])
        
        # 4. [關鍵順序]：先彈出成功提示
        # 這是一個模態窗口，代碼會在這裡暫停，直到用戶點擊 OK
        QMessageBox.information(self, "成功", f"已應用 {len(modified_rows)} 條修改。")
        
        # 5. [關鍵順序]：用戶關閉彈窗後，再刷新界面
        self.edit_cache = {}
        self.render_table() # 如果表格變空，這裡面會發出 audit_finished_signal -> 觸發 MainWindow 的解鎖彈窗
        self.data_changed_signal.emit()

    def undo_last_apply(self):
        """撤銷上一步應用"""
        if not self.history_stack: return
        
        # 1. 彈出歷史快照
        prev_snapshot = self.history_stack.pop()
        
        # 2. 還原數據
        self.wit_data['results'] = prev_snapshot
        
        # 刷新統計
        from collections import Counter
        self.wit_data['stats'] = Counter([r['type'] for r in prev_snapshot if r['type'] != 'Match'])
        
        # 3. 刷新界面
        self.edit_cache = {}
        self.render_table()
        self.data_changed_signal.emit()
        
        # 更新按鈕狀態
        if not self.history_stack:
            self.btn_undo.setEnabled(False)
            
        QMessageBox.information(self, "撤銷", "已撤銷上一步應用，條目已恢復。")

class DisplacementMatchPanel(QWidget):
    """
    [修改版]：錯簡匹配
    功能：在同一個校本內，匹配相距較遠的「大段脫文」與「大段衍文」。
    修改內容：優化表格列結構（增加語境，刪除狀態），顯示省略號上下文。
    """
    def __init__(self):
        super().__init__()
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout(self)
        
        info = QLabel("【錯簡識別】系統自動匹配「脫文」與「衍文」。長度 > 10 字且文字匹配度 > 80%的片段將被列出。")
        info.setStyleSheet("color: #5D4037; font-family: 'KaiTi';")
        layout.addWidget(info)
        
        self.table = QTableWidget()
        # [修改]：改為 5 列，刪除狀態列，增加語境列
        self.table.setColumnCount(5)
        self.table.setHorizontalHeaderLabels(["匹配度", "脫文語境 (底本)", "脫文內容", "衍文語境 (校本)", "衍文內容"])
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        # 匹配度列可以設置窄一點
        self.table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.ResizeToContents)
        layout.addWidget(self.table)

    def load_data(self, wit_name, full_data, base_text, wit_text):
        """
        [修改]：新增 base_text 和 wit_text 參數用於提取上下文
        """
        self.table.setRowCount(0)
        if wit_name not in full_data: return
        results = full_data[wit_name]['results']
        
        # 1. 收集所有長脫文與長衍文
        MIN_LEN = 10
        deletions = [] 
        insertions = [] 
        
        for r in results:
            if r['type'] == '脫' and len(r['base_clean']) >= MIN_LEN:
                deletions.append(r)
            elif r['type'] == '衍' and len(r['wit_clean']) >= MIN_LEN:
                insertions.append(r)
                
        # 2. 兩兩比對
        matches = []
        for d in deletions:
            for i in insertions:
                s = difflib.SequenceMatcher(None, d['base_clean'], i['wit_clean'])
                ratio = s.ratio()
                
                if ratio > 0.8: 
                    matches.append({
                        'ratio': ratio,
                        'del': d,
                        'ins': i
                    })
                    
        # 3. 按相似度降序排列
        matches.sort(key=lambda x: x['ratio'], reverse=True)
        
        # 4. 渲染表格
        for row_idx, m in enumerate(matches):
            self.table.insertRow(row_idx)
            
            # --- 0. 相似度 ---
            item_ratio = QTableWidgetItem(f"{m['ratio']:.1%}")
            item_ratio.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            item_ratio.setForeground(QColor("#2E8B57"))
            item_ratio.setFont(QFont("Arial", 10, QFont.Weight.Bold))
            self.table.setItem(row_idx, 0, item_ratio)
            
            # --- 1 & 2. 脫文 (底本上下文 + 內容) ---
            d_rec = m['del']
            d_start = d_rec['idx']
            d_end = d_start + d_rec['base_clean_len']
            
            # 獲取底本上下文 (前5 + …… + 後5)
            d_pre = base_text[max(0, d_start - 5) : d_start]
            d_post = base_text[d_end : min(len(base_text), d_end + 5)]
            d_context = f"{d_pre}……{d_post}"
            
            self.table.setItem(row_idx, 1, QTableWidgetItem(d_context))
            # 內容完整顯示
            item_d_content = QTableWidgetItem(d_rec['base_clean'])
            item_d_content.setToolTip(d_rec['base_clean']) # 鼠標懸停顯示全文
            self.table.setItem(row_idx, 2, item_d_content)
            
            # --- 3 & 4. 衍文 (校本上下文 + 內容) ---
            i_rec = m['ins']
            i_start = i_rec['wit_idx']
            i_end = i_start + i_rec['wit_clean_len']
            
            # 獲取校本上下文 (前5 + …… + 後5)
            i_pre = wit_text[max(0, i_start - 5) : i_start]
            i_post = wit_text[i_end : min(len(wit_text), i_end + 5)]
            i_context = f"{i_pre}……{i_post}"
            
            self.table.setItem(row_idx, 3, QTableWidgetItem(i_context))
            # 內容完整顯示
            item_i_content = QTableWidgetItem(i_rec['wit_clean'])
            item_i_content.setToolTip(i_rec['wit_clean'])
            self.table.setItem(row_idx, 4, item_i_content)

# ==========================================
# [重構] 版本源流考察面板 (Scipy + NetworkX 增強版)
# ==========================================
import itertools
import math
import networkx as nx
from scipy.cluster import hierarchy
from scipy.spatial.distance import squareform

class LineageHelpDialog(QDialog):
    """
    [新增組件] 圖譜解讀說明彈窗
    """
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("📊 如何解讀版本源流圖譜？")
        self.resize(700, 600)
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout(self)
        
        browser = QTextBrowser()
        browser.setOpenExternalLinks(True)
        browser.setStyleSheet("font-family: 'SimSun'; font-size: 15px; line-height: 1.6;")
        
        # HTML 內容：使用傳統校勘學術語 + 現代統計輔助
        html_content = """
        <h2 style="color:#5D4037;">版本源流圖譜解讀指南</h2>
        <p>本模塊基於各版本間的<b>「共同訛誤」（Shared Errors）</b>進行統計。在校勘學上，若兩個版本在同一位置出現完全相同的錯誤（如特定的避諱缺筆、獨特的訛字），通常意味著它們源自同一個<b>「母本」（Common Source）</b>或屬於同一個<b>「傳承系統」（Lineage）</b>。</p>
        <hr>
        
        <h3 style="color:#B74639;">1. 熱力矩陣 (Heatmap Matrix)</h3>
        <p><b>含義：</b>直觀展示版本間的「親疏關係」。</p>
        <ul>
            <li><b>顏色越深（紅）：</b>代表兩個版本間的共同訛誤數量越多，關係越親密，可能互為父子本或兄弟本。</li>
            <li><b>顏色越淺（白/黃）：</b>代表共同訛誤很少，關係疏遠，可能屬於不同的版本系統。</li>
        </ul>
        <p><i>舉例：若「甲本」與「乙本」的交匯處為深紅色，說明二者高度一致，校勘時可視為同一系統的代表。</i></p>
        
        <h3 style="color:#2E5C8A;">2. 譜系聚類樹 (Dendrogram)</h3>
        <p><b>含義：</b>模擬生物演化樹，推測版本的<b>「同源關係」</b>與<b>「分化路徑」</b>。</p>
        <ul>
            <li><b>橫軸（遺傳距離/差異度）：</b>線條越長，代表差異越大。</li>
            <li><b>分叉點（Clade）：</b>線條在左側越早匯合，代表它們越早歸為一類（聚類）。</li>
        </ul>
        <p><i>舉例：若樹狀圖顯示 (甲, 乙) 先匯合，然後再與 (丙) 匯合。這暗示「甲、乙」可能源自同一母本，而「丙」則是另一個系統的本子。</i></p>
        
        <h3 style="color:#5C7A62;">3. 親緣網絡 (Network Graph)</h3>
        <p><b>含義：</b>基於力導向算法（Force-directed Layout）的動態拓撲圖。</p>
        <ul>
            <li><b>結點吸附：</b>共享訛誤多的版本會自動「吸」在一起，形成<b>「家族叢集」</b>。</li>
            <li><b>連線粗細：</b>線條越粗，代表共同訛誤越多，證據越確鑿。</li>
        </ul>
        <p><i>用途：快速識別哪些版本是核心主流（位於中心），哪些是邊緣版本或混合版本。</i></p>
        
        <hr>
        <p style="color:#666; font-size:13px;">* 註：統計僅包含您在左側勾選的錯誤類型（如訛、脫等）。異體字是否計入源流判斷，視具體研究需求而定。</p>
        """
        
        browser.setHtml(html_content)
        layout.addWidget(browser)
        
        btn_close = QPushButton("我明白了")
        btn_close.clicked.connect(self.accept)
        layout.addWidget(btn_close, alignment=Qt.AlignmentFlag.AlignRight)

# ==========================================
# [重構] 版本源流考察面板 (最終修正版)
# ==========================================
import itertools
import math
import networkx as nx
from scipy.cluster import hierarchy
from scipy.spatial.distance import squareform

class LineageAnalysisPanel(QWidget):
    """
    [功能]：版本源流考察與可視化 (專業版)
    [邏輯]：計算多版本間的「共同訛誤」，繪製熱力圖、樹狀圖、網絡圖。
    [依賴]：matplotlib, scipy, networkx
    """
    def __init__(self):
        super().__init__()
        self.all_data = {}
        self.base_clean_text = "" # 這是底本純文本，必須確保被正確傳入
        # 顏色常量
        self.colors = {'訛': '#B74639', '脫': '#2E5C8A', '衍': '#5C7A62', '倒': '#E6B450', '異': '#8E44AD'}
        self.init_ui()

    def init_ui(self):
        layout = QHBoxLayout(self)
        
        # === 左側：控制區 ===
        ctrl_group = QGroupBox("設置與篩選")
        ctrl_group.setFixedWidth(220)
        ctrl_layout = QVBoxLayout(ctrl_group)
        
        ctrl_layout.addWidget(QLabel("1. 選擇參與比對的校本："))
        self.list_versions = QListWidget()
        self.list_versions.setSelectionMode(QListWidget.SelectionMode.MultiSelection)
        self.list_versions.itemSelectionChanged.connect(self.run_analysis)
        ctrl_layout.addWidget(self.list_versions)
        
        ctrl_layout.addWidget(QLabel("2. 納入統計的類型："))
        self.cbs_types = {}
        for t in ['訛', '脫', '衍', '倒', '異']:
            cb = QCheckBox(t)
            cb.setChecked(True)
            cb.stateChanged.connect(self.run_analysis) # 狀態改變即刷新
            self.cbs_types[t] = cb
            ctrl_layout.addWidget(cb)
            
        ctrl_layout.addStretch()
        
        # [新增] 說明按鈕 (左側底部)
        self.btn_help = QPushButton("📊 如何解讀這些圖譜？")
        self.btn_help.setStyleSheet("""
            QPushButton { color: #2E5C8A; font-weight: bold; border: 1px solid #2E5C8A; border-radius: 4px; padding: 5px; }
            QPushButton:hover { background-color: #EBF5FB; }
        """)
        self.btn_help.clicked.connect(self.show_help)
        ctrl_layout.addWidget(self.btn_help)
        
        layout.addWidget(ctrl_group)
        
        # === 右側：可視化與詳情 ===
        right_splitter = QSplitter(Qt.Orientation.Vertical)
        
        # 上半部：圖表 Tab
        self.viz_tabs = QTabWidget()
        self.tab_heatmap = QWidget(); self.viz_tabs.addTab(self.tab_heatmap, "熱力矩陣 (親疏關係)")
        self.tab_tree = QWidget();    self.viz_tabs.addTab(self.tab_tree, "譜系聚類 (Dendrogram)")
        self.tab_network = QWidget(); self.viz_tabs.addTab(self.tab_network, "親緣網絡 (Force Layout)")
        
        # 初始化畫布
        self.init_canvases()
        
        right_splitter.addWidget(self.viz_tabs)
        
        # 下半部：詳細列表
        detail_group = QGroupBox("共同訛誤詳細條目 (選中版本間的交集)")
        detail_layout = QVBoxLayout(detail_group)
        self.table_detail = QTableWidget()
        self.table_detail.setColumnCount(5)
        # [修改] 第一列改為「上下文」，第三列「底本」
        self.table_detail.setHorizontalHeaderLabels(["上下文 (定位)", "類型", "底本", "共同校本內容", "涉及版本"])
        self.table_detail.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        self.table_detail.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.ResizeToContents) # 上下文自適應
        self.table_detail.setAlternatingRowColors(True)
        detail_layout.addWidget(self.table_detail)
        
        right_splitter.addWidget(detail_group)
        
        # 設置右側比例
        right_splitter.setStretchFactor(0, 6)
        right_splitter.setStretchFactor(1, 4)
        
        layout.addWidget(right_splitter)

    def init_canvases(self):
        """初始化三個 Matplotlib 畫布"""
        if not HAS_MATPLOTLIB: return
        
        # 1. 熱力圖
        self.fig_heat = Figure(figsize=(5, 4), dpi=100)
        self.canvas_heat = FigureCanvas(self.fig_heat)
        l1 = QVBoxLayout(self.tab_heatmap); l1.addWidget(self.canvas_heat)
        
        # 2. 樹狀圖
        self.fig_tree = Figure(figsize=(5, 4), dpi=100)
        self.canvas_tree = FigureCanvas(self.fig_tree)
        l2 = QVBoxLayout(self.tab_tree); l2.addWidget(self.canvas_tree)
        
        # 3. 網絡圖
        self.fig_net = Figure(figsize=(5, 4), dpi=100)
        self.canvas_net = FigureCanvas(self.fig_net)
        l3 = QVBoxLayout(self.tab_network); l3.addWidget(self.canvas_net)

    def show_help(self):
        """顯示幫助彈窗"""
        dlg = LineageHelpDialog(self)
        dlg.exec()

    def load_data(self, all_data, base_text):
        """外部調用入口，強制更新 base_text"""
        self.all_data = all_data
        self.base_clean_text = base_text # [關鍵修復] 確保拿到最新的底本
        
        # 刷新列表，保持原有選中狀態
        selected_names = [item.text() for item in self.list_versions.selectedItems()]
        self.list_versions.clear()
        
        for name in all_data.keys():
            item = QListWidgetItem(name)
            self.list_versions.addItem(item)
            if name in selected_names:
                item.setSelected(True)
        
        # 默認全選（如果是第一次加載）
        if not selected_names and self.list_versions.count() > 0:
            self.list_versions.selectAll()
            
        self.run_analysis()

    def get_fingerprints(self, wit_name, active_types):
        """生成指紋: (idx, type, content)"""
        data = self.all_data[wit_name]
        fingerprints = set()
        raw_list = data['results']
        
        for r in raw_list:
            t = r['type']
            if t == 'Match': continue
            if t not in active_types: continue
            
            # [嚴格定義]：位置 + 類型 + 內容
            content = r['wit_clean']
            # 如果是「脫」，內容可能為空，這也是一種特徵
            key = (r['idx'], t, content)
            fingerprints.add(key)
            
        return fingerprints

    def run_analysis(self):
        """執行核心分析流程"""
        if not HAS_MATPLOTLIB or not self.all_data: return
        
        # 1. 獲取選中的版本
        selected_items = self.list_versions.selectedItems()
        if len(selected_items) < 2:
            self.fig_heat.clear(); self.canvas_heat.draw()
            self.fig_tree.clear(); self.canvas_tree.draw()
            self.fig_net.clear(); self.canvas_net.draw()
            self.table_detail.setRowCount(0)
            return
            
        versions = [item.text() for item in selected_items]
        n = len(versions)
        
        # 2. 獲取激活的類型
        active_types = [t for t, cb in self.cbs_types.items() if cb.isChecked()]
        
        # 3. 提取指紋
        ver_fps = {v: self.get_fingerprints(v, active_types) for v in versions}
        
        # 4. 計算矩陣
        sim_matrix = [[0]*n for _ in range(n)]
        max_shared = 0
        
        for i in range(n):
            for j in range(n):
                if i == j:
                    sim_matrix[i][j] = len(ver_fps[versions[i]]) 
                else:
                    common = ver_fps[versions[i]] & ver_fps[versions[j]]
                    count = len(common)
                    sim_matrix[i][j] = count
                    if count > max_shared: max_shared = count
        
        dist_matrix = []
        for i in range(n):
            row = []
            for j in range(n):
                if i == j: row.append(0)
                else: row.append((max_shared + 1) - sim_matrix[i][j])
            dist_matrix.append(row)

        # 5. 繪製圖表
        self.draw_heatmap(versions, sim_matrix)
        self.draw_tree_scipy(versions, dist_matrix)
        self.draw_network_nx(versions, sim_matrix)
        
        # 6. 更新詳情表
        self.update_detail_table(versions, ver_fps)

    def draw_heatmap(self, names, matrix):
        self.fig_heat.clear()
        ax = self.fig_heat.add_subplot(111)
        im = ax.imshow(matrix, cmap='OrRd')
        
        ax.set_xticks(range(len(names)))
        ax.set_yticks(range(len(names)))
        ax.set_xticklabels(names, rotation=45, ha="right", fontfamily='KaiTi')
        ax.set_yticklabels(names, fontfamily='KaiTi')
        
        for i in range(len(names)):
            for j in range(len(names)):
                text = ax.text(j, i, matrix[i][j], ha="center", va="center", color="black")
                
        ax.set_title("版本共同訛誤數量矩陣", fontfamily='KaiTi', fontsize=12)
        self.fig_heat.tight_layout()
        self.canvas_heat.draw()

    def draw_tree_scipy(self, names, dist_matrix_full):
        self.fig_tree.clear()
        ax = self.fig_tree.add_subplot(111)
        
        try:
            condensed_dist = squareform(dist_matrix_full)
        except:
            ax.text(0.5, 0.5, "數據不足以構建樹狀圖", ha='center', fontfamily='KaiTi')
            self.canvas_tree.draw()
            return

        linkage_matrix = hierarchy.linkage(condensed_dist, method='average')
        
        hierarchy.dendrogram(
            linkage_matrix, 
            labels=names, 
            ax=ax, 
            orientation='right',
            leaf_font_size=12
        )
        
        ax.set_title("版本同源關係樹 (UPGMA)", fontfamily='KaiTi', pad=15)
        ax.set_xlabel("遺傳距離 (差異度)", fontfamily='KaiTi')
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.spines['left'].set_visible(False)
        
        self.fig_tree.tight_layout()
        self.canvas_tree.draw()

    def draw_network_nx(self, names, sim_matrix):
        self.fig_net.clear()
        ax = self.fig_net.add_subplot(111)
        ax.axis('off')
        
        n = len(names)
        G = nx.Graph()
        G.add_nodes_from(names)
        
        max_shared = 0
        for i in range(n):
            for j in range(i+1, n):
                weight = sim_matrix[i][j]
                if weight > 0:
                    G.add_edge(names[i], names[j], weight=weight)
                    if weight > max_shared: max_shared = weight
        
        if max_shared == 0: max_shared = 1

        pos = nx.spring_layout(G, k=1.5, weight='weight', seed=42, iterations=50)
        
        nx.draw_networkx_nodes(G, pos, ax=ax, node_size=800, node_color='#D5F5E3', edgecolors='#5C7A62')
        
        for node, (x, y) in pos.items():
            ax.text(x, y, node, fontsize=10, ha='center', va='center', 
                    fontfamily='KaiTi', fontweight='bold', zorder=10)

        edges = G.edges(data=True)
        for u, v, d in edges:
            w = d['weight']
            width = (w / max_shared) * 4 + 0.5
            alpha = 0.3 + (w / max_shared) * 0.6
            nx.draw_networkx_edges(G, pos, ax=ax, edgelist=[(u, v)], 
                                   width=width, alpha=alpha, edge_color='#B74639')
            if w > max_shared * 0.1:
                (x1, y1) = pos[u]
                (x2, y2) = pos[v]
                mid_x, mid_y = (x1+x2)/2, (y1+y2)/2
                ax.text(mid_x, mid_y, str(w), fontsize=8, color='#2E5C8A', 
                        ha='center', va='center', 
                        bbox=dict(boxstyle='round,pad=0.1', fc='white', alpha=0.8, ec='none'))

        ax.set_title("共同訛誤家族叢集 (力導向)", fontfamily='KaiTi', y=0.95)
        self.canvas_net.draw()

    def update_detail_table(self, versions, ver_fps):
        """
        [核心修復] 更新詳情表
        1. 使用 Scheme A (前5...後5)
        2. 衍文顯示【衍文無需底本】
        3. 底本列顯示正確文字
        """
        self.table_detail.setRowCount(0)
        global_counter = defaultdict(list)
        for v in versions:
            for fp in ver_fps[v]:
                global_counter[fp].append(v)
        
        common_fps = [fp for fp, vs in global_counter.items() if len(vs) >= 2]
        common_fps.sort(key=lambda x: x[0])
        
        # 緩存底本長度，避免循環計算
        base_len = len(self.base_clean_text)
        
        for fp in common_fps:
            idx, t, content = fp
            v_list = global_counter[fp]
            row = self.table_detail.rowCount()
            self.table_detail.insertRow(row)
            
            # --- 1. 上下文 (Context) 計算 ---
            # 方案 A: 前五字……後五字 (中間不包含錯誤本身)
            start_pre = max(0, idx - 5)
            end_pre = idx
            
            # 對於非衍文，錯誤本身佔據了 base_clean 的位置，需要跳過
            # 但這裡我們沒有存 r['base_clean_len']，這是一個小麻煩
            # 解決方案：去原始數據裡反查一下 (雖然效率稍低，但為了準確性)
            # 或者：假設如果是 '衍'，base_len=0；如果是 '脫/訛/異/倒'，默認跳過1個字(不準確)
            # 最佳方案：從原始數據獲取長度。
            # 這裡我們做一個快速查找：
            base_occupy_len = 0
            if t == '衍':
                base_occupy_len = 0
            else:
                # 嘗試在第一個版本的原始結果中找到這個指紋對應的 base_len
                # 這是為了精確跳過底本上的錯誤字
                sample_ver = v_list[0]
                sample_res = self.all_data[sample_ver]['results']
                for sr in sample_res:
                    if sr['idx'] == idx and sr['type'] == t and sr['wit_clean'] == content:
                        base_occupy_len = sr['base_clean_len']
                        break
            
            start_post = idx + base_occupy_len
            end_post = min(base_len, start_post + 5)
            
            pre_text = self.base_clean_text[start_pre : end_pre]
            post_text = self.base_clean_text[start_post : end_post]
            
            context_str = f"{pre_text}……{post_text}"
            
            item_ctx = QTableWidgetItem(context_str)
            item_ctx.setToolTip(f"索引位置: {idx}") # 懸停顯示索引
            self.table_detail.setItem(row, 0, item_ctx)
            
            # --- 2. 類型 ---
            self.table_detail.setItem(row, 1, QTableWidgetItem(t))
            
            # --- 3. 底本文字 ---
            if t == '衍':
                base_char = "【衍文無需底本】"
                item_base = QTableWidgetItem(base_char)
                item_base.setForeground(QColor("#5C7A62")) # 綠色
                item_base.setFont(QFont("SimSun", 9, QFont.Weight.Bold))
            else:
                # 截取底本實際文字
                if idx < base_len:
                    base_char = self.base_clean_text[idx : idx + base_occupy_len]
                else:
                    base_char = "【無】" # 理論上不應發生
                item_base = QTableWidgetItem(base_char)
            
            self.table_detail.setItem(row, 2, item_base)
            
            # --- 4. 共同內容 ---
            self.table_detail.setItem(row, 3, QTableWidgetItem(content if content else "【無】"))
            
            # --- 5. 涉及版本 ---
            self.table_detail.setItem(row, 4, QTableWidgetItem("、".join(v_list)))

# ==========================================
# 5. 統計面板組件 
# ==========================================
class StatsPanel(QWidget):
    def __init__(self):
        super().__init__()
        self.all_data = {} 
        self.has_variant_split = False # 標記是否包含異體字數據
        # 定義顏色常量 
        self.colors = {
            '訛': '#B74639', 
            '脫': '#2E5C8A', 
            '衍': '#5C7A62', 
            '倒': '#E6B450',
            '異': '#8E44AD' 
        }
        self.init_ui()

# [新增] 監聽標籤頁切換，控制底部表格顯隱
    def on_tab_changed(self, index):
        # 獲取當前選中的 Tab 標題
        title = self.chart_tabs.tabText(index)
        
        # 如果標題包含 "源流"，就隱藏底部的 table_group
        if "版本源流考察" in title:
            if hasattr(self, 'table_group'):
                self.table_group.hide()
        else:
            # 切換回其他頁面時，顯示表格
            if hasattr(self, 'table_group'):
                self.table_group.show()

    def init_ui(self):
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(10, 10, 10, 10)
        main_layout.setSpacing(10)

        # ==========================================
        # 1. 頂部交互篩選區 
        # ==========================================
        filter_layout = QHBoxLayout()
        filter_layout.addWidget(QLabel("顯示類型篩選："))
        
        self.checkboxes = {}
        # 按順序創建複選框
        for key in ['訛', '脫', '衍', '倒', '異']:
            cb = QCheckBox(key)
            cb.setChecked(True)
            cb.setStyleSheet(f"font-weight: bold; color: {self.colors[key]}; font-family: 'KaiTi'; font-size: 14px;")
            # 連接信號：當勾選狀態改變時，刷新圖表
            cb.stateChanged.connect(self.refresh_charts)
            self.checkboxes[key] = cb
            filter_layout.addWidget(cb)
        
        filter_layout.addStretch()
        main_layout.addLayout(filter_layout)
        
        # ==========================================
        # 2. 主體分割區
        # ==========================================
        splitter = QSplitter(Qt.Orientation.Vertical)
        splitter.setHandleWidth(1) 
        
        gb_style = """
            QGroupBox {
                font-size: 14px;
                border: 1px solid #E0E0E0;
                border-radius: 6px;
                margin-top: 15px; 
                background-color: #FFFFFF; 
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                subcontrol-position: top left;
                padding: 0 5px;
                left: 10px;
                font-family: "KaiTi";
                font-weight: bold;
                color: #5D4037;
            }
        """

        # --- 上半部分：圖表區 ---
        self.chart_group = QGroupBox("  校勘數據可視化 ")
        self.chart_group.setStyleSheet(gb_style)
        self.chart_group.setMinimumHeight(350) 
        
        chart_layout = QVBoxLayout(self.chart_group)
        chart_layout.setContentsMargins(10, 30, 10, 10) 
        
        self.chart_tabs = QTabWidget()
        self.chart_tabs.setStyleSheet("QTabWidget::pane { border: 0; }")
        self.chart_tabs.currentChanged.connect(self.on_tab_changed)
        chart_layout.addWidget(self.chart_tabs)
        
        # Tab 1: 數量 (Bar Chart)
        self.tab_bar_chart = QWidget()
        layout_bar = QVBoxLayout(self.tab_bar_chart)
        layout_bar.setContentsMargins(0,0,0,0)
        if HAS_MATPLOTLIB:
            self.fig_bar = Figure(figsize=(5, 4), dpi=100)
            self.fig_bar.patch.set_facecolor('#FFFFFF')
            self.canvas_bar = FigureCanvas(self.fig_bar)
            layout_bar.addWidget(self.canvas_bar)
        else:
            layout_bar.addWidget(QLabel("未安裝 matplotlib"))
        self.chart_tabs.addTab(self.tab_bar_chart, "數量對比")

        # Tab 2: 分布 (Scatter Chart)
        self.tab_dist_chart = QWidget()
        layout_dist = QVBoxLayout(self.tab_dist_chart)
        layout_dist.setContentsMargins(0,0,0,0)
        if HAS_MATPLOTLIB:
            self.fig_dist = Figure(figsize=(5, 4), dpi=100)
            self.fig_dist.patch.set_facecolor('#FFFFFF')
            self.canvas_dist = FigureCanvas(self.fig_dist)
            layout_dist.addWidget(self.canvas_dist)
        else:
            layout_dist.addWidget(QLabel("未安裝 matplotlib"))
        self.chart_tabs.addTab(self.tab_dist_chart, "異文分布")

        # Tab 3: 綜合評估 (Radar Charts) 
        self.tab_radar_chart = QWidget()
        layout_radar = QHBoxLayout(self.tab_radar_chart) # 改為水平佈局
        layout_radar.setContentsMargins(0,0,0,0)
        
        if HAS_MATPLOTLIB:
            # 左側：4軸雷達圖
            self.fig_radar_4 = Figure(figsize=(4, 4), dpi=100)
            self.fig_radar_4.patch.set_facecolor('#FFFFFF')
            self.canvas_radar_4 = FigureCanvas(self.fig_radar_4)
            layout_radar.addWidget(self.canvas_radar_4, stretch=1)
            
            # 右側：5軸雷達圖
            self.fig_radar_5 = Figure(figsize=(4, 4), dpi=100)
            self.fig_radar_5.patch.set_facecolor('#FFFFFF')
            self.canvas_radar_5 = FigureCanvas(self.fig_radar_5)
            self.canvas_radar_5.setVisible(False) # 默認隱藏
            layout_radar.addWidget(self.canvas_radar_5, stretch=1)
        else:
            layout_radar.addWidget(QLabel("未安裝 matplotlib"))
            
        self.chart_tabs.addTab(self.tab_radar_chart, "綜合評估")

        self.lineage_panel = LineageAnalysisPanel()
        self.chart_tabs.addTab(self.lineage_panel, "版本源流考察")


        # --- 下半部分：表格區 ---
        self.table_group = QGroupBox("  詳細統計數據 ")
        self.table_group.setStyleSheet(gb_style)
        table_layout = QVBoxLayout(self.table_group)
        table_layout.setContentsMargins(10, 30, 10, 10)
        
        self.stats_table = QTableWidget()
        # 列數稍後動態設置
        self.stats_table.setAlternatingRowColors(True) 
        self.stats_table.verticalHeader().setVisible(False)
        self.stats_table.setStyleSheet("border: none; gridline-color: #E0E0E0;") 
        table_layout.addWidget(self.stats_table)

        splitter.addWidget(self.chart_group)
        splitter.addWidget(self.table_group)
        splitter.setStretchFactor(0, 3)
        splitter.setStretchFactor(1, 2)

        main_layout.addWidget(splitter)

        self.chart_tabs.currentChanged.connect(self.on_tab_changed)

    def update_data(self, all_data, has_variant_split=False, base_text=""):
        """
        更新數據入口
        :param all_data: 校勘數據
        :param has_variant_split: 是否勾選了「異體字單獨出校」
        """
        if not all_data: return
        self.all_data = all_data
        self.has_variant_split = has_variant_split

        # 1. 控制「異」複選框的顯示狀態
        if has_variant_split:
            self.checkboxes['異'].setVisible(True)
            self.checkboxes['異'].setChecked(True)
        else:
            self.checkboxes['異'].setChecked(False)
            self.checkboxes['異'].setVisible(False)

        # 2. 更新表格 (Table)
        self.update_table_content()

        # 3. 刷新圖表 (Charts)
        self.refresh_charts()

        # 4. [新增] 將最新的底本和數據傳遞給源流分析面板
        # 確保 LineagePanel 總是拿到最新的 base_text
        if hasattr(self, 'lineage_panel'):
            self.lineage_panel.load_data(all_data, base_text)

    def update_table_content(self):
        """根據數據動態刷新表格列和內容"""
        self.stats_table.setRowCount(0)
        
        # 根據是否分異體字，決定表頭
        if self.has_variant_split:
            headers = ["校本名稱", "訛", "脫", "衍", "倒", "異", "總計"]
            keys = ['訛', '脫', '衍', '倒', '異']
        else:
            headers = ["校本名稱", "訛", "脫", "衍", "倒", "總計"]
            keys = ['訛', '脫', '衍', '倒']
            
        self.stats_table.setColumnCount(len(headers))
        self.stats_table.setHorizontalHeaderLabels(headers)
        self.stats_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)

        for wit_name, data in self.all_data.items():
            stats = data['stats']
            row = self.stats_table.rowCount()
            self.stats_table.insertRow(row)
            
            # 名稱
            name_item = QTableWidgetItem(wit_name)
            name_item.setFont(QFont("KaiTi", 11, QFont.Weight.Bold)) 
            self.stats_table.setItem(row, 0, name_item)
            
            total = 0
            for i, k in enumerate(keys):
                val = stats.get(k, 0)
                item = QTableWidgetItem(str(val))
                item.setTextAlignment(Qt.AlignmentFlag.AlignCenter) 
                self.stats_table.setItem(row, i + 1, item)
                total += val
            
            # 總計
            total_item = QTableWidgetItem(str(total))
            total_item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            total_item.setFont(QFont("SimSun", 10, QFont.Weight.Bold))
            self.stats_table.setItem(row, len(headers)-1, total_item)

    def refresh_charts(self):
        """
        [可視化渲染]：根據當前複選框狀態，重新繪製所有圖表。
        [技術細節]：
        使用 Matplotlib 的面向對象 API (Object-Oriented API) 進行繪圖。
        針對中文字符在圖表中的顯示問題（如被切斷、重疊），使用了 `subplots_adjust` 進行版面微調。
        """
        if not HAS_MATPLOTLIB or not self.all_data: return
        
        names = list(self.all_data.keys())
        
        # 獲取當前激活的類型
        active_types = []
        possible_types = ['訛', '脫', '衍', '倒']
        if self.has_variant_split:
            possible_types.append('異')
            
        for t in possible_types:
            if self.checkboxes[t].isChecked():
                active_types.append(t)

        # 準備字體配置
        FONT_TITLE = {'family': 'KaiTi', 'size': 12}
        FONT_LABEL = {'family': 'KaiTi', 'size': 10}
        
        # ==========================
        # 1. 柱狀圖 (Bar Chart) 
        # ==========================
        self.fig_bar.clear()
        ax1 = self.fig_bar.add_subplot(111)
        x = np.arange(len(names))
        
        n_bars = len(active_types)
        if n_bars > 0:
            total_width = 0.8
            bar_width = total_width / n_bars
            
            for i, t in enumerate(active_types):
                vals = [self.all_data[name]['stats'].get(t, 0) for name in names]
                offset = (i - n_bars / 2 + 0.5) * bar_width
                ax1.bar(x + offset, vals, bar_width, label=t, color=self.colors[t], edgecolor='white')
        
            ax1.set_xticks(x)
            ax1.set_xticklabels(names, rotation=0, fontdict=FONT_LABEL)
            ax1.legend(loc='upper center', bbox_to_anchor=(0.5, -0.15), 
                       ncol=5, frameon=False, prop={'family': 'KaiTi'})
        else:
            ax1.text(0.5, 0.5, "未選擇數據類型", ha='center', va='center', fontdict=FONT_TITLE)

        ax1.set_ylabel('數量', fontname='SimSun')
        ax1.set_title('各本異文數量對比', fontdict=FONT_TITLE, pad=15)
        ax1.grid(axis='y', linestyle=':', alpha=0.5, color='#CCCCCC')
        ax1.spines['top'].set_visible(False)
        ax1.spines['right'].set_visible(False)
        
        self.fig_bar.subplots_adjust(left=0.15, right=0.95, top=0.9, bottom=0.25)
        self.canvas_bar.draw()

        # ==========================
        # 2. 分布圖 (Scatter Chart)
        # ==========================
        self.fig_dist.clear()
        ax2 = self.fig_dist.add_subplot(111)
        
        if n_bars > 0:
            for i, name in enumerate(names):
                results = self.all_data[name]['results']
                variants_list = self.all_data[name].get('variants', [])
                
                for t in active_types:
                    if t == '異':
                        indices = [v['idx'] for v in variants_list]
                    else:
                        indices = [r['idx'] for r in results if r['type'] == t]
                    
                    if indices:
                        ax2.scatter(indices, [i]*len(indices), marker='|', c=self.colors[t], s=120, alpha=0.9, label=t if i==0 else "")
                
                ax2.axhline(y=i, color='#E0E0E0', alpha=0.5, linewidth=8, zorder=0)

            ax2.set_yticks(range(len(names)))
            # 【優化】如果有文件名很長，這裡會自動處理，但最好依賴下方的 subplots_adjust
            ax2.set_yticklabels(names, fontdict=FONT_LABEL)
            ax2.set_xlabel('文字位置索引', fontname='SimSun')
            
            handles, labels = ax2.get_legend_handles_labels()
            by_label = dict(zip(labels, handles))
            sorted_handles = []
            sorted_labels = []
            for t in possible_types:
                if t in by_label and t in active_types:
                    sorted_handles.append(by_label[t])
                    sorted_labels.append(t)

            ax2.legend(sorted_handles, sorted_labels, loc='upper center', bbox_to_anchor=(0.5, -0.2), 
                       ncol=5, frameon=False, prop={'family': 'KaiTi'})
        else:
             ax2.text(0.5, 0.5, "未選擇數據類型", ha='center', va='center', fontdict=FONT_TITLE)

        ax2.set_title('異文位置分布概覽', fontdict=FONT_TITLE, pad=15)
        ax2.spines['top'].set_visible(False)
        ax2.spines['right'].set_visible(False)
        ax2.spines['left'].set_visible(False)
        
        # 【修改點 2】大幅增加底部留白(bottom=0.3)給圖例，增加左側留白(left=0.15)給文件名
        self.fig_dist.subplots_adjust(left=0.15, right=0.95, top=0.9, bottom=0.30)
        self.canvas_dist.draw()

        # ==========================
        # 3. 雷達圖 (Radar Charts)
        # ==========================
        # A. 4軸雷達圖
        self.draw_radar(self.fig_radar_4, names, ['訛', '脫', '衍', '倒'], "")
        self.canvas_radar_4.draw()

        # B. 5軸雷達圖
        if self.has_variant_split:
            self.canvas_radar_5.setVisible(True)
            self.draw_radar(self.fig_radar_5, names, ['訛', '脫', '衍', '倒', '異'], "")
            self.canvas_radar_5.draw()
        else:
            self.canvas_radar_5.setVisible(False)

    def draw_radar(self, fig, names, categories, title):
        """繪製雷達圖的通用函數 """
        fig.clear()
        ax = fig.add_subplot(111, polar=True)
        
        num_vars = len(categories)
        angles = np.linspace(0, 2 * np.pi, num_vars, endpoint=False).tolist()
        angles += angles[:1] # 閉環
        
        prop_cycle = plt.rcParams['axes.prop_cycle']
        default_colors = prop_cycle.by_key()['color']
        
        for i, name in enumerate(names):
            stats = self.all_data[name]['stats']
            values = [stats.get(cat, 0) for cat in categories]
            values += values[:1]
            
            color = default_colors[i % len(default_colors)]
            ax.plot(angles, values, linewidth=1.5, linestyle='-', label=name, color=color)
            ax.fill(angles, values, color=color, alpha=0.05)
            
        ax.set_theta_offset(np.pi / 2)
        ax.set_theta_direction(-1)
        ax.set_xticks(angles[:-1])
        ax.set_xticklabels(categories, size=11, fontname='KaiTi', color='#333333')
        ax.set_rlabel_position(0)
        
        # 設置標題
        ax.set_title(title, y=1.0, fontname='KaiTi', fontsize=12, color='#5D4037')
        
        ax.grid(color='#E0E0E0', linestyle='--')
        ax.spines['polar'].set_visible(False)
        
        ax.tick_params(pad=10) 
        
        ax.legend(loc='upper center', bbox_to_anchor=(0.5, -0.15), 
                  frameon=False, prop={'family': 'KaiTi', 'size': 9}, ncol=3)
        
        fig.subplots_adjust(top=0.85, bottom=0.25, left=0.10, right=0.90)

# ==========================================
# 6. 異體字字典窗口 
# ==========================================
class DictionaryWindow(QWidget):
    """
    [類別說明]：獨立的異體字查詢窗口。
    [功能]：允許用戶手動查詢系統數據庫中的正異關係，輔助人工校對。
    """
    def __init__(self, engine):
        super().__init__()
        self.engine = engine
        self.setWindowTitle("異體字字典")
        self.resize(500, 600)
        self.setStyleSheet("""
            QWidget { background-color: #FAF9F6; font-family: "SimSun", serif; }
            QLineEdit { 
                padding: 8px; border: 1px solid #B0B0B0; border-radius: 4px; 
                background: #FFF; font-size: 16px; 
            }
            QPushButton { 
                background-color: #B74639; color: #FFF; border-radius: 4px; 
                padding: 8px 15px; font-weight: bold; font-family: "KaiTi"; font-size: 15px;
            }
            QPushButton:hover { background-color: #C85446; }
            QScrollArea { border: none; background: transparent; }
            
            /* 字卡樣式 */
            QFrame#CharCard {
                background-color: #FFFFFF;
                border: 1px solid #E0E0E0;
                border-radius: 8px;
                margin-bottom: 15px;
            }
            QLabel#CardTitle { font-family: "KaiTi"; font-size: 28px; font-weight: bold; color: #B74639; }
            QLabel#CardLabel { font-weight: bold; color: #5D4037; font-size: 14px; }
            QLabel#CardValue { color: #333; font-size: 14px; }
        """)
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(15)

        # 1. 頂部搜索區
        search_layout = QHBoxLayout()
        self.input_search = QLineEdit()
        self.input_search.setPlaceholderText("請輸入單個漢字進行查詢")
        self.input_search.returnPressed.connect(self.do_search)
        
        btn_search = QPushButton("查詢")
        btn_search.clicked.connect(self.do_search)
        
        search_layout.addWidget(self.input_search)
        search_layout.addWidget(btn_search)
        layout.addLayout(search_layout)

        # 2. 結果顯示區 (滾動區域)
        self.scroll = QScrollArea()
        self.scroll.setWidgetResizable(True)
        self.result_container = QWidget()
        self.result_layout = QVBoxLayout(self.result_container)
        self.result_layout.setAlignment(Qt.AlignmentFlag.AlignTop) # 頂部對齊
        self.scroll.setWidget(self.result_container)
        
        layout.addWidget(self.scroll)

    def do_search(self):
        text = self.input_search.text().strip()
        if not text: return
        
        # 清空舊結果
        for i in reversed(range(self.result_layout.count())): 
            w = self.result_layout.itemAt(i).widget()
            if w: w.setParent(None)

        # 逐字查詢並生成卡片
        found_any = False
        db = self.engine.variant_db
        
        for char in text:
            if char in db:
                found_any = True
                self.create_card(char, db[char])
            else:
                self.create_empty_card(char)

    def create_card(self, char, info):
        card = QFrame()
        card.setObjectName("CharCard")
        # 設置卡片陰影效果 
        # shadow = QGraphicsDropShadowEffect(); shadow.setBlurRadius(10); shadow.setColor(QColor(0,0,0,20)); card.setGraphicsEffect(shadow)
        
        c_layout = QVBoxLayout(card)
        c_layout.setContentsMargins(20, 20, 20, 20)
        
        # A. 標題行 (字 + 拼音/編碼如果有的話，這裡只顯示字)
        title_line = QHBoxLayout()
        lbl_char = QLabel(char)
        lbl_char.setObjectName("CardTitle")
        title_line.addWidget(lbl_char)
        title_line.addStretch()
        c_layout.addLayout(title_line)
        
        # 分割線
        line = QFrame(); line.setFrameShape(QFrame.Shape.HLine); line.setStyleSheet("color: #EEE;")
        c_layout.addWidget(line)

        # B. 屬性解析
        standards = info.get('standards', [])
        is_std = info.get('is_standard', False)
        variants = info.get('variants', [])
        
        # 1. 身份屬性
        status_text = "規範正體" if is_std else "異體 / 非標注正體"
        color = "green" if is_std else "#666"
        self.add_row(c_layout, "身份屬性：", status_text, color)

        # 2. 對應關係邏輯
        other_standards = [s for s in standards if s != char]
        
        if is_std and not other_standards:
            self.add_row(c_layout, "狀態說明：", "此字為獨立正體，無其他對應關係。")
        elif is_std and other_standards:
            self.add_row(c_layout, "狀態說明：", "【雙重身份】本身是正體，亦是以下字的異體。", "orange")
            self.add_row(c_layout, "兼作異體：", "、".join(other_standards))
        elif not is_std and standards:
            self.add_row(c_layout, "對應正體：", f"【{'、'.join(standards)}】", "#B74639", True)

        # 3. 關聯異體
        display_variants = [v for v in variants if v != char]
        if display_variants:
            self.add_row(c_layout, "下屬異體：", "  ".join(display_variants))
        else:
            self.add_row(c_layout, "下屬異體：", "(無)")

        self.result_layout.addWidget(card)

    def create_empty_card(self, char):
        card = QFrame(); card.setObjectName("CharCard")
        l = QVBoxLayout(card); l.setContentsMargins(20, 20, 20, 20)
        
        title = QLabel(f"{char} 【未收錄】"); title.setStyleSheet("font-size: 20px; color: #888; font-family: 'KaiTi';")
        l.addWidget(title)
        desc = QLabel("可能是訛字、生僻字或數據庫未涵蓋。")
        desc.setStyleSheet("color: #999; margin-top: 5px;")
        l.addWidget(desc)
        self.result_layout.addWidget(card)

    def add_row(self, layout, label, value, color="#333", bold=False):
        row = QHBoxLayout()
        lbl = QLabel(label); lbl.setObjectName("CardLabel"); lbl.setFixedWidth(80)
        val = QLabel(value); val.setObjectName("CardValue")
        val.setStyleSheet(f"color: {color};" + ("font-weight: bold;" if bold else ""))
        val.setWordWrap(True) # 允許換行
        row.addWidget(lbl, 0, Qt.AlignmentFlag.AlignTop)
        row.addWidget(val, 1)
        layout.addLayout(row)

# ==========================================
# 【新增代碼塊】自定義字典窗口
# ==========================================
class CustomDictWindow(QDialog):
    def __init__(self, engine, parent=None):
        super().__init__(parent)
        self.engine = engine
        self.setWindowTitle("自定義異體字/訛誤規則表")
        self.resize(700, 500)
        self.init_ui()
        self.load_from_engine()

    def init_ui(self):
        layout = QVBoxLayout(self)
        
        info = QLabel("【說明】左側定義【自定義異體】(如 A=B)，右側定義【自定義訛文】(如 A!=B)。\n"
                      "支持導入txt規則文件和導出爲 txt規則文件。\n"
                      "同時，您可以在「異文對照」中右鍵指定條目快速自定義規則。\n"
                      "規則優先級高於系統內置臺灣地區教育部異體字字典（第七版），您可以利用該版塊處理一些因刻工習慣或寫本特徵出現的特殊情況。\n"
                      )
        info.setStyleSheet("color: #555; font-style: italic; margin-bottom: 10px;")
        layout.addWidget(info)
        
        edit_layout = QHBoxLayout()
        
        # 左側：異體字
        grp_var = QGroupBox("自定義異體字 (語法：A=B)")
        l1 = QVBoxLayout(grp_var)
        self.txt_variants = QTextEdit()
        self.txt_variants.setPlaceholderText("每行一條，例如：已=己 厶=私")
        l1.addWidget(self.txt_variants)
        edit_layout.addWidget(grp_var)
        
        # 右側：強制訛誤
        grp_ex = QGroupBox("自定義訛文 (語法：A!=B)")
        l2 = QVBoxLayout(grp_ex)
        self.txt_excludes = QTextEdit()
        self.txt_excludes.setPlaceholderText("每行一條，例如：雲!=云 後!=后")
        l2.addWidget(self.txt_excludes)
        edit_layout.addWidget(grp_ex)
        
        layout.addLayout(edit_layout)
        
        btn_layout = QHBoxLayout()
        btn_load = QPushButton(" 導入規則(txt)")
        btn_load.clicked.connect(self.import_file)
        btn_save = QPushButton(" 導出規則(txt)")
        btn_save.clicked.connect(self.export_file)
        
        btn_apply = QPushButton("應用並生效")
        btn_apply.clicked.connect(self.apply_rules)
        btn_apply.setStyleSheet("background-color: #B74639; color: white; font-weight: bold;")
        
        btn_layout.addWidget(btn_load)
        btn_layout.addWidget(btn_save)
        btn_layout.addStretch()
        btn_layout.addWidget(btn_apply)
        layout.addLayout(btn_layout)

    def load_from_engine(self):
        v_text = []
        seen = set()
        for k, v in self.engine.custom_variants.items():
            pair = tuple(sorted((k, v)))
            if pair not in seen:
                v_text.append(f"{k}={v}")
                seen.add(pair)
        self.txt_variants.setPlainText("\n".join(v_text))
        
        e_text = []
        for a, b in self.engine.custom_excludes:
            e_text.append(f"{a}!={b}")
        self.txt_excludes.setPlainText("\n".join(e_text))

    def apply_rules(self):
        var_map = {}
        excludes = set()
        
        # 解析異體
        for line in self.txt_variants.toPlainText().split('\n'):
            line = line.strip()
            if '=' in line and '!=' not in line:
                parts = line.split('=')
                if len(parts) == 2:
                    a, b = parts[0].strip(), parts[1].strip()
                    if a and b: var_map[a] = b; var_map[b] = a
        
        # 解析屏蔽
        for line in self.txt_excludes.toPlainText().split('\n'):
            line = line.strip()
            if '!=' in line:
                parts = line.split('!=')
                if len(parts) == 2:
                    a, b = parts[0].strip(), parts[1].strip()
                    if a and b: excludes.add((a, b)); excludes.add((b, a))
        
        self.engine.update_custom_dict(var_map, excludes)
        QMessageBox.information(self, "成功", "規則已更新至內存！\n請重新運行校勘以查看最新效果。")
        self.accept()

    def import_file(self):
        path, _ = QFileDialog.getOpenFileName(self, "導入規則", "", "Text Files (*.txt)")
        if not path: return
        try:
            with open(path, 'r', encoding='utf-8') as f:
                content = f.read()
            v_c, e_c, mode = "", "", None
            for line in content.split('\n'):
                l = line.strip()
                if l == "[VARIANTS]": mode = 'V'; continue
                elif l == "[EXCLUDES]": mode = 'E'; continue
                if not l: continue
                if mode == 'V': v_c += l + "\n"
                elif mode == 'E': e_c += l + "\n"
            
            if mode is None and '!=' not in content: v_c = content # 兼容舊格式
            
            if v_c: self.txt_variants.append("\n" + v_c.strip())
            if e_c: self.txt_excludes.append("\n" + e_c.strip())
        except Exception as e: QMessageBox.warning(self, "錯誤", str(e))

    def export_file(self):
        path, _ = QFileDialog.getSaveFileName(self, "導出規則", "custom_rules.txt", "Text Files (*.txt)")
        if not path: return
        try:
            with open(path, 'w', encoding='utf-8') as f:
                f.write("[VARIANTS]\n" + self.txt_variants.toPlainText() + "\n\n[EXCLUDES]\n" + self.txt_excludes.toPlainText())
            QMessageBox.information(self, "成功", "規則已保存！")
        except Exception as e: QMessageBox.warning(self, "錯誤", str(e))

# ==========================================
# [新增組件] MDict 字典閱讀器 (完整依賴修復版)
# ==========================================
from PyQt6.QtWebEngineCore import QWebEnginePage # 確保導入這個

class MDictSchemeHandler(QWebEngineUrlSchemeHandler):
    """[1. 圖片攔截器] 從內存讀取圖片"""
    def __init__(self, resource_index, mdd_data):
        super().__init__()
        self.resource_index = resource_index
        self.mdd_data = mdd_data

    def requestStarted(self, job: QWebEngineUrlRequestJob):
        path = job.requestUrl().path()
        # 清洗文件名
        clean_name = path.replace("\\", "/").split("/")[-1].lower()
        
        data = None
        if self.resource_index and clean_name in self.resource_index:
            real_key = self.resource_index[clean_name]
            data = self.mdd_data.get(real_key)

        if data:
            buf = QBuffer(self)
            buf.setData(data)
            buf.open(QIODevice.OpenModeFlag.ReadOnly)
            job.reply(b"image/jpeg", buf)
        else:
            job.fail(QWebEngineUrlRequestJob.Error.UrlNotFound)

class MDictPage(QWebEnginePage):
    """[2. 鏈接攔截器] 處理 entry:// 點擊和跳轉"""
    def __init__(self, parent_window):
        super().__init__(parent_window)
        self.parent_window = parent_window

    def acceptNavigationRequest(self, url, nav_type, is_main_frame):
        # 攔截點擊事件
        if nav_type == QWebEnginePage.NavigationType.NavigationTypeLinkClicked:
            schema = url.scheme()
            if schema in ["entry", "mdict"]:
                target = url.path().strip("/")
                # 讓主窗口跳轉
                self.parent_window.show_entry(target)
                return False
        return True

class MDictBackend:
    """[3. 數據後端] 內存全量加載"""
    def __init__(self, mdx_path, mdd_path):
        self.ready = False
        self.content_db = {} 
        self.display_keys = []
        self.mdd_map = {}     
        self.resource_index = {} 
        
        print(f"[-] 正在加載 MDX: {mdx_path}")
        if os.path.exists(mdx_path):
            try:
                mdx_obj = MDX(mdx_path)
                # 全量讀取並解碼
                for key_bytes, val_bytes in mdx_obj.items():
                    k_str = ""
                    try: k_str = key_bytes.decode('utf-8').strip()
                    except:
                        try: k_str = key_bytes.decode('gbk').strip()
                        except: k_str = str(key_bytes)
                    
                    v_str = ""
                    try: v_str = val_bytes.decode('utf-8')
                    except:
                        try: v_str = val_bytes.decode('gbk')
                        except: v_str = str(val_bytes)

                    if k_str:
                        self.content_db[k_str] = v_str
                        self.display_keys.append(k_str)
                
                self.display_keys.sort()
                print(f"[+] MDX 加載完畢，索引: {len(self.display_keys)}")
                self.ready = True
            except Exception as e:
                print(f"[!] MDX 錯誤: {e}")

        print(f"[-] 正在加載 MDD: {mdd_path}")
        if os.path.exists(mdd_path):
            try:
                mdd_obj = MDD(mdd_path)
                self.mdd_map = dict(mdd_obj.items())
                for key in self.mdd_map.keys():
                    try:
                        k_str = key.decode('utf-8', errors='ignore') if isinstance(key, bytes) else str(key)
                        clean = k_str.replace("\\", "/").split("/")[-1].lower()
                        self.resource_index[clean] = key
                    except: pass
                print(f"[+] MDD 資源就緒: {len(self.resource_index)}")
            except Exception as e:
                print(f"[!] MDD 錯誤: {e}")

    def get_content(self, key):
        # 處理 @@@LINK= 跳轉
        content = self.content_db.get(key, None)
        if content and content.startswith("@@@LINK="):
            target = content.split("=")[1].strip()
            return self.content_db.get(target, None)
        return content

class MDictWindow(QWidget):
    """[4. 瀏覽窗口] UI優化版"""
    def __init__(self, resources_path):
        super().__init__()
        self.setWindowTitle("歷代避諱字匯典 (閲讀優化版)")
        self.resize(1200, 900)
        
        mdx = os.path.join(resources_path, "歷代避諱字匯典.mdx")
        mdd = os.path.join(resources_path, "歷代避諱字匯典.mdd")
        
        self.backend = MDictBackend(mdx, mdd)
        self.init_ui()
        
        # 設置攔截器和自定義頁面
        self.custom_page = MDictPage(self)
        self.web.setPage(self.custom_page)

        if self.backend.mdd_map:
            self.handler = MDictSchemeHandler(self.backend.resource_index, self.backend.mdd_map)
            try:
                self.custom_page.profile().installUrlSchemeHandler(b"mdict", self.handler)
            except: pass 

        if self.backend.ready:
            self.populate_list()

    def init_ui(self):
        outer_layout = QVBoxLayout(self)
        outer_layout.setContentsMargins(0, 0, 0, 0)
        outer_layout.setSpacing(0)

        # 中間部分
        splitter = QSplitter(Qt.Orientation.Horizontal)
        
        left_panel = QWidget()
        left_layout = QVBoxLayout(left_panel)
        left_layout.setContentsMargins(5, 5, 5, 5)
        
        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText("🔍 搜索...")
        self.search_input.setFixedHeight(30)
        self.search_input.textChanged.connect(self.filter_list)
        left_layout.addWidget(self.search_input)
        
        self.list_widget = QListWidget()
        self.list_widget.setStyleSheet("font-size: 14px; border: 1px solid #ccc;")
        self.list_widget.itemClicked.connect(self.on_item_clicked)
        left_layout.addWidget(self.list_widget)
        
        self.web = QWebEngineView()
        self.web.setStyleSheet("background-color: #555;") 
        
        splitter.addWidget(left_panel)
        splitter.addWidget(self.web)
        splitter.setStretchFactor(0, 2)
        splitter.setStretchFactor(1, 8)
        
        outer_layout.addWidget(splitter)

        # 底部控制欄
        bottom_bar = QWidget()
        bottom_bar.setFixedHeight(60)
        bottom_bar.setStyleSheet("background-color: #EAECEE; border-top: 1px solid #ccc;")
        bottom_layout = QHBoxLayout(bottom_bar)
        
        self.btn_prev = QPushButton("◀ 上一頁")
        self.btn_prev.setFixedSize(120, 40)
        self.btn_prev.setStyleSheet("QPushButton { background-color: #fff; border: 1px solid #bbb; border-radius: 5px; font-size: 16px; font-weight: bold; color: #333; } QPushButton:hover { background-color: #f0f0f0; }")
        self.btn_prev.clicked.connect(self.go_prev)
        
        self.lbl_status = QLabel("準備就緒")
        self.lbl_status.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.lbl_status.setStyleSheet("font-size: 16px; color: #5D4037; font-weight: bold;")
        
        self.btn_next = QPushButton("下一頁 ▶")
        self.btn_next.setFixedSize(120, 40)
        self.btn_next.setStyleSheet(self.btn_prev.styleSheet())
        self.btn_next.clicked.connect(self.go_next)
        
        bottom_layout.addWidget(self.btn_prev)
        bottom_layout.addStretch()
        bottom_layout.addWidget(self.lbl_status)
        bottom_layout.addStretch()
        bottom_layout.addWidget(self.btn_next)
        
        outer_layout.addWidget(bottom_bar)

    def populate_list(self):
        self.list_widget.clear()
        self.list_widget.addItems(self.backend.display_keys)
        if self.list_widget.count() > 0:
            self.list_widget.setCurrentRow(0)
            self.on_item_clicked(self.list_widget.item(0))

    def filter_list(self, text):
        text = text.strip()
        for i in range(self.list_widget.count()):
            item = self.list_widget.item(i)
            item.setHidden(text not in item.text())

    def on_item_clicked(self, item):
        key = item.text()
        self.update_status_label(key)
        self.show_entry(key)

    def show_entry(self, key):
        """顯示內容並優化圖片"""
        # 如果是內部跳轉，可能需要更新選中項
        items = self.list_widget.findItems(key, Qt.MatchFlag.MatchExactly)
        if items and items[0] != self.list_widget.currentItem():
            self.list_widget.setCurrentItem(items[0])
            self.update_status_label(key)

        html = self.backend.get_content(key)
        if html:
            # CSS 強制圖片自適應 + 隱藏原生翻頁文字
            full_html = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <style>
                    html, body {{
                        margin: 0; padding: 0; height: 100vh; width: 100%;
                        background-color: #FDFEFE;
                        display: flex; flex-direction: column;
                        align-items: center; justify-content: center;
                        overflow: hidden; 
                    }}
                    img {{ 
                        max-height: 98vh; 
                        max-width: 98vw; 
                        object-fit: contain; 
                        display: block; 
                        box-shadow: 0 5px 15px rgba(0,0,0,0.2);
                    }}
                    center, a {{ display: none !important; }}
                </style>
            </head>
            <body>
                {html}
            </body>
            </html>
            """
            self.web.setHtml(full_html, baseUrl=QUrl("mdict://"))
        else:
            self.web.setHtml(f"<h3 style='color:red;'>無法讀取數據: {key}</h3>")

    def go_prev(self):
        row = self.list_widget.currentRow()
        if row > 0:
            self.list_widget.setCurrentRow(row - 1)
            self.on_item_clicked(self.list_widget.currentItem())

    def go_next(self):
        row = self.list_widget.currentRow()
        if row < self.list_widget.count() - 1:
            self.list_widget.setCurrentRow(row + 1)
            self.on_item_clicked(self.list_widget.currentItem())

    def update_status_label(self, key):
        row = self.list_widget.currentRow() + 1
        total = self.list_widget.count()
        self.lbl_status.setText(f"第 {row} / {total} 條 ： {key}")

# ==========================================
# 避諱檢測面板 (UI 優化版 - 全可拖動)
# ==========================================
from PyQt6.QtWidgets import (QTreeWidget, QTreeWidgetItem, QTextBrowser, 
                             QHeaderView, QSplitter)

class TabooCheckerPanel(QWidget):
    def __init__(self, taboo_db_path):
        super().__init__()
        self.taboo_db = []
        self.full_data_map = {} 
        self.current_text = ""  
        
        # 加載數據
        self.load_db(taboo_db_path)
        self.init_ui()
        self.mdict_window = None # [新增] 初始化字典窗口變量

    def load_db(self, path):
        if not os.path.exists(path):
            return
        try:
            with open(path, 'r', encoding='utf-8') as f:
                self.taboo_db = json.load(f)
        except Exception as e:
            print(f"避諱數據加載失敗: {e}")

    def init_ui(self):
        layout = QVBoxLayout(self)
        
        # 1. 頂部工具欄 (版本切換)
        top_bar = QHBoxLayout()
        top_bar.addWidget(QLabel("當前檢測版本："))
        self.combo_versions = QComboBox()
        self.combo_versions.currentIndexChanged.connect(self.on_version_changed)
        top_bar.addWidget(self.combo_versions)
        btn_dict = QPushButton("檢測避諱字一覽表 ")
        btn_dict.setStyleSheet("""
            QPushButton {
                background-color: #EAECEE; border: 1px solid #BDC3C7; 
                border-radius: 4px; padding: 4px 10px; font-family: 'KaiTi'; font-weight: bold; color: #5D4037;
            }
            QPushButton:hover { background-color: #D5D8DC; }
        """)
        btn_dict.clicked.connect(self.open_taboo_dictionary)
        top_bar.addWidget(btn_dict)

        # [新增] 查閱字典按鈕
        btn_mdict = QPushButton("📖《歷代避諱字匯典》")
        btn_mdict.setStyleSheet("""
            QPushButton {
                background-color: #D4E6F1; border: 1px solid #A9CCE3; 
                border-radius: 4px; padding: 4px 10px; font-family: 'KaiTi'; font-weight: bold; color: #154360;
            }
            QPushButton:hover { background-color: #AED6F1; }
        """)
        btn_mdict.clicked.connect(self.open_mdict_window)
        top_bar.addWidget(btn_mdict)
        
        top_bar.addStretch()
        layout.addLayout(top_bar)



        # 2. 主體內容 (主分割器：左右分割)
        # 用於分隔「左側列表區」和「右側預覽區」
        main_splitter = QSplitter(Qt.Orientation.Horizontal)
        
        # === 左側區域 (改為垂直分割器：樹狀圖 + 詳細備注) ===
        # 【修改點】這裡原本是 QWidget+Layout，現在改為 QSplitter，實現上下拖動
        left_splitter = QSplitter(Qt.Orientation.Vertical)

        # A. 檢測結果樹
        self.tree = QTreeWidget()
        self.tree.setHeaderLabels(["朝代 / 避諱對象", "       避諱字說明"])
        self.tree.header().setSectionResizeMode(0, QHeaderView.ResizeMode.ResizeToContents)
        self.tree.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self.tree.itemClicked.connect(self.on_tree_item_clicked)
        
        # 將樹加入左側垂直分割器
        left_splitter.addWidget(self.tree)

        # B. 詳細備注區
        detail_group = QGroupBox("詳細信息 ")
        detail_layout = QVBoxLayout(detail_group)
        detail_layout.setContentsMargins(5, 15, 5, 5) # 微調邊距適應分割器
        
        self.detail_text = QTextBrowser() 
        self.detail_text.setStyleSheet("background-color: #F9F9F9; color: #333; font-size: 14px; border: none;")
        self.detail_text.setPlaceholderText("點擊上方條目，此處顯示詳細備注...")
        detail_layout.addWidget(self.detail_text)
        
        # 將備注區加入左側垂直分割器
        left_splitter.addWidget(detail_group)

        # 設置左側垂直分割器的初始比例 (樹佔 2/3，備注佔 1/3)
        left_splitter.setStretchFactor(0, 2)
        left_splitter.setStretchFactor(1, 1)

        # 將左側垂直分割器加入主水平分割器
        main_splitter.addWidget(left_splitter)

        # === 右側區域：上下文預覽 ===
        right_panel = QGroupBox("上下文預覽 (點擊左側命中字查看)")
        right_layout = QVBoxLayout(right_panel)
        
        self.preview_browser = QTextBrowser()
        self.preview_browser.setStyleSheet("""
            QTextBrowser { 
                font-size: 16px; 
                line-height: 1.8; 
                background-color: #FFF;
                border: 1px solid #DDD;
            }
        """)
        self.preview_browser.setPlaceholderText("暫無預覽內容")
        right_layout.addWidget(self.preview_browser)
        
        # 將右側區域加入主水平分割器
        main_splitter.addWidget(right_panel)
        
        # 設置主分割器左右寬度比例 (4:6)
        main_splitter.setStretchFactor(0, 4)
        main_splitter.setStretchFactor(1, 6)
        
        layout.addWidget(main_splitter)

    def load_text_data(self, base_text, wit_texts_dict):
        """由 MainWindow 調用"""
        self.full_data_map = {}
        self.full_data_map["底本"] = base_text
        for name, text in wit_texts_dict.items():
            self.full_data_map[name] = text
            
        self.combo_versions.blockSignals(True)
        self.combo_versions.clear()
        self.combo_versions.addItems(self.full_data_map.keys())
        self.combo_versions.blockSignals(False)
        
        self.combo_versions.setCurrentIndex(0)
        self.run_detection("底本")

    def on_version_changed(self):
        ver_name = self.combo_versions.currentText()
        if ver_name:
            self.run_detection(ver_name)

    def run_detection(self, ver_name):
        """執行檢測邏輯"""
        self.tree.clear()
        self.preview_browser.clear()
        self.detail_text.clear()
        
        text = self.full_data_map.get(ver_name, "")
        self.current_text = text
        
        if not text: return
        if not self.taboo_db:
            QTreeWidgetItem(self.tree, ["無數據", "請檢查 taboo_data.json"])
            return

        dynasty_map = defaultdict(lambda: {'hits': [], 'misses': []})
        
        for entry in self.taboo_db:
            if not entry['chars']: continue 
            if not entry.get('name') or not str(entry['name']).strip(): continue

            found_chars = [c for c in entry['chars'] if c in text]
            
            if found_chars:
                entry_copy = entry.copy()
                entry_copy['found'] = found_chars
                dynasty_map[entry['dynasty']]['hits'].append(entry_copy)
            else:
                dynasty_map[entry['dynasty']]['misses'].append(entry)

        seen_dynasties = []
        for entry in self.taboo_db:
            if entry['dynasty'] not in seen_dynasties:
                seen_dynasties.append(entry['dynasty'])
                
        for dyn in seen_dynasties:
            data = dynasty_map.get(dyn)
            if not data or (not data['hits'] and not data['misses']): 
                continue
            
            hits = data['hits']
            misses = data['misses']
            
            root = QTreeWidgetItem(self.tree, [f"{dyn} (發現:{len(hits)} / 未發現:{len(misses)})", ""])
            root.setExpanded(True)
            root.setForeground(0, QColor("#5D4037"))
            root.setFont(0, QFont("KaiTi", 12, QFont.Weight.Bold))

            if hits:
                hit_node = QTreeWidgetItem(root, [f"⚠️ 疑似未避諱 ({len(hits)})", ""])
                hit_node.setForeground(0, QColor("#B74639"))
                hit_node.setExpanded(True)
                
                for h in hits:
                    found_str = "、".join(h['found'])
                    msg = f"文中有【{found_str}】字未避諱"
                    item = QTreeWidgetItem(hit_node, [f"{h['temple']} - {h['name']}", msg])
                    item.setData(0, Qt.ItemDataRole.UserRole, h)
                    item.setToolTip(0, "點擊查看詳情與上下文")

            if misses:
                miss_node = QTreeWidgetItem(root, [f"✅ 符合避諱 ({len(misses)})", ""])
                miss_node.setForeground(0, QColor("#5C7A62"))
                
                for m in misses:
                    msg = f"全書無【{m['name']}】避諱字"
                    item = QTreeWidgetItem(miss_node, [f"{m['temple']} - {m['name']}", msg])
                    item.setData(0, Qt.ItemDataRole.UserRole, m)

    def on_tree_item_clicked(self, item, column):
        """點擊樹節點，更新下方的備注和右側的預覽"""
        data = item.data(0, Qt.ItemDataRole.UserRole)
        
        if not data: 
            self.detail_text.clear()
            self.preview_browser.clear()
            return
        
        # 更新下方詳細備注區
        note_content = (
            f"<b>朝代：</b>{data['dynasty']}<br>"
            f"<b>廟號：</b>{data['temple']}<br>"
            f"<b>避諱對象：</b>{data['name']}<br>"
            f"<b>避諱字概覽：</b>{'、'.join(data['chars'])}<br>"
            f"<hr>"
            f"<b>備注說明：</b><br>{data['note']}"
        )
        self.detail_text.setHtml(note_content)

        # 更新右側上下文預覽
        self.preview_browser.clear()
        
        if 'found' not in data:
            self.preview_browser.setHtml("<p style='color:#666; text-align:center;'>（此條目未在文中發現避諱字，無上下文預覽）</p>")
            return

        found_chars = data['found']
        full_text = self.current_text
        total_len = len(full_text)
        
        html_output = []
        
        for char in found_chars:
            import re
            matches = [m.start() for m in re.finditer(re.escape(char), full_text)]
            
            for idx in matches:
                start = max(0, idx - 10)
                end = min(total_len, idx + 11)
                
                pre = full_text[start:idx]
                target = full_text[idx]
                post = full_text[idx+1:end]
                
                line_html = (
                    f"<div style='margin-bottom: 10px; border-bottom: 1px solid #EEE; padding-bottom: 5px;'>"
                    f"<span style='color:#888'>...{pre}</span>"
                    f"<span style='color:#B74639; font-weight:bold; font-size:18px; background-color:#FADBD8;'> {target} </span>"
                    f"<span style='color:#888'>{post}...</span>"
                    f"</div>"
                )
                html_output.append(line_html)
        
        if html_output:
            self.preview_browser.setHtml("".join(html_output))
        else:
            self.preview_browser.setHtml("無預覽")
            
    def open_taboo_dictionary(self):
        """打開字典窗口"""
        # 檢查是否已經打開，防止重複創建
        if not hasattr(self, 'dict_window') or self.dict_window is None:
            self.dict_window = TabooDictionaryWindow(self.taboo_db, self)
            # 窗口關閉時清理引用
            self.dict_window.finished.connect(lambda: setattr(self, 'dict_window', None))
        
        self.dict_window.show()
        self.dict_window.raise_() # 確保在最上層
        self.dict_window.activateWindow()

    def open_mdict_window(self):
        """[新增] 打開 MDict 字典窗口"""
        if not self.mdict_window:
            # 使用您的 get_resource_path 確保打包後路徑正確
            res_path = get_resource_path("resources")
            self.mdict_window = MDictWindow(res_path)
        
        self.mdict_window.show()
        self.mdict_window.raise_()
        self.mdict_window.activateWindow()


# ==========================================
# 避諱字字典窗口 (樣式衝突修復版：自適應高度 + 完整顯示)
# ==========================================
class TabooDictionaryWindow(QDialog):
    def __init__(self, taboo_db, parent=None):
        super().__init__(parent)
        self.setWindowTitle("檢測避諱字一覽表")
        self.resize(1100, 800) 
        self.taboo_db = taboo_db
        self.setModal(False) 
        
        self.init_ui()
        self.load_data()

    def init_ui(self):
        layout = QVBoxLayout(self)
        
        # 1. 頂部搜索區
        search_layout = QHBoxLayout()
        search_lbl = QLabel("🔍 檢索：")
        search_lbl.setStyleSheet("font-weight: bold; color: #5D4037;")
        
        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText("輸入關鍵字（朝代、廟號、人名、避諱字）進行實時過濾...")
        self.search_input.textChanged.connect(self.filter_data)
        
        search_layout.addWidget(search_lbl)
        search_layout.addWidget(self.search_input)
        layout.addLayout(search_layout)

        # 2. 數據展示樹
        self.tree = QTreeWidget()
        self.tree.setHeaderLabels(["廟號 / 避諱對象", "避諱字概覽", "備注"])
        
# === 【核心優化：應用自適應代理】 ===
        # 應用我們剛寫的代理，設置上下留白為 24px (更舒適)，左右 10px
        delegate = AutoAdaptDelegate(self.tree, vertical_padding=24, horizontal_padding=10)
        self.tree.setItemDelegate(delegate)

        # === 基礎屬性設置 ===
        self.tree.setWordWrap(True)            # 必須開啟，配合代理計算
        self.tree.setUniformRowHeights(False)  # 必須關閉，允許行高不一致
        self.tree.setTextElideMode(Qt.TextElideMode.ElideNone) # 關閉省略號
        self.tree.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff) # 關閉橫向滾動
        
        # === 列寬比例設置 ===
        header = self.tree.header()
        
        # 第0列 (名稱): 交互式調整，給一個固定初始寬度，避免太寬擠壓備注
        header.setSectionResizeMode(0, QHeaderView.ResizeMode.Interactive)
        self.tree.setColumnWidth(0, 250) 
        
        # 第1列 (避諱字): 自動拉伸
        header.setSectionResizeMode(1, QHeaderView.ResizeMode.Stretch)
        
        # 第2列 (備注): 自動拉伸 (因為有了代理，這裡會自動撐高)
        header.setSectionResizeMode(2, QHeaderView.ResizeMode.Stretch)
        
        # === 樣式微調 ===
        # 移除之前的 padding 設置，因為現在由 Delegate 控制佈局
        # 增加 alternate-background-color 讓長表格閱讀更舒適
        self.tree.setAlternatingRowColors(True)
    # === 【視覺優化：莫蘭迪撞色系 QSS】 ===
        self.tree.setStyleSheet("""
            QTreeWidget { 
                font-size: 15px; 
                line-height: 1.6;
                outline: 0;  /* 去除虛線框 */
                
                /* 邊框：莫蘭迪灰 */
                border: 1px solid #BDC3C7;
                
                /* 全局背景：極淡的暖白（米白） */
                background-color: #FDFEFE;
                
                /* 斑馬線底色：極淡的冷灰，營造紙張層次感 */
                alternate-background-color: #F4F6F6; 
            }

            /* 2. 鼠標懸停效果：莫蘭迪·霧霾藍 (明顯但柔和) */
            QTreeWidget::item:hover {
                /* 背景：霧霾藍 (明顯的區分色) */
                background-color: #D4E6F1; 
                /* 文字：深岩藍 (加深對比，確保清晰) */
                color: #154360;            
                border: none;
            }

            /* 無論是否激活窗口，選中時都讓背景透出底色，文字保持原樣 */
            QTreeWidget::item:selected,
            QTreeWidget::item:selected:active,
            QTreeWidget::item:selected:!active {
                background-color: transparent; /* 關鍵：背景透明 */
                color: #333333;                /* 關鍵：文字顏色恢復默認 */
                border: none;
            }

            /* 3. 細節修正：選中後如果鼠標還停留在上面，仍然顯示懸停色 */
            /* 如果不加這一條，點擊後懸停效果會消失，看起來會很怪 */
            QTreeWidget::item:selected:hover {
                background-color: #D4E6F1;
                color: #154360;
            }
        """)
        
        layout.addWidget(self.tree)

    def load_data(self):
        self.tree.clear()
        if not self.taboo_db: return

        dynasty_groups = defaultdict(lambda: {"overview": [], "entries": []})
        seen_dynasties = []
        
        for entry in self.taboo_db:
            dyn = entry['dynasty']
            if dyn not in seen_dynasties:
                seen_dynasties.append(dyn)
            
            has_chars = bool(entry['chars'])
            has_name = bool(entry.get('name') and str(entry['name']).strip())
            
            if not has_chars and not has_name:
                dynasty_groups[dyn]['overview'].append(entry)
            else:
                dynasty_groups[dyn]['entries'].append(entry)

        for dyn in seen_dynasties:
            group = dynasty_groups[dyn]
            
            # 朝代根節點
            root = QTreeWidgetItem(self.tree)
            root.setText(0, dyn)
            root.setFont(0, QFont("KaiTi", 16, QFont.Weight.Bold))
            root.setForeground(0, QColor("#5D4037"))
            root.setExpanded(True)
            
            # --- 處理總體概說 ---
            overviews = group['overview']
            overview_text = ""
            if overviews:
                notes = [ov['note'] for ov in overviews if ov['note'] and str(ov['note']).strip()]
                if notes:
                    overview_text = "\n".join(notes)
            
            ov_node = QTreeWidgetItem(root)
            
            # 構造概說內容
            if overview_text:
                full_text = f"【{dyn}代避諱概說】\n{overview_text}"
                # 恢復朱砂紅字體
                ov_node.setForeground(0, QColor("#BC8E87")) 
            else:
                full_text = f"【{dyn}代避諱概說】\n（暫無總體情況說明）"
                ov_node.setForeground(0, QColor("#888"))
            
            ov_node.setText(0, full_text)
            ov_node.setFirstColumnSpanned(True)
            
            # 恢復淡灰底紋 (現在移除 QSS item 樣式後，這裡就會生效了)
            for c in range(3):
                ov_node.setBackground(c, QColor("#F2F4F4"))
            
            # 設置字體
            ov_node.setFont(0, QFont("SimSun", 11, QFont.Weight.Bold))
            # 增加一點行高 (通過在文本中加換行符無法完全控制，依賴內容自適應)

            # --- 處理具體條目 ---
            for entry in group['entries']:
                if not entry['chars'] and not entry.get('name'): continue
                
                # 恢復連接符
                name_str = f"{entry['temple']} - {entry['name']}" 
                if not entry['temple']: 
                    name_str = entry['name']
                
                chars_str = "、".join(entry['chars'])
                note_str = entry['note']
                
                item = QTreeWidgetItem(root)
                item.setText(0, name_str)
                item.setText(1, chars_str)
                item.setText(2, note_str)
                
                # 頂部對齊 (確保長內容時整齊)
                for c in range(3):
                    item.setTextAlignment(c, Qt.AlignmentFlag.AlignTop | Qt.AlignmentFlag.AlignLeft)
                
                # 字體設置
                item.setFont(0, QFont("SimSun", 12)) 
                item.setFont(1, QFont("SimSun", 12, QFont.Weight.Bold))
                
                # 避諱字綠色
                item.setForeground(1, QColor("#9CAF88")) 

    def filter_data(self, text):
        """實時過濾邏輯"""
        search_text = text.strip().lower()
        
        for i in range(self.tree.topLevelItemCount()):
            dynasty_item = self.tree.topLevelItem(i)
            dynasty_visible = False
            
            if search_text in dynasty_item.text(0).lower():
                dynasty_visible = True
                dynasty_item.setHidden(False)
                for j in range(dynasty_item.childCount()):
                    dynasty_item.child(j).setHidden(False)
                continue

            for j in range(dynasty_item.childCount()):
                child = dynasty_item.child(j)
                child_content = (child.text(0) + child.text(1) + child.text(2)).lower()
                
                if search_text in child_content:
                    child.setHidden(False)
                    dynasty_visible = True
                else:
                    child.setHidden(True)
            
            dynasty_item.setHidden(not dynasty_visible)

# ==========================================
# [GUI 組件] 形近誤字分析面板
# ==========================================
class VisualAnalysisPanel(QWidget):
    def __init__(self, engine):
        super().__init__()
        self.engine = engine
        self.all_data = {}
        self.base_text = ""
        self.current_wit = ""
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout(self)
        
        # --- 頂部控制欄 ---
        top_group = QGroupBox("分析參數設置")
        top_layout = QHBoxLayout(top_group)
        
        top_layout.addWidget(QLabel("權重配置："))
        self.slider_ids = QSlider(Qt.Orientation.Horizontal)
        self.slider_ids.setRange(0, 100); self.slider_ids.setValue(30)
        self.slider_ids.setFixedWidth(100)
        top_layout.addWidget(QLabel("部件"))
        top_layout.addWidget(self.slider_ids)
        
        self.slider_pixel = QSlider(Qt.Orientation.Horizontal)
        self.slider_pixel.setRange(0, 100); self.slider_pixel.setValue(70)
        self.slider_pixel.setFixedWidth(100)
        top_layout.addWidget(QLabel("像素"))
        top_layout.addWidget(self.slider_pixel)
        
        top_layout.addSpacing(20)
        
        top_layout.addWidget(QLabel("閾值過濾："))
        self.spin_threshold = QComboBox()
        self.spin_threshold.addItems([f"{i/10:.1f}" for i in range(1, 10)])
        self.spin_threshold.setCurrentText("0.5")
        top_layout.addWidget(self.spin_threshold)
        
        top_layout.addStretch()
        
        self.btn_rule = QPushButton("🛠 設置易混淆部件表")
        self.btn_rule.clicked.connect(self.open_rule_editor)
        top_layout.addWidget(self.btn_rule)
        
        self.btn_run = QPushButton(" 開始形近分析 ")
        self.btn_run.setStyleSheet("background-color: #2E8B57; color: white; font-weight: bold;")
        self.btn_run.clicked.connect(self.run_analysis)
        top_layout.addWidget(self.btn_run)
        
        layout.addWidget(top_group)
        
        # --- 主表格 ---
        self.table = QTableWidget()
        self.table.setColumnCount(6)
        self.table.setHorizontalHeaderLabels(["綜合相似度", "底本字", "校本字", "部件相似度", "像素重合度", "判定說明"])
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        self.table.setAlternatingRowColors(True)
        self.table.setSortingEnabled(True)
        layout.addWidget(self.table)
        
        # --- 狀態欄 ---
        self.lbl_status = QLabel("請點擊「開始形近分析」以計算當前「訛」類條目的相似度。")
        self.lbl_status.setStyleSheet("color: #666; font-style: italic;")
        layout.addWidget(self.lbl_status)

    def load_data(self, wit_name, full_data, base_text, wit_text):
        """從主窗口接收數據"""
        self.current_wit = wit_name
        self.all_data = full_data
        self.base_text = base_text
        # 清空表格，等待用戶點擊運行
        self.table.setRowCount(0)
        self.lbl_status.setText(f"已就緒。分析對象：{wit_name} 中的所有「訛」文。")

    def run_analysis(self):
        if not self.current_wit or self.current_wit not in self.all_data:
            return
            
        results = self.all_data[self.current_wit]['results']
        # 只過濾 "訛"
        errors = [r for r in results if r['type'] == '訛' and r['base_clean_len'] == 1 and r['wit_clean_len'] == 1]
        
        if not errors:
            QMessageBox.information(self, "提示", "當前校本沒有長度為 1:1 的訛文可供分析。")
            return

        # 獲取權重
        w_ids = self.slider_ids.value() / 100.0
        w_pix = self.slider_pixel.value() / 100.0
        threshold = float(self.spin_threshold.currentText())
        
        self.table.setRowCount(0)
        self.lbl_status.setText("正在進行像素級渲染與計算，請稍候...")
        QApplication.processEvents() # 刷新界面防止卡死
        
        count = 0
        for r in errors:
            b_char = r['base_clean']
            w_char = r['wit_clean']
            
            # 核心計算
            score, desc = self.engine.check_similarity(b_char, w_char, weights=(w_ids, w_pix), threshold=threshold)
            
            if score >= threshold:
                row = self.table.rowCount()
                self.table.insertRow(row)
                
                # 分數 (用於排序)
                item_score = QTableWidgetItem(f"{score:.2f}")
                if score > 0.8: 
                    item_score.setBackground(QColor("#D5F5E3")) # 高相似度標綠
                    item_score.setFont(QFont("Arial", 10, QFont.Weight.Bold))
                self.table.setItem(row, 0, item_score)
                
                self.table.setItem(row, 1, QTableWidgetItem(b_char))
                self.table.setItem(row, 2, QTableWidgetItem(w_char))
                
                # 獲取詳細分項 (為了顯示方便，重新調用一次獲取純數值)
                ids_v = self.engine.get_component_similarity(b_char, w_char)
                pix_v = self.engine.get_pixel_iou(b_char, w_char)
                
                self.table.setItem(row, 3, QTableWidgetItem(f"{ids_v:.2f}"))
                self.table.setItem(row, 4, QTableWidgetItem(f"{pix_v:.2f}"))
                self.table.setItem(row, 5, QTableWidgetItem(desc))
                count += 1
        
        self.lbl_status.setText(f"分析完成。共發現 {count} 組形近訛誤 (閾值 >= {threshold})。")

    def open_rule_editor(self):
        """打開規則編輯器 (複用之前的 CustomDictWindow 邏輯，稍作修改)"""
        # 這裡為了簡化，直接使用一個簡單的輸入框，實際可複用
        dialog = QDialog(self)
        dialog.setWindowTitle("設置易混淆部件/字形表")
        dialog.resize(500, 400)
        l = QVBoxLayout(dialog)
        
        txt_edit = QTextEdit()
        txt_edit.setPlaceholderText("格式說明：\n1. 單向替換 (OCR常錯)：rn->m\n2. 雙向形近：日=曰\n\n每行一條。")
        l.addWidget(txt_edit)
        
        btn_box = QHBoxLayout()
        btn_save = QPushButton("保存生效")
        btn_save.clicked.connect(lambda: self.save_rules(txt_edit.toPlainText(), dialog))
        btn_box.addWidget(btn_save)
        l.addLayout(btn_box)
        dialog.exec()

    def save_rules(self, text, dialog):
        repl = {}
        equiv = set()
        for line in text.split('\n'):
            line = line.strip()
            if '->' in line:
                a, b = line.split('->')
                repl[a.strip()] = b.strip()
            elif '=' in line:
                a, b = line.split('=')
                equiv.add((a.strip(), b.strip()))
        
        self.engine.update_rules(repl, equiv)
        QMessageBox.information(self, "成功", "規則已更新至形近字引擎！")
        dialog.accept()

# ==========================================
# [新增組件] 數據源選擇彈窗 (本地 vs OCR)
# ==========================================
class DataSourceSelectionDialog(QDialog):
    """
    [功能]：提供「選擇本地文件」或「啟動 OCR」的入口。
    [返回值]：done(1) = 本地, done(2) = OCR
    """
    def __init__(self, title="選擇數據來源", parent=None):
        super().__init__(parent)
        self.setWindowTitle(title)
        self.resize(400, 200)
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout(self)
        layout.setSpacing(20)
        layout.setContentsMargins(30, 30, 30, 30)

        lbl = QLabel("請選擇數據來源方式：")
        lbl.setStyleSheet("font-size: 16px; font-weight: bold; color: #5D4037;")
        lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(lbl)

        btn_layout = QHBoxLayout()
        
        # 按鈕 1: 本地文件
        self.btn_local = QPushButton("📂 選擇本地文件")
        self.btn_local.setMinimumHeight(60)
        self.btn_local.setStyleSheet("""
            QPushButton { 
                font-size: 16px; font-family: 'KaiTi'; background-color: #EBF5FB; 
                border: 1px solid #AED6F1; border-radius: 8px; color: #2874A6;
            }
            QPushButton:hover { background-color: #D4E6F1; }
        """)
        self.btn_local.clicked.connect(lambda: self.done(1))
        
        # 按鈕 2: 啟動 OCR
        self.btn_ocr = QPushButton("📷 啟動 OCR 整理")
        self.btn_ocr.setMinimumHeight(60)
        self.btn_ocr.setStyleSheet("""
            QPushButton { 
                font-size: 16px; font-family: 'KaiTi'; background-color: #FDEDEC; 
                border: 1px solid #F5B7B1; border-radius: 8px; color: #943126;
            }
            QPushButton:hover { background-color: #FADBD8; }
        """)
        self.btn_ocr.clicked.connect(lambda: self.done(2))

        btn_layout.addWidget(self.btn_local)
        btn_layout.addWidget(self.btn_ocr)
        layout.addLayout(btn_layout)

# ==========================================
# [新增組件] 圖版映射與可視化模塊
# ==========================================
from PyQt6.QtWidgets import QDockWidget, QToolBar

class ImageMapping:
    """
    [核心類] 圖版映射器
    功能：解析 OCR 生成的 JSON，提供從「文本索引」到「圖片路徑」的查找服務。
    """
    def __init__(self, json_path, media_root, is_pdf=False):
        self.json_path = json_path
        self.media_root = media_root # 圖片文件夾路徑 或 PDF文件路徑
        self.is_pdf = is_pdf
        self.data = []
        self.pdf_doc = None
        
        self.load_json()
        
        # 如果是 PDF 模式，預先打開文檔
        if self.is_pdf and os.path.exists(self.media_root):
            import fitz # PyMuPDF
            self.pdf_doc = fitz.open(self.media_root)

    def load_json(self):
        """
        [解析器邏輯 - 嚴格註釋版]
        適配結構：List[Dict]
        Item 結構：
        {
            "file_name": "xxx.jpg", 
            "word_index": { "start": int, "end": int },
            ...
        }
        """
        try:
            with open(self.json_path, 'r', encoding='utf-8') as f:
                raw_data = json.load(f)
                
            # 數據清洗與排序：確保按 start 索引排序，方便二分查找或順序查找
            # 注意：這裡假設 JSON 是一個列表。如果是 Dict，需要調整遍歷邏輯。
            if isinstance(raw_data, list):
                # 過濾掉沒有 word_index 的無效條目
                valid_data = [item for item in raw_data if 'word_index' in item]
                # 按起始位置排序
                self.data = sorted(valid_data, key=lambda x: x['word_index']['start'])
            else:
                print("警告：JSON 根結構不是 List，請檢查解析邏輯。")
                
        except Exception as e:
            print(f"JSON 加載失敗: {e}")

    def get_image_source(self, char_index):
        """
        輸入：全文原始字符索引 (Raw Index)
        輸出：(文件名/頁碼, 圖片對象/PDF頁對象)
        """
        # 遍歷查找 (對於幾百頁的書，順序查找足夠快；若書非常厚，可改為二分查找)
        target_item = None
        for item in self.data:
            idx_info = item.get('word_index', {})
            start = idx_info.get('start', -1)
            end = idx_info.get('end', -1)
            
            # [判定邏輯]：索引落在 [start, end) 區間內
            if start <= char_index < end:
                target_item = item
                break
        
        if not target_item:
            return None, None
            
        file_name = target_item['file_name']
        
        # 資源獲取邏輯
        if self.is_pdf and self.pdf_doc:
            # PDF 模式：嘗試從 file_name 解析頁碼，或者直接按列表索引
            # 假設：OCR 導出的 JSON 順序與 PDF 頁碼順序一致
            # 這裡我們使用 target_item 在 self.data 中的索引作為頁碼 (最穩妥)
            page_idx = self.data.index(target_item)
            try:
                # 返回 pixmap 圖像數據
                page = self.pdf_doc.load_page(page_idx)
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2)) # 2倍縮放保證清晰度
                # 轉換為 QImage 所需格式
                img_data = (pix.samples, pix.width, pix.height, pix.stride)
                return f"PDF 第 {page_idx+1} 頁", img_data
            except:
                return None, None
        else:
            # 圖片文件夾模式
            full_path = os.path.join(self.media_root, file_name)
            return file_name, full_path

class SmartImageViewer(QGraphicsView):
    view_changed_signal = pyqtSignal(object, object)

    def __init__(self, parent=None):
        super().__init__(parent)
        self.scene = QGraphicsScene(self)
        self.setScene(self.scene)
        
        # 基礎設置
        self.setRenderHint(QPainter.RenderHint.Antialiasing) 
        self.setRenderHint(QPainter.RenderHint.SmoothPixmapTransform)
        self.setDragMode(QGraphicsView.DragMode.ScrollHandDrag)
        self.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.setTransformationAnchor(QGraphicsView.ViewportAnchor.AnchorUnderMouse)
        self.setResizeAnchor(QGraphicsView.ViewportAnchor.AnchorUnderMouse)
        
        self.has_image = False
        self._is_syncing = False
        self._auto_fit_mode = True 

        # 創建懸浮控制條
        self.setup_hud()

    def setup_hud(self):
        self.hud_widget = QFrame(self)
        self.hud_widget.setStyleSheet("""
            QFrame {
                background-color: rgba(40, 40, 40, 220);
                border-radius: 18px;
                border: 1px solid rgba(255, 255, 255, 50);
            }
            QPushButton {
                background-color: transparent;
                border: none;
                font-family: Arial;
                font-weight: 900;
                color: #FFFFFF;
                font-size: 18px;
                min-width: 36px;  /* 使用 min-width 確保不被壓縮 */
                min-height: 36px;
                margin: 0px;
                padding: 0px;
            }
            QPushButton:hover { background-color: rgba(255, 255, 255, 50); border-radius: 18px; }
            QPushButton:pressed { background-color: rgba(255, 255, 255, 80); }
        """)
        
        layout = QHBoxLayout(self.hud_widget)
        # 增加一點左右邊距，防止文字貼邊
        layout.setContentsMargins(5, 2, 5, 2)
        layout.setSpacing(0)
        
        btn_out = QPushButton("－")
        btn_out.setToolTip("縮小")
        btn_out.clicked.connect(lambda: self.zoom_step(0.8))
        
        btn_fit = QPushButton("⛶")
        # 使用兼容性 API
        if btn_fit.fontMetrics().horizontalAdvance("⛶") > 30: 
            btn_fit.setText("Fit")
            btn_fit.setStyleSheet("font-size: 14px;") # 文字模式下字體稍小
        
        btn_fit.setToolTip("適應窗口 (自動跟隨大小變化)")
        btn_fit.clicked.connect(self.enable_auto_fit)
        
        btn_in = QPushButton("＋")
        btn_in.setToolTip("放大")
        btn_in.clicked.connect(lambda: self.zoom_step(1.25))
        
        layout.addWidget(btn_out)
        layout.addWidget(btn_fit)
        layout.addWidget(btn_in)
        
        # === 【關鍵修改】不設置固定 resize，而是自適應大小 ===
        self.hud_widget.adjustSize()

    def resizeEvent(self, event):
        """窗口大小改變時，重新計算 HUD 位置"""
        super().resizeEvent(event)
        
        # 1. 始終吸附在右下角 (距離右邊 10px，距離底部 10px)
        if hasattr(self, 'hud_widget'):
            # 確保先讓它算好自己的大小
            self.hud_widget.adjustSize() 
            
            # 使用 viewport 的尺寸來計算，確保不會被滾動條區域影響（雖然我們隱藏了滾動條）
            vp_rect = self.viewport().rect()
            x = vp_rect.width() - self.hud_widget.width() - 10
            y = vp_rect.height() - self.hud_widget.height() - 10
            self.hud_widget.move(x, y)
            
        # 2. 自動適應模式
        if self._auto_fit_mode and self.has_image:
            self.fitInView(self.scene.itemsBoundingRect(), Qt.AspectRatioMode.KeepAspectRatio)

    def wheelEvent(self, event):
        if event.modifiers() & Qt.KeyboardModifier.ControlModifier:
            zoom_in = event.angleDelta().y() > 0
            factor = 1.25 if zoom_in else 0.8
            self.zoom_step(factor)
            event.accept()
        else:
            super().wheelEvent(event)

    def zoom_step(self, factor):
        self._auto_fit_mode = False 
        self.scale(factor, factor)
        self.notify_sync()

    def enable_auto_fit(self):
        self._auto_fit_mode = True
        self.fit_to_window()

    def fit_to_window(self):
        if not self.scene.itemsBoundingRect().isEmpty():
            self.fitInView(self.scene.itemsBoundingRect(), Qt.AspectRatioMode.KeepAspectRatio)
            self.notify_sync()

    def set_image(self, pixmap):
        self.scene.clear()
        if pixmap:
            self.scene.addPixmap(pixmap)
            self.has_image = True
            self.enable_auto_fit()
        else:
            self.has_image = False
            self.scene.addText("無圖像數據")

    def notify_sync(self):
        if not self._is_syncing:
            self.view_changed_signal.emit(self.transform(), self.mapToScene(self.viewport().rect().center()))

    def sync_from_other(self, transform, center_point):
        self._is_syncing = True
        self._auto_fit_mode = False
        self.setTransform(transform)
        self.centerOn(center_point)
        self._is_syncing = False

class EvidenceDock(QDockWidget):
    """
    [UI 組件] 圖版對照停靠窗口 (強力修正版)
    修正點：
    1. 強制解鎖主窗口中央區域的最小寬度限制，確保能佔據 50% 空間。
    2. 每次顯示時都重置大小。
    3. 允許自由拖拽調節。
    """
    def __init__(self, parent=None):
        super().__init__("圖版互證", parent)
        self.setObjectName("EvidenceDock")
        
        # 1. 設置功能特性 (允許浮動、關閉、移動)
        self.setFeatures(
            QDockWidget.DockWidgetFeature.DockWidgetClosable | 
            QDockWidget.DockWidgetFeature.DockWidgetMovable | 
            QDockWidget.DockWidgetFeature.DockWidgetFloatable
        )
        
        # 2. 設置允許停靠的區域
        self.setAllowedAreas(Qt.DockWidgetArea.LeftDockWidgetArea | Qt.DockWidgetArea.RightDockWidgetArea)
        
        # 3. 監聽浮動信號
        self.topLevelChanged.connect(self.on_floating_changed)
        
        # === 界面佈局初始化 ===
        container = QWidget()
        layout = QVBoxLayout(container)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)
        
        # 頂部工具欄
        top_bar = QFrame()
        top_bar.setFixedHeight(34)
        top_bar.setStyleSheet("background-color: #F0F0F0; border-bottom: 1px solid #CCC;")
        
        tb_layout = QHBoxLayout(top_bar)
        tb_layout.setContentsMargins(5, 0, 5, 0)
        tb_layout.setSpacing(10)
        
        self.btn_layout = QPushButton("◫ 切換佈局")
        self.btn_layout.setCursor(Qt.CursorShape.PointingHandCursor)
        self.btn_layout.setStyleSheet("""
            QPushButton { border: 1px solid #AAA; border-radius: 3px; padding: 2px 8px; background: white; color: #333; }
            QPushButton:hover { background: #E0E0E0; }
        """)
        self.btn_layout.clicked.connect(self.toggle_layout)
        
        self.cb_sync = QCheckBox("🔗 同步視圖")
        self.cb_sync.setStyleSheet("font-weight: bold; color: #2E5C8A;")
        self.cb_sync.setChecked(False) 
        
        self.lbl_info = QLabel("就緒")
        self.lbl_info.setStyleSheet("color: #666; font-size: 11px;")
        
        tb_layout.addWidget(self.btn_layout)
        tb_layout.addWidget(self.cb_sync)
        tb_layout.addStretch()
        tb_layout.addWidget(self.lbl_info)
        
        layout.addWidget(top_bar)
        
        # 分割器
        self.splitter = QSplitter(Qt.Orientation.Horizontal)
        self.splitter.setHandleWidth(4)
        
        self.view_base = SmartImageViewer()
        self.view_base.scene.addText("底本視圖區域")
        
        self.view_wit = SmartImageViewer()
        self.view_wit.scene.addText("校本視圖區域")
        
        self.splitter.addWidget(self.view_base)
        self.splitter.addWidget(self.view_wit)
        self.splitter.setStretchFactor(0, 1)
        self.splitter.setStretchFactor(1, 1)
        
        layout.addWidget(self.splitter)
        self.setWidget(container)

        self.view_base.view_changed_signal.connect(lambda t, c: self.sync_views(t, c, source="base"))
        self.view_wit.view_changed_signal.connect(lambda t, c: self.sync_views(t, c, source="wit"))

        # ========================================================
        # 【修正點 1】設置一個極小的最小寬度
        # 這保證了您想把它拖多小就能拖多小，不會被卡住
        # ========================================================
        self.setMinimumWidth(50) 
        
        # 設置 Expanding 策略，確保它有資格和主窗口搶地盤
        policy = QSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        policy.setHorizontalStretch(1) 
        policy.setVerticalStretch(1)
        self.setSizePolicy(policy)

    def showEvent(self, event):
        """
        [交互核心] 每次窗口顯示時觸發，強制執行 50% 分割
        """
        super().showEvent(event)
        
        # 獲取父窗口 (MainWindow)
        main_window = self.parent()
        
        # 僅當停靠在主窗口上時才執行 resizeDocks
        if main_window and isinstance(main_window, QMainWindow) and not self.isFloating():
            
            # ========================================================
            # 【強力修正點 2】解鎖中央部件的寬度限制！
            # 您的截圖顯示右側推不過去，是因為右側部件(CentralWidget)覺得自己"不能再小了"。
            # 這行代碼強制告訴它："你可以變得非常小"，從而允許 Dock 搶佔 50% 的空間。
            # ========================================================
            if main_window.centralWidget():
                main_window.centralWidget().setMinimumWidth(50)

            # 獲取主窗口當前總寬度
            total_width = main_window.width()
            
            # 計算 40% 寬度
            target_width = int(total_width * 0.3)
            
            # 使用 QTimer.singleShot 延遲執行，確保佈局系統準備就緒後再調整
            QTimer.singleShot(10, lambda: main_window.resizeDocks([self], [target_width], Qt.Orientation.Horizontal))

    def on_floating_changed(self, is_floating):
        """處理獨立窗口邏輯"""
        if is_floating:
            flags = Qt.WindowType.Window | \
                    Qt.WindowType.CustomizeWindowHint | \
                    Qt.WindowType.WindowTitleHint | \
                    Qt.WindowType.WindowSystemMenuHint | \
                    Qt.WindowType.WindowMinimizeButtonHint | \
                    Qt.WindowType.WindowMaximizeButtonHint | \
                    Qt.WindowType.WindowCloseButtonHint
            self.setWindowFlags(flags)
            QTimer.singleShot(10, self.show)

    def toggle_layout(self):
        if self.splitter.orientation() == Qt.Orientation.Horizontal:
            self.splitter.setOrientation(Qt.Orientation.Vertical)
            self.btn_layout.setText("⊟ 切換佈局")
        else:
            self.splitter.setOrientation(Qt.Orientation.Horizontal)
            self.btn_layout.setText("◫ 切換佈局")

    def show_image(self, view_type, name, img_source):
        viewer = self.view_base if view_type == 'base' else self.view_wit
        target_name = "底本" if view_type == 'base' else "校本"
        
        pixmap = None
        if isinstance(img_source, str):
            if os.path.exists(img_source):
                pixmap = QPixmap(img_source)
        elif isinstance(img_source, tuple):
            samples, w, h, stride = img_source
            fmt = QImage.Format.Format_RGB888
            qimg = QImage(samples, w, h, stride, fmt)
            pixmap = QPixmap.fromImage(qimg)
            
        if pixmap:
            viewer.set_image(pixmap)
            self.lbl_info.setText(f"{target_name}: {name}")
        else:
            viewer.scene.clear()
            viewer.scene.addText(f"無法加載：{name}")

    def sync_views(self, transform, center_point, source):
        if not self.cb_sync.isChecked(): return
        if source == "base":
            if self.view_wit.has_image: self.view_wit.sync_from_other(transform, center_point)
        else:
            if self.view_base.has_image: self.view_base.sync_from_other(transform, center_point)

# ==========================================
# 7. 主窗口 (UI)
# ==========================================
class MainWindow(QMainWindow):
    """
    [類別說明]：應用程式主窗口。
    [職責]：集成所有功能模塊，處理 UI 佈局、用戶交互、線程調度和數據流轉。
    """
    def __init__(self):
        super().__init__()
        self.setWindowTitle("一人持本 · 古籍自動校勘系統")

        # =================================================
        # 【新增】設置窗口圖標 (左上角)
        # =================================================
        icon_path = get_resource_path(os.path.join("resources", "logo.ico"))
        if os.path.exists(icon_path):
            self.setWindowIcon(QIcon(icon_path))
        # =================================================
        self.suppress_unlock_popup = False
        
        # 自動獲取屏幕大小，並設置為屏幕的 80%
        screen = QApplication.primaryScreen()
        size = screen.availableGeometry()
        self.resize(int(size.width() * 0.8), int(size.height() * 0.8))
        
        # 加載異體字數據庫
        db_path = get_resource_path(os.path.join("resources", "variant_data.json.gz"))
        self.engine = CollationEngine(db_path)
        # 1. 加載異體字數據庫
        db_path = get_resource_path(os.path.join("resources", "variant_data.json.gz"))
        self.engine = CollationEngine(db_path)
        # 2. 加載嵌入式字體
        self.load_embedded_fonts()  
        self.ocr_windows = []
        self.all_data = {} 
        self.base_full_text = ""   
        self.base_clean_text = "" 
        self.wit_full_texts = {} 
        self.wit_clean_texts = {} 
        self.base_map = []       
        self.wit_maps = {} 
        self.wit_cleans = {}      
        
        self.base_path = ""
        self.wit_paths = [] 
        self.cur_variants = [] # 用於存儲當前選中校本的異體字列表
        
        # 應用全局樣式
        self.setStyleSheet(STYLESHEET)
        
        # 設置默認字體
        font = QFont("SimSun", 11)
        QApplication.setFont(font)

        # 加載形近字 IDS 數據
        ids_path = get_resource_path(os.path.join("resources", "ids.txt"))
        # 傳入 embedded_fonts 列表，用於在像素計算時回退
        self.visual_engine = VisualCheckEngine(ids_path, self.font_map, self.available_font_families)

        # 【新增】初始化音韻加載器
        gy_path = get_resource_path(os.path.join("resources"))
        self.phonetic_loader = GuangYunLoader(gy_path)

        self.base_mapper = None # 底本映射器
        self.wit_mappers = {}   # 校本映射器 { '校本名': ImageMapping對象 }

        self.init_ui()
        
        # === [新增] 初始化圖版停靠窗口 (默認隱藏，有數據時顯示) ===
        self.evidence_dock = EvidenceDock(self)
        self.addDockWidget(Qt.DockWidgetArea.LeftDockWidgetArea, self.evidence_dock)
        self.evidence_dock.hide()

    def get_font_for_char(self, char):
        """
        [關鍵技術]：字符級字體回退 (Character-Level Font Fallback)。
        
        [問題背景]：
        古籍中大量存在未進行編碼的生僻字。
        普通的 `QFont` 或 Python 的字符串函數在處理這些需要 4 字節 (Surrogate Pairs) 存儲的字符時，
        往往會出現長度計算錯誤或無法正確匹配字體，導致界面顯示為方框。
        
        [解決方案]：
        使用 Qt 底層的 `QRawFont` 接口。
        不同於 `QFont`，`QRawFont` 直接與底層字體引擎交互，支持通過 `supportsCharacter(char_code)` 
        精確查詢某個具體的 TTF/OTF 文件是否包含該 Unicode 編碼的字形。
        
        [算法流程]：
        1. 優先檢查系統默認宋體 (SimSun)，因為它渲染最快。
        2. 若默認字體不支持，則遍歷 `load_embedded_fonts` 加載的專用字體列表（全宋體）。
        3. 一旦找到支持該字符的字體，立即返回字體名稱。
        """
        if not char: return None
        
        # 獲取該字符的 Unicode 編碼 (整數)
        # 這樣做可以繞過 Python 字符串長度的限制問題
        char_code = ord(char)

        # 1. 為了性能，先看看普通宋體能不能顯示 (針對常用字優化)
        # 如果是普通字，沒必要去輪詢那一大堆字體，直接返回 None (用默認)
        try:
            # 只針對 BMP 範圍內的字 (普通字) 做快速檢查
            if char_code <= 0xFFFF:
                fm_default = QFontMetrics(QFont("SimSun"))
                if fm_default.inFont(char):
                    return None 
        except:
            pass # 如果出錯就忽略，進入下面的強力搜索模式

        # 2. 強力搜索：輪詢我們加載的"字體軍火庫"
        for font_name in self.available_font_families:
            try:
                # 創建底層字體對象
                raw_font = QRawFont.fromFont(QFont(font_name))
                
                # 直接詢問：這個字體裡有沒有這個編碼對應的字形？
                if raw_font.supportsCharacter(char_code):
                    # 找到了！就是這個字體！
                    return font_name
            except Exception:
                continue # 萬一某個字體加載出問題，跳過，試下一個

        # 3. 如果遍歷了所有字體都沒找到，返回 None (Word 會顯示方框)
        return None
    
    def insert_smart_text(self, paragraph, text, color_rgb=None, is_bold=False):
        """
        [輔助函數]：Word導出的智能写入封装。
        [功能]：
        1. 逐字檢查字符需要的字體。
        2. 自動應用字體映射（英文名轉中文名）。
        3. 應用顏色和粗體。
        這樣保證文檔中【所有】文字（包括普通文本）都能正確顯示生僻字。
        """
        if not text: return

        for char in text:
            run = paragraph.add_run(char)
            
            # 1. 設置樣式
            if color_rgb:
                run.font.color.rgb = color_rgb
            if is_bold:
                run.font.bold = True
            
            # 2. 智能字體匹配
            best_font = self.get_font_for_char(char)
            
            if best_font:
                # 如果找到了特殊字體，使用映射表轉換成 Word 能認的中文名
                final_name = self.font_map.get(best_font, best_font)
                run.font.name = final_name
                run.element.rPr.rFonts.set(qn('w:eastAsia'), final_name)
            else:
                # 普通字體，回退到標準宋體
                run.font.name = "Times New Roman"
                run.element.rPr.rFonts.set(qn('w:eastAsia'), "宋體")



    def load_embedded_fonts(self):
        """加載所有字體，並建立可用字體列表及映射表"""
        font_dir = get_resource_path(os.path.join("resources", "fonts"))
        
        self.available_font_families = [] 
        # === 字體名稱映射表 ===
        # 這是爲了解決 Word 不認英文名，只認中文名的問題
        self.font_map = {
            "FSung-1": "全宋体-1",
            "FSung-2": "全宋体-2",
            "FSung-3": "全宋体-3",
            "FSung-F": "全宋体-F",
            "FSung-m": "全宋体(等宽)",
            "FSung-p": "全宋体(调和)",
            "FSung-X": "全宋体-X",
        }
        
        if not os.path.exists(font_dir):
            print(f"警告：找不到字體目錄 {font_dir}")
            return

        loaded_count = 0
        print("-" * 30)
        print("正在分析字體文件...")
        try:
            for filename in os.listdir(font_dir):
                if filename.lower().endswith((".ttf", ".otf", ".woff")):
                    font_path = os.path.join(font_dir, filename)
                    font_id = QFontDatabase.addApplicationFont(font_path)
                    
                    if font_id != -1:
                        loaded_count += 1
                        families = QFontDatabase.applicationFontFamilies(font_id)
                        if families:
                            real_name = families[0]
                            if real_name not in self.available_font_families:
                                self.available_font_families.append(real_name)
                            print(f"文件名: {filename}  --->  Word識別名: 【{real_name}】")
            
            # 追加系統後備字體
            system_backups = ["HanaMinB", "SimSun-ExtB", "Microsoft YaHei"]
            for f in system_backups:
                 if f not in self.available_font_families:
                     self.available_font_families.append(f)
                     
            print("-" * 30)
            print(f"最終可用字體庫: {self.available_font_families}")
            print("-" * 30)
        except Exception as e:
            print(f"字體加載過程出錯: {e}")

    def init_ui(self):
        main_widget = QWidget()
        self.setCentralWidget(main_widget)
        main_layout = QVBoxLayout(main_widget)
        main_layout.setContentsMargins(20, 20, 20, 20)
        main_layout.setSpacing(15)

        # 頂部標題區 
        title_label = QLabel("一人持本 · 古籍自動校勘系統")
        title_label.setFont(QFont("KaiTi", 24, QFont.Weight.Bold))
        title_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        title_label.setStyleSheet("color: #2B2B2B; margin-bottom: 5px;")
        main_layout.addWidget(title_label)

        # A. 文件選擇區
        file_group = QGroupBox(" 選擇底本與校本 ")
        file_layout = QHBoxLayout(file_group)
        file_layout.setContentsMargins(15, 25, 15, 15)
        file_layout.setSpacing(15)
        
        # 左側：底本
        base_layout = QVBoxLayout()
        self.btn_base = QPushButton(" 📂 選擇底本文件 ")
        self.btn_base.setMinimumHeight(40)
        self.btn_base.clicked.connect(self.upload_base)
        self.lbl_base = QLabel("（尚未選擇）")
        self.lbl_base.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.lbl_base.setStyleSheet("color: #888; font-style: italic;")
        base_layout.addWidget(self.btn_base)
        base_layout.addWidget(self.lbl_base)
        file_layout.addLayout(base_layout, stretch=1)

        # 分割線
        line = QFrame()
        line.setFrameShape(QFrame.Shape.VLine)
        line.setFrameShadow(QFrame.Shadow.Sunken)
        line.setStyleSheet("color: #DDD;")
        file_layout.addWidget(line)

        # 中間：校本列表
        wit_layout = QVBoxLayout()
        self.list_wits = QListWidget()
        self.list_wits.setStyleSheet("border: 1px solid #DDD; border-radius: 4px; background: #FFF;")
        wit_layout.addWidget(QLabel("校本列表："))
        wit_layout.addWidget(self.list_wits)
        
        btn_tools = QHBoxLayout()
        self.btn_add = QPushButton(" ＋ 添加 ")
        self.btn_add.clicked.connect(self.add_wit)
        self.btn_del = QPushButton(" － 刪除 ")
        self.btn_del.clicked.connect(self.remove_wit)
        self.btn_clear = QPushButton(" 🗑 清空 ")
        self.btn_clear.clicked.connect(self.clear_wits)
        btn_tools.addWidget(self.btn_add)
        btn_tools.addWidget(self.btn_del)
        btn_tools.addWidget(self.btn_clear)
        wit_layout.addLayout(btn_tools)
        file_layout.addLayout(wit_layout, stretch=2)

        # 右側：執行區
        run_layout = QVBoxLayout()
        self.cb_variant_filter = QCheckBox("異體字單獨出校")
        self.cb_variant_filter.setChecked(True)
        
        self.btn_run = QPushButton("開始\n 校勘 ")
        self.btn_run.setObjectName("primary_btn")
        self.btn_run.setFixedSize(70, 70)
        self.btn_run.clicked.connect(self.start_collation)
        self.btn_run.setEnabled(False)
        
        self.lbl_status = QLabel("請先添加文件")
        self.lbl_status.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.lbl_status.setStyleSheet("color: #666; margin-top: 5px;")
        
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        self.progress_bar.setFixedWidth(85)
        
        run_container = QVBoxLayout()
        run_container.addStretch()
        run_container.addWidget(self.cb_variant_filter, alignment=Qt.AlignmentFlag.AlignCenter)
        run_container.addWidget(self.btn_run, alignment=Qt.AlignmentFlag.AlignCenter)
        run_container.addWidget(self.lbl_status, alignment=Qt.AlignmentFlag.AlignCenter)
        run_container.addWidget(self.progress_bar, alignment=Qt.AlignmentFlag.AlignCenter)
        run_container.addStretch()
        
        file_layout.addLayout(run_container, stretch=1)
        main_layout.addWidget(file_group)

        # C. 內容展示區 (Tab)
        self.tabs = QTabWidget()
        self.tabs.currentChanged.connect(self.on_tab_changed)
 
        # === 【插入點：右上角自定義按鈕】 ===
        top_btn_widget = QWidget()
        top_btn_layout = QHBoxLayout(top_btn_widget)
        top_btn_layout.setContentsMargins(0, 0, 0, 0)
        
        # --- [新增] 移動來的校本切換控件 ---
        self.lbl_view_hint = QLabel("當前查看校本：")
        self.lbl_view_hint.setStyleSheet("font-family: 'KaiTi'; font-weight: bold; color: #5D4037;")
        
        self.combo_wits = QComboBox()
        self.combo_wits.setMinimumWidth(150)
        # 連接切換信號
        self.combo_wits.currentIndexChanged.connect(self.switch_view)
        
        # 初始化指針 (這是原 Tab 1 裡的代碼，移到這裡)
        self.last_wit_index = 0

        # 將它們加入右上角佈局 (放在最前面)
        top_btn_layout.addWidget(self.lbl_view_hint)
        top_btn_layout.addWidget(self.combo_wits)
        # --------------------------------

        self.btn_custom_dict = QPushButton("⚙ 自定義規則")
        self.btn_custom_dict.setStyleSheet("font-family: 'KaiTi'; font-weight: bold; color: #5D4037;")
        self.btn_custom_dict.clicked.connect(self.open_custom_dict)
        top_btn_layout.addWidget(self.btn_custom_dict)

        # 1. 創建按鈕
        self.btn_open_dict = QPushButton(" 查閱異體字字典 ")
        self.btn_open_dict.setFixedSize(140, 30) 
        self.btn_open_dict.setCursor(Qt.CursorShape.PointingHandCursor) # 讓鼠標放上去變手型

        # 2. 設置樣式
        self.btn_open_dict.setStyleSheet("""
            QPushButton {
                background-color: transparent; 
                color: #5D4037; 
                border: 1px solid #C0C0C0; 
                border-radius: 4px;
                font-family: 'KaiTi'; font-weight: bold;
                margin-right: 5px; 
                margin-top: 2px;
            }
            QPushButton:hover {
                background-color: #E0E8F0;
                color: #2E5C8A;
                border-color: #2E5C8A;
            }
        """)

        # 3. 連接功能
        self.btn_open_dict.clicked.connect(self.open_dictionary)

        # 4. 【關鍵修改】將查閱字典按鈕也加入到這個佈局中！
        top_btn_layout.addWidget(self.btn_open_dict) 

        # 5. 【關鍵修改】設置角落控件為這個「包含兩個按鈕的容器」，而不是單個按鈕
        self.tabs.setCornerWidget(top_btn_widget, Qt.Corner.TopRightCorner)

        # 6. 加入主佈局
        main_layout.addWidget(self.tabs)
        # ========================================================

        # --- Tab 1: 異文與異體 (二級 Tab) ---
        self.tab_diff_root = QWidget()
        root_layout = QVBoxLayout(self.tab_diff_root)

        self.sub_tabs = QTabWidget()
        self.sub_tabs.setObjectName("SubTabs")
        root_layout.addWidget(self.sub_tabs)    
        # 子 Tab 1: 訛脫衍倒
        main_diff_widget = QWidget()
        main_diff_layout = QVBoxLayout(main_diff_widget)
        filter_layout = QHBoxLayout()
        filter_lbl = QLabel("顯示類型篩選：")
        filter_lbl.setStyleSheet("color: #5D4037; font-weight: bold;")
        filter_layout.addWidget(filter_lbl)

        self.type_filters = {} # 用於存儲複選框對象
        filter_types = ['訛', '脫', '衍', '倒', '異']
        for t in filter_types:
            cb = QCheckBox(t)
            cb.setChecked(True) # 默認全部選中
            cb.setStyleSheet("font-family: 'KaiTi'; font-size: 14px; margin-right: 10px;")
            # 當狀態改變時，觸發 switch_view 重新渲染表格
            # 使用 lambda 傳遞當前 combo_wits 的索引，避免參數錯誤
            cb.stateChanged.connect(lambda: self.switch_view(self.combo_wits.currentIndex()))
            
            self.type_filters[t] = cb
            filter_layout.addWidget(cb)        
        filter_layout.addStretch() # 彈簧，將複選框頂到左邊
        main_diff_layout.addLayout(filter_layout)
        self.table_main = QTableWidget()
        self.table_main.setAlternatingRowColors(True)
        self.table_main.setColumnCount(5)
        self.table_main.setHorizontalHeaderLabels(["底本原文", "校本原文", "類型", "底本", "校本"])
        self.table_main.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        self.table_main.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.table_main.customContextMenuRequested.connect(self.show_main_table_context_menu)
        self.table_main.cellDoubleClicked.connect(self.on_diff_table_double_click)
        main_diff_layout.addWidget(self.table_main)
        
        self.btn_export_excel = QPushButton("  導出異文對照表 (Excel) ")
        self.btn_export_excel.setObjectName("export_excel")
        self.btn_export_excel.clicked.connect(self.export_excel)
        main_diff_layout.addWidget(self.btn_export_excel, alignment=Qt.AlignmentFlag.AlignRight)
        
        self.sub_tabs.addTab(main_diff_widget, "異文對照")

        # 子 Tab 2: 異體字對照
        self.var_panel = QWidget()
        var_layout = QVBoxLayout(self.var_panel)
        
        # 頻次統計區 (橫向滾動)
        self.stats_scroll = QScrollArea()
        self.stats_scroll.setFixedHeight(80)
        self.stats_scroll.setWidgetResizable(True)
        self.stats_cont = QWidget()
        self.stats_lay = QHBoxLayout(self.stats_cont)
        self.stats_lay.setSpacing(10)
        self.stats_scroll.setWidget(self.stats_cont)
        var_layout.addWidget(QLabel("異體字頻次統計 (點擊可過濾)："))
        var_layout.addWidget(self.stats_scroll)
        
        self.btn_sort = QPushButton(" 按頻次排序表格內容")
        self.btn_sort.clicked.connect(self.sort_var_table)
        var_layout.addWidget(self.btn_sort)
        
        self.table_var = QTableWidget()
        self.table_var.setAlternatingRowColors(True)
        self.table_var.setColumnCount(9)
        self.table_var.setHorizontalHeaderLabels(["序號", "底本上下文", "底本文字", "文字屬性", "校本上下文", "校本文字", "文字屬性", "次數", "情況說明"])
        self.table_var.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        var_layout.addWidget(self.table_var)
        self.sub_tabs.addTab(self.var_panel, "異體字對照")
        # --- 新增 Tab 3: 複雜訛誤審覈 ---
        self.complex_panel = ComplexAuditPanel()
        # 連接數據變更信號 -> 重新加載所有視圖
        self.complex_panel.data_changed_signal.connect(self.reload_all_views)
        # [新增]：連接解鎖信號
        self.complex_panel.audit_finished_signal.connect(self.toggle_displacement_tab)
        self.sub_tabs.addTab(self.complex_panel, "複雜訛誤審覈")

        # 2. 添加錯簡匹配面板
        self.disp_panel = DisplacementMatchPanel()
        idx = self.sub_tabs.addTab(self.disp_panel, "錯簡匹配 🔒")
        self.sub_tabs.setTabEnabled(idx, False) 
        # [新增代碼]：安裝事件過濾器 (這是監聽禁用 Tab 的唯一方法)
        self.sub_tabs.tabBar().installEventFilter(self)

        # === 【初始化避諱檢測 Tab】 ===
        # 獲取 JSON 路徑
        taboo_json_path = get_resource_path(os.path.join("resources", "taboo_data.json"))
        # 傳入路徑進行初始化
        self.taboo_panel = TabooCheckerPanel(taboo_json_path)
        self.sub_tabs.addTab(self.taboo_panel, "避諱檢測")

        # --- 新增 Tab 6: 形近誤字 ---
        self.visual_panel = VisualAnalysisPanel(self.visual_engine)
        self.sub_tabs.addTab(self.visual_panel, "形近誤字")

        # --- 新增 Tab 7: 音近通假 ---
        self.phonetic_panel = PhoneticAnalysisPanel(self.phonetic_loader, self.engine) 
        self.sub_tabs.addTab(self.phonetic_panel, "音近通假")

        self.tabs.addTab(self.tab_diff_root, "  異文對照  ")

        # --- Tab 2: 原文校勘標注 ---
        tab2_widget = QWidget()
        tab2_layout = QVBoxLayout(tab2_widget)
        # 頂部工具欄：顯示異體字標注開關
        tool_bar_t2 = QHBoxLayout()
        self.cb_show_variants_in_text = QCheckBox("在文中顯示異體字標注 (綠字顯示)")
        self.cb_show_variants_in_text.setChecked(True) # 默認開啟
        self.cb_show_variants_in_text.setStyleSheet("font-family: 'KaiTi'; font-size: 14px; color: #5D4037;")
        # 連接信號：點擊後直接刷新預覽，無需傳參
        self.cb_show_variants_in_text.stateChanged.connect(lambda: self.render_text_preview())
        tool_bar_t2.addWidget(self.cb_show_variants_in_text)
        tool_bar_t2.addStretch() # 彈簧，讓勾選框靠左
        tab2_layout.addLayout(tool_bar_t2)
        self.text_preview = QTextEdit()
        self.text_preview.setReadOnly(True)
        self.text_preview.setFont(QFont("SimSun", 12))
        self.text_preview.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.text_preview.customContextMenuRequested.connect(self.show_text_preview_menu)
        tab2_layout.addWidget(self.text_preview)
        
        self.btn_export_word_anno = QPushButton(" 📝 導出標注文檔 (Word) ")
        self.btn_export_word_anno.clicked.connect(self.export_annotated_word)
        tab2_layout.addWidget(self.btn_export_word_anno, alignment=Qt.AlignmentFlag.AlignRight)
        self.tabs.addTab(tab2_widget, "  原文標注  ")

        # --- Tab 3: 統計數據 ---
        self.stats_panel = StatsPanel()
        self.tabs.addTab(self.stats_panel, "  數據統計  ")

        # --- Tab 4: 校勘記管理與預覽  ---
        self.tab4_widget = QWidget()
        tab4_layout = QVBoxLayout(self.tab4_widget)
        
        # 【新增】初始化撤銷棧
        self.undo_stack = QUndoStack(self)

        # 1. 頂部工具欄
        t4_toolbar = QHBoxLayout()
        
        self.cb_show_diff_notes = QCheckBox("顯示異體字校勘記")
        self.cb_show_diff_notes.setChecked(True)
        self.cb_show_diff_notes.stateChanged.connect(self.on_display_option_changed)
        
        self.cb_highlight_diff = QCheckBox("異體字紅字高亮")
        self.cb_highlight_diff.setChecked(False)
        self.cb_highlight_diff.stateChanged.connect(self.refresh_note_preview)
        
        t4_toolbar.addWidget(self.cb_show_diff_notes)
        t4_toolbar.addWidget(self.cb_highlight_diff)
        
        # --- 撤銷/恢復按鈕 ---
        line_split = QFrame(); line_split.setFrameShape(QFrame.Shape.VLine); line_split.setFrameShadow(QFrame.Shadow.Sunken)
        t4_toolbar.addWidget(line_split)
        
        # 創建 Action 方便綁定快捷鍵
        self.action_undo = self.undo_stack.createUndoAction(self, "撤銷")
        self.action_undo.setShortcut(QKeySequence.StandardKey.Undo)
        self.action_redo = self.undo_stack.createRedoAction(self, "恢復")
        self.action_redo.setShortcut(QKeySequence.StandardKey.Redo)
        
        btn_undo = QPushButton(" ↶ 撤銷 ")
        btn_undo.addAction(self.action_undo) 
        btn_undo.clicked.connect(self.action_undo.trigger) # 連接點擊事件
        
        btn_redo = QPushButton(" ↷ 恢復 ")
        btn_redo.addAction(self.action_redo)
        btn_redo.clicked.connect(self.action_redo.trigger)
        
        t4_toolbar.addWidget(btn_undo)
        t4_toolbar.addWidget(btn_redo)
        t4_toolbar.addStretch()
        # -----------------------------
        
        self.btn_export_word_note = QPushButton(" 📖 導出底本附校勘記 (Word) ")
        self.btn_export_word_note.clicked.connect(self.export_real_footnote)
        t4_toolbar.addWidget(self.btn_export_word_note)
        
        tab4_layout.addLayout(t4_toolbar)

        # 2. 左右分欄 (Splitter) 
        self.note_splitter = QSplitter(Qt.Orientation.Horizontal)
        
        # 左側：預覽區
        preview_group = QGroupBox("效果預覽")
        preview_group.setStyleSheet("QGroupBox { font-family: 'KaiTi'; font-weight: bold; color: #5D4037; }")
        p_layout = QVBoxLayout(preview_group)
        self.notes_preview = QTextEdit()
        self.notes_preview.setReadOnly(True)
        self.notes_preview.setFont(QFont("SimSun", 11))
        p_layout.addWidget(self.notes_preview)
        
        # 右側：管理區
        manage_group = QGroupBox("校勘記數據管理 (雙擊修改，右鍵刪除)")
        manage_group.setStyleSheet("QGroupBox { font-family: 'KaiTi'; font-weight: bold; color: #5D4037; }")
        m_layout = QVBoxLayout(manage_group)
        
        self.table_notes = QTableWidget()
        self.table_notes.setColumnCount(4)
        self.table_notes.setHorizontalHeaderLabels(["ID", "位置", "類型", "校勘記內容"])
        self.table_notes.horizontalHeader().setSectionResizeMode(3, QHeaderView.ResizeMode.Stretch)
        self.table_notes.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        self.table_notes.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.table_notes.customContextMenuRequested.connect(self.show_note_context_menu)
        
        # 連接新的信號槽
        self.table_notes.itemChanged.connect(self.on_note_content_changed)
        
        m_layout.addWidget(self.table_notes)
        
        self.note_splitter.addWidget(preview_group)
        self.note_splitter.addWidget(manage_group)
        self.note_splitter.setStretchFactor(0, 4)
        self.note_splitter.setStretchFactor(1, 6)
        
        tab4_layout.addWidget(self.note_splitter)
        self.tabs.addTab(self.tab4_widget, " 校勘記管理 ")

    def on_tab_changed(self, index):
        tab_title = self.tabs.tabText(index)
        # 簡單的判斷：如果是後兩個 Tab，隱藏頂部下拉框；否則顯示
        if "校勘記管理" in tab_title or "數據統計" in tab_title:
            self.lbl_view_hint.setVisible(False)
            self.combo_wits.setVisible(False)
        else:
            self.lbl_view_hint.setVisible(True)
            self.combo_wits.setVisible(True)

    # --- 邏輯函數 ---
    def upload_base(self):
        # 1. 彈出選擇對話框
        dlg = DataSourceSelectionDialog("選擇底本來源", self)
        res = dlg.exec()
        
        if res == 1: # 選擇本地文件
            path, _ = QFileDialog.getOpenFileName(self, "選擇底本", "", "Word (*.docx)")
            if path:
                self.base_path = path
                self.lbl_base.setText(os.path.basename(path))
                self.lbl_base.setStyleSheet("color: #2B2B2B; font-weight: bold;")
                self.check_ready()
                self.ask_load_mapping('base', path)

        elif res == 2: # 啟動 OCR
            self.launch_ocr_for_target("base")

    # === [修改] 添加校本 (含 OCR 入口) ===
    def add_wit(self):
        dlg = DataSourceSelectionDialog("添加校本來源", self)
        res = dlg.exec()
        
        if res == 1: # 選擇本地文件
            paths, _ = QFileDialog.getOpenFileNames(self, "添加校本", "", "Word (*.docx)")
            if paths:
                for p in paths:
                    self.add_wit_from_path(p)
                    self.ask_load_mapping('wit', p)
                self.check_ready()
                
        elif res == 2: # 啟動 OCR
            self.launch_ocr_for_target("wit")

    def ask_load_mapping(self, target_type, doc_path):
        """
        詢問用戶是否加載圖版映射 JSON
        target_type: 'base' or 'wit'
        doc_path: Word 文件路徑 (用於自動猜測同名 json)
        """
        # 自動猜測路徑: word文件名_info.json
        guess_json = os.path.splitext(doc_path)[0] + "_info.json"
        
        msg = "是否關聯圖版映射文件 (JSON) 以啟用「圖文互證」功能？\n"
        if os.path.exists(guess_json):
            msg += f"\n檢測到同目錄下存在：{os.path.basename(guess_json)}\n點擊「Yes」自動加載，點擊「No」手動選擇。"
        
        reply = QMessageBox.question(self, "關聯圖版", msg, 
                                     QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No | QMessageBox.StandardButton.Cancel)
        
        if reply == QMessageBox.StandardButton.Cancel: return
        
        json_path = ""
        media_path = "" # 圖片文件夾 或 PDF路徑
        is_pdf = False
        
        # A. 確定 JSON 路徑
        if reply == QMessageBox.StandardButton.Yes and os.path.exists(guess_json):
            json_path = guess_json
        else:
            json_path, _ = QFileDialog.getOpenFileName(self, "選擇映射 JSON", os.path.dirname(doc_path), "JSON (*.json)")
            
        if not json_path: return
        
        # B. 確定 圖片/PDF 源
        # 詢問資源類型
        type_reply = QMessageBox.question(self, "資源類型", "對應的圖版資源是 PDF 文件嗎？\n(Yes=PDF文件, No=圖片文件夾)",
                                          QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        is_pdf = (type_reply == QMessageBox.StandardButton.Yes)
        
        if is_pdf:
            media_path, _ = QFileDialog.getOpenFileName(self, "選擇對應 PDF", os.path.dirname(doc_path), "PDF (*.pdf)")
        else:
            media_path = QFileDialog.getExistingDirectory(self, "選擇圖片文件夾", os.path.dirname(doc_path))
            
        if not media_path: return
        
        # C. 創建映射器
        mapper = ImageMapping(json_path, media_path, is_pdf)
        
        if target_type == 'base':
            self.base_mapper = mapper
        else:
            # 使用文件名作為 Key
            fname = os.path.basename(doc_path)
            # 注意：這裡需要和 WorkerThread 裡的 clean_filename 邏輯保持一致
            # 最好在 add_wit_from_path 時就存儲這個映射名，這裡簡化處理
            if not fname.endswith("本"): fname_key = os.path.splitext(fname)[0] + "本"
            else: fname_key = os.path.splitext(fname)[0]
            
            self.wit_mappers[fname_key] = mapper
            
        self.evidence_dock.show() # 成功加載後顯示 Dock
        QMessageBox.information(self, "成功", "圖版關聯成功！\n在異文對照表中雙擊文字即可回查。")

    # === [新增] 內部輔助：直接通過路徑添加校本 ===
    def add_wit_from_path(self, path):
        # 避免重複添加 (可選)
        # existing = [self.list_wits.item(i).data(Qt.ItemDataRole.UserRole) for i in range(self.list_wits.count())]
        # if path in existing: return 
        
        self.list_wits.addItem(os.path.basename(path))
        self.list_wits.item(self.list_wits.count()-1).setData(Qt.ItemDataRole.UserRole, path)

    # === [新增] 啟動 OCR 窗口邏輯 ===
    def launch_ocr_for_target(self, target_type):
        """
        啟動一個新的 OCR 窗口，並根據目標類型(base/wit)綁定不同的回調槽。
        target_type: 'base' or 'wit'
        """
        ocr_win = OCRMainWindow()
        
        # 綁定信號：當 OCR 導出文件時，調用對應的處理函數
        if target_type == "base":
            ocr_win.file_exported.connect(self.on_ocr_base_exported)
        else:
            ocr_win.file_exported.connect(self.on_ocr_wit_exported)
            
        # 顯示窗口 (非模態)
        ocr_win.show()
        
        # 存入列表防止回收
        self.ocr_windows.append(ocr_win)

    # === [新增] OCR 回調槽：處理底本 ===
    def on_ocr_base_exported(self, file_path):
        # 詢問用戶是否自動對接
        reply = QMessageBox.question(
            self, "OCR 導出成功", 
            f"OCR 系統已導出文件：\n{os.path.basename(file_path)}\n\n是否將其設置為【底本】？",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )
        
        if reply == QMessageBox.StandardButton.Yes:
            self.base_path = file_path
            self.lbl_base.setText(os.path.basename(file_path))
            self.lbl_base.setStyleSheet("color: #2B2B2B; font-weight: bold;")
            self.check_ready()
            self.ask_load_mapping('base', file_path)
            QMessageBox.information(self, "已加載", "已自動設置為底本。")

    # === [新增] OCR 回調槽：處理校本 ===
    def on_ocr_wit_exported(self, file_path):
        # 詢問用戶是否自動對接
        reply = QMessageBox.question(
            self, "OCR 導出成功", 
            f"OCR 系統已導出文件：\n{os.path.basename(file_path)}\n\n是否將其追加到【校本列表】？",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )
        
        if reply == QMessageBox.StandardButton.Yes:
            self.add_wit_from_path(file_path)
            self.check_ready()
            self.ask_load_mapping('wit', file_path)
            QMessageBox.information(self, "已加載", "已自動追加到校本列表。")

    def remove_wit(self):
        for item in self.list_wits.selectedItems():
            self.list_wits.takeItem(self.list_wits.row(item))
        self.check_ready()

    def clear_wits(self):
        self.list_wits.clear()
        self.check_ready()

    def check_ready(self):
        self.wit_paths = [self.list_wits.item(i).data(Qt.ItemDataRole.UserRole) for i in range(self.list_wits.count())]
        if self.base_path and self.wit_paths:
            self.btn_run.setEnabled(True)
            self.lbl_status.setText("準備就緒")
            self.lbl_status.setStyleSheet("color: #5C7A62; font-weight: bold;")
        else:
            self.btn_run.setEnabled(False)
            self.lbl_status.setText("請添加文件")
            self.lbl_status.setStyleSheet("color: #666;")

    def start_collation(self):
        self.btn_run.setEnabled(False)
        self.tabs.setCurrentIndex(0)
        self.worker = WorkerThread(self.base_path, self.wit_paths, self.engine, self.cb_variant_filter.isChecked())
        self.worker.progress_update.connect(lambda s: self.lbl_status.setText(s))
        self.worker.finished.connect(self.on_finished)
        self.worker.error.connect(self.on_error)
        self.worker.start()

    def on_finished(self, all_data, base_raw, base_clean, wit_originals, base_map, wit_maps, wit_cleans):
        self.all_data = all_data
        self.base_full_text = base_raw   
        self.base_clean_text = base_clean 
        self.wit_full_texts = wit_originals
        self.wit_clean_texts = wit_cleans
        self.base_map = base_map
        self.wit_maps = wit_maps
        self.wit_cleans = wit_cleans
        self.taboo_panel.load_text_data(self.base_clean_text, self.wit_cleans)
        
        self.lbl_status.setText("校勘完成")
        self.lbl_status.setStyleSheet("color: #B74639; font-weight: bold;")
        self.btn_run.setEnabled(True)
        
        self.combo_wits.clear()
        self.combo_wits.addItems(list(all_data.keys()))
        # === 【修改開始】 ===
        # 1. 開啟抑制標誌 (防止 complex_panel 加載數據時觸發的信號導致彈窗)
        self.suppress_unlock_popup = True
        
        curr_wit_clean = self.wit_clean_texts.get(self.combo_wits.currentText(), "")
        
        # 2. 加載複雜訛誤板塊 (這可能會發出 audit_finished_signal -> toggle_displacement_tab)
        self.complex_panel.load_data(self.combo_wits.currentText(), self.all_data, self.base_clean_text, curr_wit_clean)
        
        # 3. 如果錯簡 Tab 被解鎖了 (Tab Enabled=True)，手動加載一次錯簡數據
        # 因為 toggle_displacement_tab 裡的加載被抑制邏輯覆蓋了或者需要確保數據最新
        if self.sub_tabs.isTabEnabled(3):
             self.disp_panel.load_data(self.combo_wits.currentText(), self.all_data, self.base_clean_text, curr_wit_clean)

        # 4. 關閉抑制標誌 (恢復正常手動操作時的彈窗)
        self.suppress_unlock_popup = False
        is_variant_enabled = self.cb_variant_filter.isChecked()
        self.stats_panel.update_data(all_data, is_variant_enabled, base_clean)
        # --- 初始化校勘記管理數據 ---
        self.init_note_records() 
        # -------------------------------
        self.preview_notes_html()
        QMessageBox.information(self, "完成", "校勘工作已完成！")

    def on_error(self, msg):
        self.lbl_status.setText("出錯")
        self.btn_run.setEnabled(True)
        QMessageBox.critical(self, "錯誤", msg)

    def open_dictionary(self):
        if not hasattr(self, 'dict_window') or self.dict_window is None:
            self.dict_window = DictionaryWindow(self.engine)
        self.dict_window.show()
        self.dict_window.raise_()
        self.dict_window.activateWindow()

    # --- 視圖切換與渲染 ---
    def switch_view(self, index):
        if index < 0: return
        # =================================================================
        # 【新增】安全攔截邏輯：放在獲取數據之前
        # =================================================================
        # 檢查是否切換了索引，且是否存在未保存的撤銷記錄
        if index != self.last_wit_index and hasattr(self, 'complex_panel') and self.complex_panel.history_stack:
            reply = QMessageBox.question(
                self, "確認切換",
                "注意：您在當前校本的「複雜訛誤審覈」中有可撤銷的操作記錄。\n\n切換校本將清空這些記錄，您將無法撤銷剛才的應用。\n確定要繼續切換嗎？",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                QMessageBox.StandardButton.No
            )

            if reply == QMessageBox.StandardButton.No:
                # 用戶後悔了 -> 恢復下拉菜單的選項
                self.combo_wits.blockSignals(True) # 暫停信號，防止遞歸
                self.combo_wits.setCurrentIndex(self.last_wit_index) # 撥回原位
                self.combo_wits.blockSignals(False) # 恢復信號
                return # 直接結束，不執行下面的代碼

        # 更新記錄指針 (這步很重要，要在攔截通過後執行)
        self.last_wit_index = index
        wit_name = self.combo_wits.currentText()
        data = self.all_data.get(wit_name)
        if not data: return
        
        results = data['results']
        variants = data['variants']
        curr_wit_clean = self.wit_clean_texts.get(wit_name, "")

        # =========================================
        # 1. 渲染主異文表 (訛脫衍倒)
        # =========================================
        self.table_main.setRowCount(0)
        # 過濾出非 Match 的異文
        diffs = [r for r in results if r['type'] != 'Match']
        
        for i, r in enumerate(diffs):
            r_type = r['type']
            # 如果該類型的複選框存在，且未被勾選，則跳過此行
            if r_type in self.type_filters:
                if not self.type_filters[r_type].isChecked():
                    continue
            row = self.table_main.rowCount()
            self.table_main.insertRow(row)
            
            # --- 上下文計算 ---
            idx = r['idx']
            w_idx = r['wit_idx']
            
            # 底本上下文
            start_pre = max(0, idx - 2)
            pre_text = self.base_clean_text[start_pre : idx]
            center_base = r['base_clean'] if r['base_clean'] else "【無】"
            start_post = idx + r['base_clean_len']
            post_text = self.base_clean_text[start_post : min(len(self.base_clean_text), start_post + 2)]
            
            # 校本上下文
            w_start_pre = max(0, w_idx - 2)
            w_pre_text = curr_wit_clean[w_start_pre : w_idx]
            center_wit = r['wit_clean'] if r['wit_clean'] else "【無】"
            w_start_post = w_idx + r['wit_clean_len']
            w_post_text = curr_wit_clean[w_start_post : min(len(curr_wit_clean), w_start_post + 2)]

            # --- 構建 HTML 樣式 (恢復原樣) ---
            # 樣式：前後文灰色(#666)，中間字朱砂紅(#B74639)加粗
            html_base = f"<html><span style='color:#666'>{pre_text}</span><span style='color:#B74639; font-weight:bold; font-size:14px;'>{center_base}</span><span style='color:#666'>{post_text}</span></html>"
            html_wit = f"<html><span style='color:#666'>{w_pre_text}</span><span style='color:#B74639; font-weight:bold; font-size:14px;'>{center_wit}</span><span style='color:#666'>{w_post_text}</span></html>"

            # 第 0 列：底本上下文 (使用 Label 渲染 HTML)
            lbl_base = QLabel(html_base)
            lbl_base.setAlignment(Qt.AlignmentFlag.AlignCenter)
            lbl_base.setStyleSheet("background-color: transparent;")
            self.table_main.setCellWidget(row, 0, lbl_base)
            
            # 第 1 列：校本上下文 (使用 Label 渲染 HTML)
            lbl_wit = QLabel(html_wit)
            lbl_wit.setAlignment(Qt.AlignmentFlag.AlignCenter)
            lbl_wit.setStyleSheet("background-color: transparent;")
            self.table_main.setCellWidget(row, 1, lbl_wit)

            # 第 2 列：類型 
            item_type = QTableWidgetItem(r['type'])
            item_type.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            if r['type'] == '訛': item_type.setBackground(QColor("#FADBD8")) # 淡紅
            elif r['type'] == '脫': item_type.setBackground(QColor("#D6EAF8")) # 淡藍
            elif r['type'] == '衍': item_type.setBackground(QColor("#D5F5E3")) # 淡綠
            elif r['type'] == '倒': item_type.setBackground(QColor("#FCF3CF")) # 淡黃
            elif r['type'] == '異': item_type.setBackground(QColor("#E8DAEF")) # 淡紫
            self.table_main.setItem(row, 2, item_type)
            item_type.setData(Qt.ItemDataRole.UserRole, r) # 綁定數據用於回查

            # 第 3 列：底本文字
            t_base = QTableWidgetItem(center_base)
            t_base.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            if center_base == "【無】": t_base.setForeground(QColor("#B74639"))
            self.table_main.setItem(row, 3, t_base)

            # 第 4 列：校本文字
            t_wit = QTableWidgetItem(center_wit)
            t_wit.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            if center_wit == "【無】": t_wit.setForeground(QColor("#B74639"))
            self.table_main.setItem(row, 4, t_wit)

        # 2. 渲染異體字表
        self.render_variants(variants, curr_wit_clean)
        
        # 3. 渲染原文預覽
        self.render_text_preview()
        # 4. 加載複雜訛誤板塊 (新增)
        self.complex_panel.load_data(wit_name, self.all_data, self.base_clean_text, curr_wit_clean)
        # [新增]：加載錯簡數據 (雖然可能被鎖定，但先加載無妨，或者判斷一下)
        if self.sub_tabs.isTabEnabled(3): 
             self.disp_panel.load_data(wit_name, self.all_data, self.base_clean_text, curr_wit_clean)

        # 5. 傳遞數據給新面板 (僅傳遞數據，不自動運行，節省性能)
        if hasattr(self, 'visual_panel'):
            self.visual_panel.load_data(wit_name, self.all_data, self.base_clean_text, curr_wit_clean)
            
        if hasattr(self, 'phonetic_panel'):
            self.phonetic_panel.load_data(wit_name, self.all_data, self.base_clean_text, curr_wit_clean)

    def render_variants(self, variants, w_clean):
        # 清空頻次按鈕
        while self.stats_lay.count():
            item = self.stats_lay.takeAt(0) # 取出第一個元素
            if item.widget():
                item.widget().deleteLater() # 如果是按鈕，刪除
        
        counts = Counter([(v['base_clean'], v['wit_clean']) for v in variants])
        
        # 創建頻次按鈕
        for pair, count in counts.items():
            _, _, desc = self.engine.get_char_attr(pair[0], pair[1])
            tag = f"{pair[0]}-{pair[1]}: {count}次" + (" (均可正)" if "均可作正" in desc else "")
            btn = QPushButton(tag)
            btn.setStyleSheet("text-align: left; padding: 5px;")
            btn.clicked.connect(lambda ch, p=pair: self.filter_var(p))
            self.stats_lay.addWidget(btn)
        self.stats_lay.addStretch()
        
        self.cur_variants = variants
        self.update_var_table(variants, w_clean, counts)

    def update_var_table(self, vs, w_clean, counts):
        self.table_var.setRowCount(0)
        for i, v in enumerate(vs):
            row = self.table_var.rowCount()
            self.table_var.insertRow(row)
            
            # 序號
            item_idx = QTableWidgetItem(str(i+1))
            item_idx.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            self.table_var.setItem(row, 0, item_idx)
            
            # --- 上下文計算 (前後各5字) ---
            idx, w_idx = v['idx'], v['wit_idx']
            
            # 底本上下文
            pre_base = self.base_clean_text[max(0, idx-5):idx]
            center_base = v['base_clean']
            post_base = self.base_clean_text[idx+1:idx+6]
            
            # 校本上下文
            pre_wit = w_clean[max(0, w_idx-5):w_idx]
            center_wit = v['wit_clean']
            post_wit = w_clean[w_idx+1:w_idx+6]
            
            # --- 構建 HTML 樣式 (應用主表的樣式) ---
            html_base = f"<html><span style='color:#666'>{pre_base}</span><span style='color:#B74639; font-weight:bold; font-size:14px;'>{center_base}</span><span style='color:#666'>{post_base}</span></html>"
            html_wit = f"<html><span style='color:#666'>{pre_wit}</span><span style='color:#B74639; font-weight:bold; font-size:14px;'>{center_wit}</span><span style='color:#666'>{post_wit}</span></html>"

            # 第 1 列：底本上下文 (HTML Label)
            lbl_base = QLabel(html_base)
            lbl_base.setAlignment(Qt.AlignmentFlag.AlignCenter)
            lbl_base.setStyleSheet("background-color: transparent;")
            self.table_var.setCellWidget(row, 1, lbl_base)

            # 第 2 列：底本文字
            item_base = QTableWidgetItem(center_base)
            item_base.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            self.table_var.setItem(row, 2, item_base)
            
            # 第 3 列：底本屬性
            a1, a2, desc = self.engine.get_char_attr(v['base_clean'], v['wit_clean'])
            item_a1 = QTableWidgetItem(a1)
            item_a1.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            # 如果屬性為空(雙重身份)，標記黃色背景
            if not a1: item_a1.setBackground(QColor("#FCF3CF")) 
            self.table_var.setItem(row, 3, item_a1)
            
            # 第 4 列：校本上下文 (HTML Label)
            lbl_wit = QLabel(html_wit)
            lbl_wit.setAlignment(Qt.AlignmentFlag.AlignCenter)
            lbl_wit.setStyleSheet("background-color: transparent;")
            self.table_var.setCellWidget(row, 4, lbl_wit)

            # 第 5 列：校本文字
            item_wit = QTableWidgetItem(center_wit)
            item_wit.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            self.table_var.setItem(row, 5, item_wit)
            
            # 第 6 列：校本屬性
            item_a2 = QTableWidgetItem(a2)
            item_a2.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            if not a2: item_a2.setBackground(QColor("#FCF3CF"))
            self.table_var.setItem(row, 6, item_a2)
            
            # 第 7 列：次數
            item_cnt = QTableWidgetItem(str(counts[(v['base_clean'], v['wit_clean'])]))
            item_cnt.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            self.table_var.setItem(row, 7, item_cnt)
            
            # 第 8 列：情況說明
            item_desc = QTableWidgetItem(desc)
            if "均可作正" in desc: 
                item_desc.setForeground(QColor("#B74639")) # 紅字警告
                item_desc.setFont(QFont("SimSun", 10, QFont.Weight.Bold))
            self.table_var.setItem(row, 8, item_desc)

    def filter_var(self, p):
        f = [v for v in self.cur_variants if v['base_clean']==p[0] and v['wit_clean']==p[1]]
        self.update_var_table(f, self.wit_cleans[self.combo_wits.currentText()], Counter([(v['base_clean'], v['wit_clean']) for v in self.cur_variants]))

    def sort_var_table(self):
        c = Counter([(v['base_clean'], v['wit_clean']) for v in self.cur_variants])
        s = sorted(self.cur_variants, key=lambda v: c[(v['base_clean'], v['wit_clean'])], reverse=True)
        self.update_var_table(s, self.wit_cleans[self.combo_wits.currentText()], c)

    def render_text_preview(self):
        """根據當前校本和勾選狀態，渲染原文標注"""
        # 1. 獲取當前數據
        wit_name = self.combo_wits.currentText()
        if not wit_name or wit_name not in self.all_data:
            self.text_preview.clear()
            return
            
        data = self.all_data[wit_name]
        # 合併主要結果和異體字結果
        # 如果沒勾選"區分訛與異體字"，data['variants'] 是空的，這裡邏輯依然成立
        all_records = data['results']
        # 按在文中的位置排序
        sorted_records = sorted(all_records, key=lambda x: x['idx'])

        self.text_preview.clear()
        cursor = self.text_preview.textCursor()
        
        # --- 定義樣式格式 ---
        # 訛誤：朱砂紅
        fmt_err_text = QTextCharFormat(); fmt_err_text.setForeground(QColor("#B74639")); fmt_err_text.setFontWeight(QFont.Weight.Bold)
        # 訛誤標注：靛藍
        fmt_err_note = QTextCharFormat(); fmt_err_note.setForeground(QColor("#3E5E79")); fmt_err_note.setFontPointSize(10)
        
        # 異體字：竹青色 (#2E8B57)
        fmt_var_text = QTextCharFormat(); fmt_var_text.setForeground(QColor("#2E8B57")); fmt_var_text.setFontWeight(QFont.Weight.Bold)
        # 異體標注：紫藤色 (#8E44AD)
        fmt_var_note = QTextCharFormat(); fmt_var_note.setForeground(QColor("#8E44AD")); fmt_var_note.setFontPointSize(10)
        
        fmt_norm = QTextCharFormat() # 普通黑色
        
        doc_cursor = 0 
        show_variants = self.cb_show_variants_in_text.isChecked()

        for r in sorted_records:
            b_clean_idx = r['idx']
            b_len = r['base_clean_len']
            
            # 計算真實文檔位置
            real_start = self.base_map[b_clean_idx]
            if b_len > 0:
                last_char_idx = self.base_map[b_clean_idx + b_len - 1]
                real_char_end = last_char_idx + 1
            else:
                real_char_end = real_start
            real_segment_end = self.base_map[b_clean_idx + b_len]
            
            # 插入跳過的普通文本
            skipped = self.base_full_text[doc_cursor : real_start]
            if skipped: cursor.insertText(skipped, fmt_norm)
            
            text_content = self.base_full_text[real_start : real_char_end]
            trailing_punct = self.base_full_text[real_char_end : real_segment_end]
            
            # --- 核心分支邏輯 ---
            if r['type'] == 'Match':
                cursor.insertText(text_content, fmt_norm)
            
            elif r['type'] == '異':
                if show_variants:
                    # 勾選顯示：綠字 + 紫色標注
                    if text_content: cursor.insertText(text_content, fmt_var_text)
                    note_text = r['wit_clean'] if r['wit_clean'] else "【無】"
                    cursor.insertText(f"[{r['type']}:{note_text}]", fmt_var_note)
                else:
                    # 未勾選：當作普通文字顯示 (隱藏標注)
                    cursor.insertText(text_content, fmt_norm)
            
            else:
                # 其他類型 (訛、脫、衍、倒)：紅字 + 藍色標注
                if text_content: cursor.insertText(text_content, fmt_err_text)
                note_text = r['wit_clean'] if r['wit_clean'] else "【無】"
                cursor.insertText(f"[{r['type']}:{note_text}]", fmt_err_note)
            
            # 插入尾部標點
            if trailing_punct: cursor.insertText(trailing_punct, fmt_norm)
            doc_cursor = real_segment_end
        
        # 插入剩餘文本
        cursor.insertText(self.base_full_text[doc_cursor:], fmt_norm)

   # ==========================================
    #  適配器：供預覽和導出使用
    # ==========================================
    def get_notes_data(self):
        """
        [適配版] 從 self.note_records 獲取供預覽和導出使用的字典格式
        Format: { pos_idx: [note_content, ...] }
        """
        if not hasattr(self, 'note_records'):
            return {}
            
        insertion_points = defaultdict(list)
        
        # 獲取過濾設置
        show_diff = self.cb_show_diff_notes.isChecked()
        
        for rec in self.note_records:
            # 簡單過濾邏輯：如果用戶選擇不顯示校勘記，且這是個潛在的異體字(非脫衍倒)，則跳過
            is_potential_variant = rec['type'] in ['異', '訛'] 
            
            # 如果用戶在頂部取消勾選了"顯示異體字校勘記"，且這條確實是異體字，則跳過
            if not show_diff and rec['type'] == '異':
                continue
                
            insertion_points[rec['pos_idx']].append(rec['content'])
            
        return insertion_points

    def preview_notes_html(self):
        if not self.all_data: return
        insertion_points = self.get_notes_data()
        html = "<html><body style='font-family:SimSun; font-size:16px; line-height:1.8; color:#333; background-color:#FFFEFA;'>"
        html += "<div style='margin-bottom:20px;'>"
        cursor = 0
        note_counter = 1
        page_notes = [] 
        while cursor <= len(self.base_full_text):
            if cursor in insertion_points:
                notes = insertion_points[cursor]
                nums = []
                for n_content in notes:
                    nums.append(note_counter)
                    page_notes.append((note_counter, n_content))
                    note_counter += 1
                nums_str = ",".join(map(str, nums))
                html += f"<sup style='color:#3E5E79; font-weight:bold; font-size:12px;'>{nums_str}</sup>"
            if cursor < len(self.base_full_text):
                html += self.base_full_text[cursor]
            cursor += 1
        html += "</div>"
        html += "<hr style='border:1px solid #D0D0D0; margin:20px 0;'>"
        html += "<div style='font-size:14px; color:#555;'>"
        for num, content in page_notes:
            html += f"<p style='margin:5px 0;'>{num}. {content}</p>"
        html += "</div></body></html>"
        self.notes_preview.setHtml(html)

    def export_excel(self):
        """
        [導出功能]：生成帶有富文本格式的 Excel 報表。
        [技術選型]：使用 `xlsxwriter` 引擎。
        [原因]：
        普通的 pandas `to_excel` 不支持在同一個單元格內混合字體（Rich String）。
        為了在 Excel 中正確顯示生僻字（使用 Ext-B 專用字體）同時保持其他文字為宋體，
        必須使用 `worksheet.write_rich_string` 進行底層寫入。
        """
        if not self.all_data: return
        path, _ = QFileDialog.getSaveFileName(self, "導出的Excel", "異文對照表.xlsx", "Excel (*.xlsx)")
        if not path: return
        
        try:
            # 必须指定 engine='xlsxwriter' 才能支持富文本写入
            with pd.ExcelWriter(path, engine='xlsxwriter') as writer:
                workbook = writer.book
                
                # 定义基础单元格格式 (垂直居中, 自动换行)
                fmt_basic = workbook.add_format({'font_name': '宋体', 'valign': 'vcenter', 'text_wrap': True})
                
                # ===========================================================
                # 【核心辅助函数】将字符串转为 Excel 富文本片段列表
                # 输入: "abc𠮷def"
                # 输出: ['abc', format_obj_hanamin, '𠮷', 'def']
                # ===========================================================
                def get_rich_segments(text):
                    if not text: return [text]
                    segments = []
                    current_font = None
                    current_buffer = []
                    
                    for char in text:
                        # 1. 逐字检测字体
                        font_name = self.get_font_for_char(char)
                        
                        # 2. 应用映射 (把 FSung-1 转为 全宋体-1)
                        if font_name:
                            font_name = self.font_map.get(font_name, font_name)
                        
                        # 3. 状态检测：字体是否改变？
                        # (None 代表默认字体/宋体)
                        if font_name != current_font:
                            # 结算缓冲区里的旧文字
                            if current_buffer:
                                if current_font:
                                    # 如果上一段是特殊字体，创建一个格式对象 (绿色加粗以示区别)
                                    fmt = workbook.add_format({'font_name': current_font, 'color': '#2E8B57', 'bold': True})
                                    segments.append(fmt)
                                segments.append("".join(current_buffer))
                            
                            # 切换到新状态
                            current_font = font_name
                            current_buffer = [char]
                        else:
                            current_buffer.append(char)
                            
                    # 4. 结算最后一段
                    if current_buffer:
                        if current_font:
                            fmt = workbook.add_format({'font_name': current_font, 'color': '#2E8B57', 'bold': True})
                            segments.append(fmt)
                        segments.append("".join(current_buffer))
                    
                    # 如果没有任何特殊字体，直接返回原字符串 (性能优化)
                    if len(segments) == 1 and isinstance(segments[0], str):
                        return segments[0]
                        
                    return segments

                # ===========================================================
                # 开始遍历数据并写入
                # ===========================================================
                for wit_name, data in self.all_data.items():
                    curr_wit_clean = self.wit_clean_texts.get(wit_name, "")
                    
                    # -------------------------------------------------------
                    # Sheet 1: 异文总表
                    # -------------------------------------------------------
                    rows = []
                    # 需要一个临时列表来存储原始文本，用于后续生成富文本
                    raw_data_cache = [] 
                    
                    diffs = [res for res in data['results'] if res['type'] != 'Match']
                    
                    for r in diffs:
                        # 筛选逻辑
                        if hasattr(self, 'type_filters') and r['type'] in self.type_filters:
                            if not self.type_filters[r['type']].isChecked():
                                continue
                        
                        idx = r['idx']
                        w_idx = r['wit_idx']
                        
                        pre = self.base_clean_text[max(0, idx-5):idx]
                        post = self.base_clean_text[idx+r['base_clean_len']:idx+r['base_clean_len']+5]
                        w_pre = curr_wit_clean[max(0, w_idx-5):w_idx]
                        w_post = curr_wit_clean[w_idx+r['wit_clean_len']:w_idx+r['wit_clean_len']+5]
                        
                        b_real = r['base_clean'] if r['base_clean'] else ""
                        w_real = r['wit_clean'] if r['wit_clean'] else "【无】"
                        
                        # 构造显示文本
                        base_full_str = f"{pre}【{b_real}】{post}"
                        wit_full_str = f"{w_pre}【{w_real}】{w_post}"
                        
                        # 存入列表 (用于 DataFrame 占位)
                        rows.append([base_full_str, wit_full_str, r['type'], b_real, w_real])
                        
                        # 存入缓存 (用于后续富文本覆盖)
                        raw_data_cache.append({
                            'A': base_full_str, # 底本原文
                            'B': wit_full_str,  # 校本原文
                            'D': b_real,        # 底本文字
                            'E': w_real         # 校本文字
                        })
                    
                    # 1. 先用 Pandas 写入基础数据和表头
                    df = pd.DataFrame(rows, columns=['底本原文', '校本原文', '类型', '底本文字', '校本文字'])
                    sheet_name = (wit_name[:20] + "_异文")
                    df.to_excel(writer, sheet_name=sheet_name, index=False)
                    
                    # 2. 获取 worksheet 对象，准备覆盖写入
                    worksheet = writer.sheets[sheet_name]
                    worksheet.set_column('A:E', 25, fmt_basic) # 设置列宽
                    
                    # 3. 逐行覆盖写入富文本 (Rich Text)
                    for i, raw_item in enumerate(raw_data_cache):
                        row_idx = i + 1 # 跳过表头(第0行)
                        
                        # A列: 底本原文
                        segs = get_rich_segments(raw_item['A'])
                        if isinstance(segs, list): 
                            worksheet.write_rich_string(row_idx, 0, *segs, fmt_basic)
                        
                        # B列: 校本原文
                        segs = get_rich_segments(raw_item['B'])
                        if isinstance(segs, list): 
                            worksheet.write_rich_string(row_idx, 1, *segs, fmt_basic)
                            
                        # D列: 底本文字
                        segs = get_rich_segments(raw_item['D'])
                        if isinstance(segs, list): 
                            worksheet.write_rich_string(row_idx, 3, *segs, fmt_basic)
                            
                        # E列: 校本文字
                        segs = get_rich_segments(raw_item['E'])
                        if isinstance(segs, list): 
                            worksheet.write_rich_string(row_idx, 4, *segs, fmt_basic)

                    # -------------------------------------------------------
                    # Sheet 2: 异体字专属表 (逻辑同上)
                    # -------------------------------------------------------
                    v_rows = []
                    v_raw_cache = []
                    
                    allow_variant = True
                    if hasattr(self, 'type_filters') and '异' in self.type_filters:
                        if not self.type_filters['异'].isChecked(): allow_variant = False
                    
                    if allow_variant:
                        variants = data['variants']
                        v_counts = Counter([(v['base_clean'], v['wit_clean']) for v in variants])
                        
                        for v in variants:
                            idx, w_idx = v['idx'], v['wit_idx']
                            b_ctx = self.base_clean_text[max(0, idx-5):idx] + f"【{v['base_clean']}】" + self.base_clean_text[idx+1:idx+6]
                            w_ctx = curr_wit_clean[max(0, w_idx-5):w_idx] + f"【{v['wit_clean']}】" + curr_wit_clean[w_idx+1:w_idx+6]
                            a1, a2, desc = self.engine.get_char_attr(v['base_clean'], v['wit_clean'])
                            count = v_counts[(v['base_clean'], v['wit_clean'])]
                            
                            v_rows.append([b_ctx, v['base_clean'], a1, w_ctx, v['wit_clean'], a2, count, desc])
                            
                            # 缓存需要富文本处理的列
                            v_raw_cache.append({
                                'A': b_ctx,          # 底本上下文
                                'B': v['base_clean'], # 底本文字
                                'D': w_ctx,          # 校本上下文
                                'E': v['wit_clean']   # 校本文字
                            })
                        
                        df_v = pd.DataFrame(v_rows, columns=['底本上下文', '底本文字', '底本属性', '校本上下文', '校本文字', '校本属性', '出现次数', '情况说明'])
                        v_sheet_name = (wit_name[:20] + "_异体")
                        df_v.to_excel(writer, sheet_name=v_sheet_name, index=False)
                        
                        # 覆盖写入 Sheet 2
                        ws_v = writer.sheets[v_sheet_name]
                        ws_v.set_column('A:H', 15, fmt_basic)
                        
                        for i, raw_item in enumerate(v_raw_cache):
                            row_idx = i + 1
                            # 覆盖 A, B, D, E 四列
                            s = get_rich_segments(raw_item['A']); 
                            if isinstance(s, list): ws_v.write_rich_string(row_idx, 0, *s, fmt_basic)
                            
                            s = get_rich_segments(raw_item['B']); 
                            if isinstance(s, list): ws_v.write_rich_string(row_idx, 1, *s, fmt_basic)
                            
                            s = get_rich_segments(raw_item['D']); 
                            if isinstance(s, list): ws_v.write_rich_string(row_idx, 3, *s, fmt_basic)
                            
                            s = get_rich_segments(raw_item['E']); 
                            if isinstance(s, list): ws_v.write_rich_string(row_idx, 4, *s, fmt_basic)
                    
            QMessageBox.information(self, "成功", "Excel 已導出！\n(生僻字應用全宋體，請自行安裝！)")
        except Exception as e:
            import traceback
            traceback.print_exc()
            QMessageBox.critical(self, "失败", str(e))

    def open_custom_dict(self):
        """打開自定義規則窗口"""
        dlg = CustomDictWindow(self.engine, self)
        dlg.exec()

    def export_annotated_word(self):
        wit_name = self.combo_wits.currentText()
        if not wit_name or not self.all_data: return
        path, _ = QFileDialog.getSaveFileName(self, "導出校勘標注文檔", f"基於{wit_name}的校勘標注", "Word (*.docx)")
        if not path: return
        
        try:
            data = self.all_data[wit_name]
            all_records = data['results'] + data['variants']
            results = sorted(all_records, key=lambda x: x['idx'])
            
            doc = Document()
            # 設置文檔默認字體
            style = doc.styles['Normal']
            style.font.name = 'Times New Roman'
            style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋體')
            
            p = doc.add_paragraph()
            doc_cursor = 0
            show_variants = self.cb_show_variants_in_text.isChecked()
            
            # 定義顏色常量
            COLOR_GREEN = DocxColor(46, 139, 87)   # 竹青 (異體字)
            COLOR_PURPLE = DocxColor(142, 68, 173) # 紫藤 (異體標注)
            COLOR_RED = DocxColor(183, 70, 57)     # 朱砂 (訛誤字)
            COLOR_BLUE = DocxColor(62, 94, 121)    # 靛藍 (訛誤標注)

            for r in results:
                b_clean_idx = r['idx']
                b_len = r['base_clean_len']
                
                # 計算真實文檔位置
                real_start = self.base_map[b_clean_idx]
                if b_len > 0:
                    last_char_idx = self.base_map[b_clean_idx + b_len - 1]
                    real_char_end = last_char_idx + 1
                else:
                    real_char_end = real_start
                real_segment_end = self.base_map[b_clean_idx + b_len]
                
                # 1. 寫入跳過的普通文本 (要防止普通文本里有生僻字)
                skipped = self.base_full_text[doc_cursor : real_start]
                self.insert_smart_text(p, skipped) 
                
                text_content = self.base_full_text[real_start : real_char_end]
                trailing_punct = self.base_full_text[real_char_end : real_segment_end]
                
                # 2. 寫入核心內容
                if r['type'] == 'Match':
                    self.insert_smart_text(p, text_content)
                
                elif r['type'] == '異':
                    if show_variants:
                        # 異體字本體：綠色
                        self.insert_smart_text(p, text_content, color_rgb=COLOR_GREEN)
                        
                        # 異體字標注：紫色
                        note_text = r['wit_clean'] if r['wit_clean'] else "【無】"
                        self.insert_smart_text(p, "[", color_rgb=COLOR_PURPLE)
                        self.insert_smart_text(p, f"{r['type']}:", color_rgb=COLOR_PURPLE) # 類型
                        self.insert_smart_text(p, note_text, color_rgb=COLOR_PURPLE)       # 內容
                        self.insert_smart_text(p, "]", color_rgb=COLOR_PURPLE)
                    else:
                        # 未勾選則顯示為普通文字
                        self.insert_smart_text(p, text_content)
                
                else:
                    # 其他訛誤：紅色文字
                    self.insert_smart_text(p, text_content, color_rgb=COLOR_RED)
                    
                    # 訛誤標注：藍色
                    note_text = r['wit_clean'] if r['wit_clean'] else "【無】"
                    self.insert_smart_text(p, "[", color_rgb=COLOR_BLUE)
                    self.insert_smart_text(p, f"{r['type']}:", color_rgb=COLOR_BLUE)
                    self.insert_smart_text(p, note_text, color_rgb=COLOR_BLUE)
                    self.insert_smart_text(p, "]", color_rgb=COLOR_BLUE)
                
                # 3. 寫入尾部標點
                self.insert_smart_text(p, trailing_punct)
                doc_cursor = real_segment_end
            
            # 4. 寫入最後剩餘的文本
            self.insert_smart_text(p, self.base_full_text[doc_cursor:])
            
            doc.save(path)
            
            # --- 嘗試調用 Word 進行字體嵌入---
            if HAS_PYWIN32:
                try:
                    pythoncom.CoInitialize()
                    word_app = win32.gencache.EnsureDispatch('Word.Application')
                    word_app.Visible = False; word_app.DisplayAlerts = False
                    abs_path = os.path.abspath(path)
                    wdoc = word_app.Documents.Open(abs_path)
                    wdoc.EmbedTrueTypeFonts = True
                    wdoc.Save(); wdoc.Close(); word_app.Quit()
                    pythoncom.CoUninitialize()
                except: pass
            QMessageBox.information(self, "成功", "校勘標注文檔已導出！")
        except Exception as e:
             QMessageBox.critical(self, "失敗", str(e))

    # ==========================================
    # 校勘記管理核心邏輯 
    # ==========================================
    def generate_note_content(self, note_type, base_char, wit_char, wit_names, prev_char=""):
        """
        [邏輯說明]：根據校勘類型和相關文字，動態生成符合古籍整理規範的校勘記文本。
        :param note_type: 類型 (訛/脫/衍/倒/異)
        :param base_char: 底本文字
        :param wit_char: 校本文字 (若是脫，此字段可能為空)
        :param wit_names: 校本名稱列表
        :return: 格式化後的校勘記字符串
        """
        # 1. 處理校本名稱：確保名稱後有"本"字，並處理多校本合併
        clean_names = []
        for w in wit_names:
            # 先去掉可能存在的"本"字，再統一加上
            name_root = w.replace("本", "")
            clean_names.append(name_root)
        
        if len(clean_names) > 1:
            # 多個校本：某、某本皆...
            combined_name = "、".join(clean_names) + "本皆"
        else:
            # 單個校本：某本...
            combined_name = clean_names[0] + "本"

        # 2. 處理缺失的校本字 (針對手動切換類型時可能出現的情況)
        final_wit_char = wit_char if wit_char else "【需補充】"

        # 3. 根據類型應用模板
        if note_type == '訛':
            return f"{base_char}，{combined_name}作{final_wit_char}。"
        
        elif note_type == '異':
            return f"{base_char}，{combined_name}作{final_wit_char}。"
        
        elif note_type == '脫':
            return f"{base_char}，{combined_name}脫。"
        
        elif note_type == '倒':
            return f"{base_char}，{combined_name}倒乙。"
        
        elif note_type == '衍':
            # 衍文邏輯：判斷是否有前字
            if prev_char:
                return f"{prev_char}字下{combined_name}衍{final_wit_char}。"
            else:
                return f"卷首{combined_name}衍{final_wit_char}。"
        
        # 默認兜底格式
        return f"{base_char}，{combined_name}{note_type}{final_wit_char}。"
    
    # ==========================================
    #  重構：數據初始化 
    # ==========================================
    def init_note_records(self):
        """
        [增強版] 初始化校勘記數據
        新增功能：記錄底本文字的起始位置 (start_idx)，用於預覽區原文高亮
        """
        if not self.all_data: return

        combined_errors = defaultdict(list)
        for wit, data in self.all_data.items():
            all_res = data['results'] # 只讀取 results，避免重複
            
            for r in all_res:
                if r['type'] != 'Match':
                    key = (r['idx'], r['base_clean_len'])
                    combined_errors[key].append({'wit_fn': wit, **r})

        self.note_records = []
        record_id = 1
        sorted_keys = sorted(combined_errors.keys())

        for (idx, base_len) in sorted_keys:
            errors = combined_errors[(idx, base_len)]
            
            base_char = errors[0]['base_clean']
            prev_char = ""
            if idx > 0 and idx <= len(self.base_clean_text):
                 prev_char = self.base_clean_text[idx - 1]

            # --- 【關鍵修改】計算原文的 起始 和 結束 位置 ---
            if base_len == 0:
                # 衍文：底本無字，無法高亮原文
                real_start_idx = -1 
                if idx == 0: real_end_idx = 0
                else: real_end_idx = self.base_map[idx - 1] + 1
            else:
                # 訛/異/脫/倒：底本有字，記錄範圍
                real_start_idx = self.base_map[idx]
                real_end_idx = self.base_map[idx + base_len - 1] + 1

            grouped = defaultdict(list)
            for e in errors:
                key = (e['type'], e['wit_clean'])
                grouped[key].append(e['wit_fn'])

            for (err_type, wit_text), wits in grouped.items():
                content = self.generate_note_content(err_type, base_char, wit_text, wits, prev_char)

                record = {
                    'id': record_id,
                    'pos_idx': real_end_idx,   # 插入位置 (校勘記放這裡)
                    'start_idx': real_start_idx, # 原文起始位置 (用於原文高亮)
                    'end_idx': real_end_idx,     # 原文結束位置
                    'type': err_type,
                    'content': content,
                    'base_char': base_char,
                    'wit_char': wit_text,
                    'wit_names': wits,
                    'prev_char': prev_char,
                    'is_manual': False,
                    'display_idx': record_id
                }
                self.note_records.append(record)
                record_id += 1

        self.update_note_table()
        self.refresh_note_preview()

    def on_display_option_changed(self):
        """當顯示選項改變時，同時刷新表格隱藏狀態和預覽區"""
        self.update_note_table()     # 刷新表格 (隱藏/顯示行)
        self.refresh_note_preview()  # 刷新左側預覽

    # ==========================================
    #  UI刷新與交互邏輯 (支持過濾隱藏)
    # ==========================================
    def update_note_table(self):
        """將 self.note_records 渲染到右側表格 (含隱藏邏輯)"""
        self.table_notes.blockSignals(True)
        self.table_notes.setRowCount(0)
        
        # 獲取當前過濾開關狀態
        show_diff = self.cb_show_diff_notes.isChecked()
        
        for row, rec in enumerate(self.note_records):
            self.table_notes.insertRow(row)
            
            # ID
            item_id = QTableWidgetItem(str(rec['id']))
            item_id.setFlags(item_id.flags() ^ Qt.ItemFlag.ItemIsEditable)
            item_id.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            self.table_notes.setItem(row, 0, item_id)
            
            # 位置
            item_pos = QTableWidgetItem(str(rec['pos_idx']))
            item_pos.setFlags(item_pos.flags() ^ Qt.ItemFlag.ItemIsEditable)
            item_pos.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            self.table_notes.setItem(row, 1, item_pos)
            
            # 類型 (下拉框)
            combo = QComboBox()
            combo.addItems(['訛', '脫', '衍', '倒', '異'])
            combo.setCurrentText(rec['type'])
            combo.currentIndexChanged.connect(lambda idx, r=row, c=combo: self.on_note_type_changed(r, c.currentText()))
            self.table_notes.setCellWidget(row, 2, combo)
            
            # 內容
            item_content = QTableWidgetItem(rec['content'])
            self.table_notes.setItem(row, 3, item_content)
            
            # 綁定數據
            item_id.setData(Qt.ItemDataRole.UserRole, rec)
            
            # --- 【關鍵修改】根據類型和勾選狀態隱藏行 ---
            # 如果用戶不看異體字，且當前條目是異體字，則隱藏該行
            if not show_diff and rec['type'] == '異':
                self.table_notes.setRowHidden(row, True)
            else:
                self.table_notes.setRowHidden(row, False)

        self.table_notes.blockSignals(False)

# --- 信號處理 ---
    def on_note_type_changed(self, row, new_type):
        """當下拉框類型改變時"""
        if row >= len(self.note_records): return
        old_type = self.note_records[row]['type']
        
        if new_type != old_type:
            # 推送命令到撤銷棧
            cmd = NoteTypeChangeCommand(self, row, new_type)
            self.undo_stack.push(cmd)

    def on_note_content_changed(self, item):
        """當表格內容被手動修改時"""
        if item.column() == 3: # 只處理內容列
            row = item.row()
            new_text = item.text()
            old_text = self.note_records[row]['content']
            
            if new_text != old_text:
                # 推送命令到撤銷棧
                cmd = NoteContentEditCommand(self, row, new_text)
                self.undo_stack.push(cmd)

    # --- 界面刷新輔助 ---
    def refresh_ui_row(self, row):
        """刷新指定行的界面元素 (用於 Undo/Redo)"""
        if row >= len(self.note_records): return
        rec = self.note_records[row]
        
        self.table_notes.blockSignals(True)
        # 更新下拉框
        widget = self.table_notes.cellWidget(row, 2)
        if isinstance(widget, QComboBox):
            widget.setCurrentText(rec['type'])
        # 更新文本
        self.table_notes.item(row, 3).setText(rec['content'])
        self.table_notes.blockSignals(False)
        
        self.refresh_note_preview()

    def refresh_preview_only(self):
        """只刷新預覽 (用於手動編輯時)"""
        self.refresh_note_preview()

    # --- 底層增刪邏輯 (供 Command 調用) ---
    def delete_note_data(self, row):
        """底層刪除"""
        del self.note_records[row]
        self.update_note_table() # 刪除涉及行號變化，必須重繪整個表
        self.refresh_note_preview()

    def insert_note_data(self, row, record):
        """底層插入"""
        self.note_records.insert(row, record)
        self.update_note_table()
        self.refresh_note_preview()
        
    # --- 右鍵刪除菜單適配 ---
    def delete_note(self, row_index):
        """右鍵刪除現在也走 Undo Stack"""
        cmd = NoteDeleteCommand(self, row_index)
        self.undo_stack.push(cmd)

    def show_note_context_menu(self, pos):
        """
        [修改] 校勘記表格右鍵菜單
        功能：集成「刪除」與「定位圖版」功能
        """
        # 1. 獲取被點擊的行
        item = self.table_notes.itemAt(pos)
        if not item: return
        
        row = item.row()
        # 獲取該行對應的完整數據記錄 (rec)
        rec = self.note_records[row]
        
        # 2. 創建菜單
        menu = QMenu()
        
        # === [新增部分 A] 添加定位選項 ===
        action_locate = menu.addAction("🔍 定位圖版 (顯示底本)")
        menu.addSeparator() # 添加一條分隔線
        # ===============================
        
        # [原有邏輯] 添加刪除選項
        action_del_one = menu.addAction("刪除此條校勘記")
        
        # [原有邏輯] 判斷是否顯示批量刪除 (非脫衍倒)
        note_content = rec['content']
        is_likely_variant = not any(x in note_content for x in ['衍', '脫', '倒'])
        action_del_batch = None
        if is_likely_variant:
            action_del_batch = menu.addAction(f"刪除所有內容為「{note_content}」的校勘記")
            
        # 3. 顯示菜單並等待用戶點擊
        action = menu.exec(self.table_notes.viewport().mapToGlobal(pos))
        
        # 4. 處理點擊結果
        if not action: return # 用戶取消了點擊

        # === [新增部分 B] 處理定位邏輯 ===
        if action == action_locate:
            # 確保有底本映射器，且該記錄有有效的原文索引
            if self.base_mapper and rec['start_idx'] != -1:
                # 調用映射器獲取圖片
                name, img = self.base_mapper.get_image_source(rec['start_idx'])
                if name:
                    # 在左側 Dock 中顯示圖片
                    self.evidence_dock.show_image('base', name, img)
                    self.evidence_dock.show() # 確保窗口展開
            else:
                QMessageBox.warning(self, "提示", "無法定位：未關聯圖版或該條目為衍文（無底本字）。")
        # ===============================

        # [原有邏輯] 處理刪除
        elif action == action_del_one:
            self.delete_note(row)
            
        elif action_del_batch and action == action_del_batch:
            self.delete_notes_batch(note_content)



    def delete_notes_batch(self, target_content):
        """批量刪除內容相同的"""
        original_count = len(self.note_records)
        self.note_records = [r for r in self.note_records if r['content'] != target_content]
        deleted_count = original_count - len(self.note_records)
        
        self.update_note_table()
        self.refresh_note_preview()
        QMessageBox.information(self, "刪除成功", f"已批量刪除 {deleted_count} 條內容為\n「{target_content}」的校勘記。")

   # ==========================================
    #  預覽渲染 (支持原文+校勘記 雙重高亮)
    # ==========================================
    def refresh_note_preview(self):
        """
        [可視化功能]：預覽渲染引擎。
        [特性]：實現了「雙重高亮」。不僅在腳注中顯示內容，
        還會根據 `start_idx` 和 `end_idx` 在原文中將對應的異文用紅色高亮標出。
        這為校對者提供了直觀的視覺反饋，快速定位異文在文中的位置。
        """
        if not hasattr(self, 'note_records'): return
        
        show_diff = self.cb_show_diff_notes.isChecked()
        highlight_red = self.cb_highlight_diff.isChecked()
        
        display_map = defaultdict(list)
        
        # 用於存儲需要變紅的原文索引集合
        highlight_indices = set()
        
        current_num = 1
        
        for rec in self.note_records:
            # 1. 過濾邏輯
            if not show_diff and rec['type'] == '異':
                continue
                
            # 2. 收集需要高亮的原文索引
            # 條件：開啟高亮 + 類型是異 + 確實有底本文字(不是衍文)
            if highlight_red and rec['type'] == '異' and rec['start_idx'] != -1:
                for i in range(rec['start_idx'], rec['end_idx']):
                    highlight_indices.add(i)

            # 3. 處理校勘記內容高亮
            content = rec['content']
            display_text = content
            if highlight_red and rec['type'] == '異':
                 display_text = f"<span style='color:#B74639; font-weight:bold;'>{content}</span>"
            
            display_map[rec['pos_idx']].append((current_num, display_text))
            current_num += 1

        # --- 生成 HTML ---
        html = "<html><body style='font-family:SimSun; font-size:16px; line-height:1.8; color:#333; background-color:#FFFEFA;'>"
        html += "<div style='margin-bottom:20px;'>"
        
        cursor = 0
        page_notes_html = []
        
        while cursor <= len(self.base_full_text):
            # 1. 插入上標 (校勘記標號)
            if cursor in display_map:
                notes = display_map[cursor]
                nums = [str(n[0]) for n in notes]
                nums_str = ",".join(nums)
                html += f"<sup style='color:#3E5E79; font-weight:bold; font-size:12px;'>{nums_str}</sup>"
                # 收集底部注釋內容
                for n_num, n_text in notes:
                    page_notes_html.append(f"<p style='margin:5px 0;'>{n_num}. {n_text}</p>")

            # 2. 插入原文文字 (含高亮處理)
            if cursor < len(self.base_full_text):
                char = self.base_full_text[cursor]
                
                # 【核心邏輯】如果當前下標在需要高亮的集合裡，包裹紅色樣式
                if cursor in highlight_indices:
                    html += f"<span style='color:#B74639; font-weight:bold;'>{char}</span>"
                else:
                    html += char
            
            cursor += 1
            
        html += "</div>"
        html += "<hr style='border:1px solid #D0D0D0; margin:20px 0;'>"
        html += "<div style='font-size:14px; color:#555;'>"
        html += "".join(page_notes_html)
        html += "</div></body></html>"
        
        self.notes_preview.setHtml(html)

    # ==========================================
    #  導出功能 (支持精準過濾)
    # ==========================================
    def export_real_footnote(self):
        """
        [導出功能]：生成帶有Word 腳注 (Footnotes) 的校勘文檔。
        [邏輯]：
        調用後台線程 WordExportThread，傳入經過用戶在前端過濾、編輯後的校勘記數據。
        此功能依賴 Windows COM 接口，確保導出的文檔格式專業、兼容性強。
        """
        if not HAS_PYWIN32:
            QMessageBox.critical(self, "缺少插件", "您還沒有安裝 pywin32！")
            return
        
        # 1. 準備數據：過濾掉用戶隱藏的內容
        show_diff = self.cb_show_diff_notes.isChecked()
        final_insertion_points = defaultdict(list)
        
        for rec in self.note_records:
            # 直接使用 type 字段進行精準過濾
            # 如果未勾選顯示，且類型為'異'，則跳過不導出
            if not show_diff and rec['type'] == '異':
                continue
                
            final_insertion_points[rec['pos_idx']].append(rec['content'])

        if not final_insertion_points:
            QMessageBox.information(self, "提示", "當前沒有可導出的校勘記內容。")
            return

        path, _ = QFileDialog.getSaveFileName(self, "導出原文附校勘記", "底本附校勘記_修訂版.docx", "Word (*.docx)")
        if not path: return
        
        self.btn_export_word_note.setEnabled(False)
        self.lbl_status.setText("正在導出修訂後的校勘記...")
        
        self.word_thread = WordExportThread(path, self.base_full_text, final_insertion_points, self.available_font_families, self.font_map)
        self.word_thread.progress_signal.connect(self.update_progress)
        self.word_thread.status_signal.connect(lambda s: self.lbl_status.setText(s))
        self.word_thread.finished_signal.connect(self.on_word_export_finished)
        self.word_thread.error_signal.connect(self.on_word_export_error)
        self.word_thread.start()

    def update_progress(self, val):
        self.progress_bar.setValue(val)

    def on_word_export_finished(self, file_path):
        self.progress_bar.setVisible(False)
        self.btn_export_word_note.setEnabled(True)
        self.lbl_status.setText("導出完成")
        QMessageBox.information(self, "成功", f"底本附校勘記已生成！！\n保存在：{file_path}")

    def on_word_export_error(self, err_msg):
        self.progress_bar.setVisible(False)
        self.btn_export_word_note.setEnabled(True)
        self.lbl_status.setText("導出失敗")
        QMessageBox.critical(self, "失敗", f"Word 遙控失敗：{err_msg}\n請確認 Word 未被其他程序獨佔。")

    def reload_all_views(self):
        """當數據被 ComplexAuditPanel 修改後，刷新所有界面"""
        # 1. 刷新統計面板
        if hasattr(self, 'stats_panel'):
            self.stats_panel.update_data(self.all_data, self.cb_variant_filter.isChecked(), self.base_clean_text)
        
        # 2. 重新生成校勘記 (這會更新 self.note_records)
        self.init_note_records()
        self.preview_notes_html()
        
        # 3. 刷新當前視圖 (異文表、異體表)
        # 獲取當前校本索引，如果是 -1 (沒選) 則不刷
        idx = self.combo_wits.currentIndex()
        if idx >= 0:
            self.switch_view(idx)

    def eventFilter(self, source, event):
        """底層事件過濾：攔截對禁用 Tab 的點擊"""
        if source == self.sub_tabs.tabBar() and event.type() == QEvent.Type.MouseButtonPress:
            # 獲取點擊位置對應的 Tab 索引
            # 注意：PyQt6 使用 position().toPoint()
            tab_index = self.sub_tabs.tabBar().tabAt(event.position().toPoint())
            
            if tab_index != -1 and not self.sub_tabs.isTabEnabled(tab_index):
                # 如果點擊了禁用 Tab，手動觸發提示
                self.on_tab_clicked(tab_index)
                return True # 攔截事件，不讓它繼續傳遞
                
        return super().eventFilter(source, event)

    def on_tab_clicked(self, index):
        """
        [新增交互]：當點擊 Tab 時觸發
        如果點擊了被鎖定的「錯簡匹配」Tab，彈出提示框。
        """
        # 檢查該 Tab 是否被禁用 (Locked)
        if not self.sub_tabs.isTabEnabled(index):
            # 獲取 Tab 標題，確認是錯簡模塊
            title = self.sub_tabs.tabText(index)
            if "錯簡匹配 🔒" in title:
                QMessageBox.warning(self, "功能鎖定", "請先清空「複雜訛誤審覈」列表以解鎖此功能。")

    def toggle_displacement_tab(self, unlocked):
        """接收信號，控制錯簡 Tab 的鎖定/解鎖"""
        TAB_INDEX = 3 # 確保索引正確
        
        # 只有當狀態真正改變（從鎖定變解鎖）時，才考慮彈窗
        was_locked = not self.sub_tabs.isTabEnabled(TAB_INDEX)
        
        self.sub_tabs.setTabEnabled(TAB_INDEX, unlocked)
        
        if unlocked:
            self.sub_tabs.setTabText(TAB_INDEX, "錯簡匹配")
            self.sub_tabs.setTabToolTip(TAB_INDEX, "") 
            
            # 自動加載數據
            wit_name = self.combo_wits.currentText()
            curr_wit_clean = self.wit_clean_texts.get(wit_name, "")
            # [修改]：這裡調用 load_data 時傳入四個參數
            if wit_name:
                self.disp_panel.load_data(wit_name, self.all_data, self.base_clean_text, curr_wit_clean)
            
            # [核心修改]：僅當「之前是鎖定狀態」且「沒有被抑制」時，才彈窗
            if was_locked and not self.suppress_unlock_popup:
                QMessageBox.information(self, "審覈階段完成", "複雜訛誤審覈已完成！\n\n「錯簡匹配」功能已解鎖。")
        else:
            self.sub_tabs.setTabText(TAB_INDEX, "錯簡匹配 🔒")
            self.sub_tabs.setTabToolTip(TAB_INDEX, "")

    def on_diff_table_double_click(self, row, col):
        """
        [交互核心] 雙擊異文表定位圖版
        """
        # 0=底本上下文, 1=校本上下文, 2=類型, 3=底本字, 4=校本字
        if col not in [0, 1, 3, 4]: return
        
        item_type = self.table_main.item(row, 2)
        if not item_type: return
        
        record = self.type_filters # 這裡之前的代碼邏輯是實時過濾，可能行號對不上
        
        r = item_type.data(Qt.ItemDataRole.UserRole)
        if not r: return
        
        # A. 定位底本
        if self.base_mapper:
            # 獲取 Raw Index (通過 clean_and_map 生成的映射表)
            clean_idx = r['idx']
            # 防止越界
            if clean_idx < len(self.base_map):
                raw_idx = self.base_map[clean_idx]
                name, img = self.base_mapper.get_image_source(raw_idx)
                if name:
                    self.evidence_dock.show_image('base', name, img)
        
        # B. 定位校本
        wit_name = self.combo_wits.currentText()
        if wit_name in self.wit_mappers:
            mapper = self.wit_mappers[wit_name]
            # 獲取該校本的映射表
            if wit_name in self.wit_maps:
                w_map = self.wit_maps[wit_name]
                clean_idx = r['wit_idx']
                if clean_idx < len(w_map):
                    raw_idx = w_map[clean_idx]
                    name, img = mapper.get_image_source(raw_idx)
                    if name:
                        self.evidence_dock.show_image('wit', name, img)
                        
        self.evidence_dock.show()

    def show_text_preview_menu(self, pos):
        """
        [新增] 原文預覽區的右鍵菜單：選中文字後定位底本圖片
        """
        # 創建標準菜單
        menu = self.text_preview.createStandardContextMenu()
        menu.addSeparator()
        action_locate = menu.addAction("🔍 定位底本圖版")
        
        # 顯示菜單
        action = menu.exec(self.text_preview.mapToGlobal(pos))
        
        if action == action_locate:
            cursor = self.text_preview.textCursor()
            selected_text = cursor.selectedText()
            
            if not selected_text: 
                QMessageBox.information(self, "提示", "請先選中一段文字。")
                return
            
            # 1. 在清洗後的底本中查找選中文字的位置
            # 注意：這裡只查找第一個匹配項作為演示。如果文字重複，可能定位不準，但在校勘場景通常足夠。
            clean_idx = self.base_clean_text.find(selected_text)
            
            if clean_idx != -1 and self.base_mapper:
                # 2. 將 清洗後索引(Clean Index) 轉換為 原始索引(Raw Index)
                # self.base_map 是我們在導入時生成的映射表
                if clean_idx < len(self.base_map):
                    raw_idx = self.base_map[clean_idx]
                    
                    # 3. 獲取圖片
                    name, img = self.base_mapper.get_image_source(raw_idx)
                    if name:
                        self.evidence_dock.show_image('base', name, img)
                        self.evidence_dock.show()
            else:
                QMessageBox.warning(self, "失敗", "無法在底本中定位該文字，或未加載底本圖版。")

    # === 【新增方法：右鍵菜單】 ===
    def show_main_table_context_menu(self, pos):
        item = self.table_main.itemAt(pos)
        if not item: return
        row = item.row()
        
        # 獲取底本和校本文字 (假設3是底本字列, 4是校本字列，請根據您實際表格確認)
        b_char = self.table_main.item(row, 3).text()
        w_char = self.table_main.item(row, 4).text()
        if not b_char or not w_char: return
        
        menu = QMenu()
        action_add_var = menu.addAction(f"設為「{b_char}={w_char}」為臨時異體")
        action_add_ex = menu.addAction(f"設為「{b_char}!={w_char}」為強制訛誤")
        
        action = menu.exec(self.table_main.viewport().mapToGlobal(pos))
        
        if action == action_add_var:
            self.add_custom_rule(b_char, w_char, is_variant=True)
        elif action == action_add_ex:
            self.add_custom_rule(b_char, w_char, is_variant=False)

    # === 【新增方法：處理規則衝突並自動更新】 ===
    def add_custom_rule(self, a, b, is_variant=True):
        engine = self.engine
        msg = ""
        
        if is_variant:
            # 自動刪除對立面 (屏蔽規則)
            if (a, b) in engine.custom_excludes: engine.custom_excludes.remove((a, b))
            if (b, a) in engine.custom_excludes: engine.custom_excludes.remove((b, a))
            # 添加
            engine.custom_variants[a] = b
            engine.custom_variants[b] = a
            msg = f"已設置：{a}={b} (異體)\n舊的衝突規則已自動清理。"
        else:
            # 自動刪除對立面 (異體規則)
            if engine.custom_variants.get(a) == b: del engine.custom_variants[a]
            if engine.custom_variants.get(b) == a: del engine.custom_variants[b]
            # 添加
            engine.custom_excludes.add((a, b))
            engine.custom_excludes.add((b, a))
            msg = f"已設置：{a}!={b} (強制訛誤)\n舊的衝突規則已自動清理。"
            
        QMessageBox.information(self, "規則更新", f"{msg}\n請重新運行校勘以生效。")


if __name__ == "__main__":
# 1. [新增] 註冊 MDict URL Scheme (必須在 QApplication 之前)
    scheme = QWebEngineUrlScheme(b"mdict")
    scheme.setSyntax(QWebEngineUrlScheme.Syntax.Path)
    
    # === 【修改點在這裏】 ===
    # 將 QWebEngineUrlScheme.PortUnspecified 改為 -1
    scheme.setDefaultPort(-1) 
    # ======================
    
    scheme.setFlags(QWebEngineUrlScheme.Flag.CorsEnabled | QWebEngineUrlScheme.Flag.LocalAccessAllowed)
    QWebEngineUrlScheme.registerScheme(scheme)

    # 1. 【新增】修復 Windows 任務欄圖標顯示為 Python 默認圖標的問題
    # 設置 AppUserModelID，讓 Windows 認為這是一個獨立的程序
    if sys.platform == 'win32':
        myappid = 'mycompany.collation_tool.version.1.0' # 任意唯一的字符串
        ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID(myappid)

    app = QApplication(sys.argv)
    
    # 2. 【新增】設置應用程序級別的圖標 (任務欄圖標)
    icon_path = get_resource_path(os.path.join("resources", "logo.ico"))
    if os.path.exists(icon_path):
        app_icon = QIcon(icon_path)
        app.setWindowIcon(app_icon) # 設置全局圖標
    else:
        print(f"警告：找不到圖標文件 -> {icon_path}")

    window = MainWindow()

    if os.path.exists(icon_path):
        window.setWindowIcon(QIcon(icon_path))

    window.show()
    sys.exit(app.exec())